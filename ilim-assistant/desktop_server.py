"""
RÜZGAR masaüstü kabuğu için yerel HTTP API (Electron).

Çalıştırma (ilim-assistant klasöründe):
  python desktop_server.py

Varsayılan: http://127.0.0.1:8779 (RUZGAR_API_PORT / .env)
"""

from __future__ import annotations

import json
import os
import shutil
import subprocess
import sys
import tempfile
import uuid
import zipfile
import re
from html import unescape
from urllib.parse import urlparse
from urllib.request import urlopen, Request
import xml.etree.ElementTree as ET
from pathlib import Path
from typing import Annotated, Any, Iterator

import uvicorn
from fastapi import FastAPI, File, Form, HTTPException, Query, UploadFile, WebSocket
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, StreamingResponse

from pydantic import BaseModel, Field
from starlette.requests import Request
from starlette.concurrency import iterate_in_threadpool, run_in_threadpool

from ilim_assistant.chat_core import (
    message_calls_wake_name,
    normalize_mode,
    prior_messages_for_turn,
    prepare_turn,
    rag_footer,
)
from ilim_assistant.stream_orchestra import prefetch_main_engine_bundle_for_stream
from ilim_assistant.hafiza_i_ruzgar import (
    genel_hafiza_lookup,
    get_hafiza_motor as _get_hafiza_motor,
)
from ilim_assistant.text_encoding import finalize_assistant_reply
from ilim_assistant.llm_brain import select_brain_chain, stream_chat_with_brain
from ilim_assistant.llm_ollama import format_llm_user_error
from ilim_assistant.motorlar.programlama_motoru import (
    apply_assistant_reply_tools,
    code_debug_max_retries,
    wants_autonomous_code_debug,
)
from ilim_assistant.stt_whisper import stt_runtime_available, transcribe_file
from ilim_assistant.video_ffmpeg import (
    MAX_SEGMENT_SECONDS,
    burn_subtitles_into_mp4,
    concat_two_files,
    export_directory,
    ffmpeg_available,
    ffprobe_available,
    ffprobe_json,
    mux_replace_audio,
    summarize_probe,
    transcode_to_mp4,
    trim_media,
)

app = FastAPI(title="RÜZGAR Desktop API")

# Video probe: tarayıcıdan gelen dosya üst sınırı (metadata okuma; yine de güvenlik)
MAX_VIDEO_PROBE_BYTES = 600 * 1024 * 1024

# Proje kökü (YAPAY ZEKA): ilim-assistant/desktop_server.py → iki üst dizin
_ILIM_ASSISTANT_ROOT = Path(__file__).resolve().parent
REPO_ROOT = _ILIM_ASSISTANT_ROOT.parent
_SKIP_LIST_NAMES = frozenset({"node_modules", "__pycache__", ".venv", "venv", ".git"})


def _load_env_file() -> None:
    """Mevcut .env / RUZGAR_BRAIN.env / google_api_key.txt — GOOGLE_GEMINI_API_KEY vb."""
    try:
        from ilim_assistant.env_bootstrap import ensure_ruzgar_env

        ensure_ruzgar_env()
    except Exception:
        path = _ILIM_ASSISTANT_ROOT / ".env"
        if not path.is_file():
            return
        try:
            for raw in path.read_text(encoding="utf-8").splitlines():
                line = raw.strip()
                if not line or line.startswith("#") or "=" not in line:
                    continue
                key, _, val = line.partition("=")
                key = key.strip()
                val = val.strip().strip('"').strip("'")
                if key and key not in os.environ:
                    os.environ[key] = val
        except OSError:
            pass


_load_env_file()
try:
    from ilim_assistant.config import apply_global_api_key_to_runtime, gemini_ready

    apply_global_api_key_to_runtime()
    if gemini_ready():
        print("[RÜZGAR] Bulut: Gemini anahtarı yüklü.", file=sys.stderr)
    else:
        print(
            "[RÜZGAR] UYARI: Gemini anahtarı yok — ilim-assistant/RUZGAR_BRAIN.env kontrol edin.",
            file=sys.stderr,
        )
except Exception:
    pass


def _gemini_startup_warmup() -> None:
    """GLOBAL_API_KEY — isteğe bağlı ping; daemon kapalıysa ağ çağrısı yapma (hızlı açılış)."""
    try:
        from ilim_assistant.config import (
            apply_global_api_key_to_runtime,
            gemini_disabled,
            gemini_ready,
            ollama_only_mode,
            suppress_cloud_runtime_keys,
        )

        suppress_cloud_runtime_keys()
        if ollama_only_mode() or gemini_disabled():
            print(
                "[RÜZGAR] Yerel Ollama modu — Gemini/Groq kapalı.",
                file=sys.stderr,
            )
            return
        from ilim_assistant.gemini_daemon import _daemon_enabled, daemon_status, start_gemini_daemon

        apply_global_api_key_to_runtime()
        if not gemini_ready():
            print(
                "[RÜZGAR] Gemini: GLOBAL_API_KEY yok (.env) — yerel Ollama kullanılır.",
                file=sys.stderr,
            )
            return
        if not _daemon_enabled():
            print(
                "[RÜZGAR] Gemini anahtarı yüklü — açılış ping atlandı (RUZGAR_GEMINI_DAEMON=0).",
                file=sys.stderr,
            )
            return
        from ilim_assistant.llm_gemini import gemini_model_ping

        ping = gemini_model_ping()
        if ping.get("ok"):
            print(
                f"[RÜZGAR] Gemini hazır: {ping.get('model')} "
                f"(GLOBAL_API_KEY yüklendi, daemon başlatılıyor)",
                file=sys.stderr,
            )
        else:
            print(
                f"[RÜZGAR] Gemini uyarı: model={ping.get('model')} "
                f"status={ping.get('status_code')} — {str(ping.get('reason', ''))[:120]}",
                file=sys.stderr,
            )
        start_gemini_daemon()
        st = daemon_status()
        if st.get("ok"):
            os.environ["RUZGAR_GEMINI_DAEMON_OK"] = "1"
    except Exception:
        pass


_gemini_startup_warmup()


def _rag_sync_warmup() -> None:
    """İlk sohbet turunda 60–120 sn RAG gecikmesini önler (async startup'tan önce)."""
    try:
        from ilim_assistant.rag_store import warmup_index

        warmup_index()
        print("[RÜZGAR] RAG indeks + gömme modeli önbelleğe alındı.", file=sys.stderr)
    except Exception as exc:
        print(f"[RÜZGAR] RAG warm-up atlandı: {exc}", file=sys.stderr)


try:
    from ilim_assistant.config import skip_rag_warmup

    if not skip_rag_warmup():
        _rag_sync_warmup()
    else:
        print("[RÜZGAR] RAG warm-up atlandı (RUZGAR_SKIP_RAG_WARMUP / hafif sohbet).", file=sys.stderr)
except Exception:
    _rag_sync_warmup()


def pdf_text_runtime_available() -> bool:
    try:
        import pypdf  # noqa: F401

        return True
    except ImportError:
        return False


def docx_text_runtime_available() -> bool:
    try:
        import docx  # noqa: F401

        return True
    except ImportError:
        return False


def _repo_resolve_under_root(rel_query: str, must_be_dir: bool = False) -> Path:
    raw = (rel_query or "").strip().replace("\\", "/").lstrip("/")
    root = REPO_ROOT.resolve()
    target = (root / raw.replace("/", os.sep)).resolve() if raw else root
    try:
        target.relative_to(root)
    except ValueError:
        raise HTTPException(status_code=400, detail="Geçersiz yol (proje dışı).")
    if must_be_dir:
        if not target.is_dir():
            raise HTTPException(status_code=404, detail="Klasör yok.")
    else:
        if not target.is_file():
            raise HTTPException(status_code=404, detail="Dosya yok.")
    return target


def _parse_bool_form(val: str | None, default: bool = True) -> bool:
    if val is None or str(val).strip() == "":
        return default
    return str(val).strip().lower() in ("1", "true", "yes", "on")


def _repo_list_children(rel_query: str) -> list[dict[str, Any]]:
    raw = (rel_query or "").strip().replace("\\", "/").lstrip("/")
    target = _repo_resolve_under_root(rel_query, must_be_dir=True)
    out: list[dict[str, Any]] = []
    for p in sorted(target.iterdir(), key=lambda x: (not x.is_dir(), x.name.lower())):
        name = p.name
        if name in _SKIP_LIST_NAMES:
            continue
        if name.startswith(".") and name != ".vscode":
            continue
        rel_full = f"{raw}/{name}" if raw else name
        rel_full = rel_full.replace("\\", "/")
        out.append({"name": name, "isDir": p.is_dir(), "rel": rel_full})
    return out


def _simple_html_to_text(raw: str) -> str:
    text = unescape(raw or "")
    text = re.sub(r"(?is)<script[^>]*>.*?</script>", " ", text)
    text = re.sub(r"(?is)<style[^>]*>.*?</style>", " ", text)
    text = re.sub(r"(?i)<br\s*/?>", "\n", text)
    text = re.sub(r"(?i)</p\s*>", "\n\n", text)
    text = re.sub(r"(?i)</h[1-6]\s*>", "\n\n", text)
    text = re.sub(r"(?s)<[^>]+>", " ", text)
    text = re.sub(r"[ \t]+\n", "\n", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r"[ \t]{2,}", " ", text)
    return text.strip()


def _rtf_to_text(raw: str) -> str:
    s = raw or ""
    s = re.sub(r"\\'[0-9a-fA-F]{2}", " ", s)
    s = re.sub(r"\\par[d]?", "\n", s)
    s = re.sub(r"\\tab", "\t", s)
    s = re.sub(r"\\[a-zA-Z]+\d* ?", "", s)
    s = s.replace("{", "").replace("}", "")
    s = re.sub(r"\n{3,}", "\n\n", s)
    s = re.sub(r"[ \t]{2,}", " ", s)
    return s.strip()


def _read_epub_text(target: Path) -> tuple[str, dict[str, Any]]:
    from ilim_assistant.motorlar.tercume_ebook_read import read_epub

    hit = read_epub(target)
    meta = {
        "chapters_read": hit.get("chapters_read") or 0,
        "title": hit.get("title") or "",
        "author": hit.get("author") or "",
    }
    return str(hit.get("text") or ""), meta


def _read_fb2_text(target: Path) -> tuple[str, dict[str, Any]]:
    from ilim_assistant.motorlar.tercume_ebook_read import read_fb2

    hit = read_fb2(target)
    meta = {
        "sections_read": hit.get("chapters_read") or 0,
        "chapters_read": hit.get("chapters_read") or 0,
        "title": hit.get("title") or "",
        "author": hit.get("author") or "",
    }
    return str(hit.get("text") or ""), meta


def _command_exists(name: str) -> bool:
    return shutil.which(name) is not None


def _convert_with_calibre_to_text(target: Path) -> str:
    from ilim_assistant.motorlar.tercume_ebook_read import convert_calibre_to_text

    return convert_calibre_to_text(target)


def _convert_djvu_to_text(target: Path) -> str:
    from ilim_assistant.motorlar.tercume_ebook_read import convert_djvu_to_text

    return convert_djvu_to_text(target)

# Electron (file://) ↔ yerel API: geçici tam açık CORS — köprü/socket kapanmasın.
# allow_credentials + "*" birlikte tarayıcıda CORS'u bozar; permissive modda credentials kapalı.
_cors_permissive = os.environ.get("RUZGAR_CORS_PERMISSIVE", "1").strip().lower() in (
    "1",
    "true",
    "yes",
    "on",
)
if _cors_permissive:
    app.add_middleware(
        CORSMiddleware,
        allow_origins=["*"],
        allow_credentials=False,
        allow_methods=["*"],
        allow_headers=["*"],
        expose_headers=["*"],
    )
else:
    app.add_middleware(
        CORSMiddleware,
        allow_origins=[
            "http://127.0.0.1:8779",
            "http://localhost:8779",
            "http://127.0.0.1:8777",
            "http://localhost:8777",
            "null",
        ],
        allow_credentials=True,
        allow_methods=["*"],
        allow_headers=["*"],
        expose_headers=["*"],
    )

from starlette.middleware.base import BaseHTTPMiddleware


class _RuzgarUserContextMiddleware(BaseHTTPMiddleware):
    """Halk modunda X-Ruzgar-User ile kişisel hafıza yolu seçilir."""

    async def dispatch(self, request, call_next):
        from ilim_assistant.ruzgar_public_mode import bind_request_user, clear_request_user

        hdr = request.headers.get("X-Ruzgar-User") or request.headers.get("X-Ruzgar-Session")
        bind_request_user(hdr)
        try:
            return await call_next(request)
        finally:
            clear_request_user()


app.add_middleware(_RuzgarUserContextMiddleware)

# Masaüstü kabuk — http://127.0.0.1:PORT/ui (file:// yok; sunucu health ile şerit yeşil)
_DESKTOP_UI_DIR = REPO_ROOT / "ruzgar-desktop"
if (_DESKTOP_UI_DIR / "index.html").is_file():
    from fastapi.staticfiles import StaticFiles
    from starlette.middleware.base import BaseHTTPMiddleware
    from starlette.requests import Request
    from starlette.responses import HTMLResponse, Response

    _UI_INDEX_PATH = _DESKTOP_UI_DIR / "index.html"
    _UI_NO_CACHE = {"Cache-Control": "no-store, no-cache, must-revalidate", "Pragma": "no-cache"}

    class _NoCacheUiMiddleware(BaseHTTPMiddleware):
        async def dispatch(self, request: Request, call_next) -> Response:
            response = await call_next(request)
            if request.url.path.startswith("/ui"):
                for k, v in _UI_NO_CACHE.items():
                    response.headers[k] = v
            return response

    def _ui_api_root_snippet() -> str:
        if os.environ.get("PORT") or os.environ.get("RUZGAR_PUBLIC_MODE", "").strip().lower() in (
            "1",
            "true",
            "yes",
            "on",
        ):
            return '<script>window.__RUZGAR_API_ROOT__=window.location.origin;</script>'
        from ilim_assistant.ruzgar_api_port import resolve_api_port

        port = resolve_api_port()
        return f'<script>window.__RUZGAR_API_ROOT__="http://127.0.0.1:{port}";</script>'

    def _ui_index_html() -> str:
        raw = _UI_INDEX_PATH.read_text(encoding="utf-8")
        ok_snip = _ui_api_root_snippet()
        if ok_snip not in raw and "__RUZGAR_API_ROOT__" not in raw:
            raw = raw.replace("</head>", f"  {ok_snip}\n  </head>", 1)
        return raw

    @app.get("/ui", include_in_schema=False)
    @app.get("/ui/", include_in_schema=False)
    @app.get("/ui/index.html", include_in_schema=False)
    def ruzgar_ui_index() -> HTMLResponse:
        return HTMLResponse(_ui_index_html(), headers=_UI_NO_CACHE)

    app.add_middleware(_NoCacheUiMiddleware)
    app.mount(
        "/ui",
        StaticFiles(directory=str(_DESKTOP_UI_DIR), html=False),
        name="ruzgar_desktop_ui",
    )


def _boot_motorlar_anaonce() -> None:
    """.cursorrules — Ana motor → çekirdek → 5 yardımcı (sıra kilitli)."""
    import importlib

    importlib.import_module("ilim_assistant.main_engine")
    importlib.import_module("ilim_assistant.motorlar.ruzgar_cekirdegi")
    for name in (
        "ilim_assistant.ses_motoru",
        "ilim_assistant.video_motoru",
        "ilim_assistant.okuma_motoru",
        "ilim_assistant.mimar_motoru",
        "ilim_assistant.tercume_motoru",
        "ilim_assistant.motorlar.hafiza_motoru",
        "ilim_assistant.programlama_motoru",
    ):
        importlib.import_module(name)


@app.on_event("startup")
async def _warmup_rag() -> None:
    """İlk sohbet turunda RAG disk okumasını ve gömme modelini önceden yükle."""
    try:
        from ilim_assistant.config import defer_motor_boot, skip_rag_warmup

        if not defer_motor_boot():
            _boot_motorlar_anaonce()
    except Exception:
        pass
    try:
        from ilim_assistant.config import skip_rag_warmup

        if not skip_rag_warmup():
            from ilim_assistant.rag_store import warmup_index

            warmup_index()
    except Exception:
        pass
    try:
        _get_hafiza_motor()
        from ilim_assistant.ruzgar_egitim import (
            ensure_canonical_egitim_pairs,
            sanitize_egitim_hafiza,
            sync_greeting_egitim_aliases,
        )

        ensure_canonical_egitim_pairs()
        ng = sync_greeting_egitim_aliases()
        if ng:
            print(f"[Rüzgar] Selam tetikleyicileri eşlendi: {ng} kayıt.", flush=True)
        n = sanitize_egitim_hafiza()
        try:
            from ilim_assistant.ruzgar_bilissel_analiz import sanitize_empati_hafiza

            ne = sanitize_empati_hafiza()
            if ne:
                print(f"[Rüzgar] Empati hafızası temizlendi: {ne} kayıt.", flush=True)
        except Exception:
            pass
        if n:
            print(f"[Rüzgar] Eğitim hafızası temizlendi: {n} hatalı kayıt.", flush=True)
    except Exception:
        pass
    try:
        from ilim_assistant.dinamit_hatirlatici import start_reminder_background_thread
        from ilim_assistant.gorev_yoneticisi import init_tasks_db

        start_reminder_background_thread()
        init_tasks_db()
    except Exception:
        pass
    try:
        from ilim_assistant.ruzgar_owner_lock import startup_owner_banner

        print(startup_owner_banner(), flush=True)
    except Exception:
        pass
    if os.environ.get("RUZGAR_PRINT_READY_SEAL", "1").strip().lower() not in (
        "0",
        "false",
        "no",
    ):
        print(
            "Rüzgar Kullanıma Hazır, Sistemi Yeniden Başlatabilirsiniz.",
            flush=True,
        )


def _default_fetch_pages() -> float:
    try:
        return float(os.environ.get("RUZGAR_DEFAULT_FETCH_PAGES", "1"))
    except ValueError:
        return 1.0


class ChatRequest(BaseModel):
    message: str = ""
    history: list[dict[str, Any]] = Field(default_factory=list)
    use_web: bool = True
    fetch_pages: float = Field(default_factory=_default_fetch_pages)
    coding_mode: bool = False
    session_wake_used: bool = False
    # genel | ses | okuma | video | uretim | programlama | gelisim | duzen | dosya | hizli
    # None/boş: eski istemciler veya hatalı JSON; Kod modu açıksa programlama varsayılır (prefetch atlanır).
    mode: str | None = None
    # Electron getRoot() — @@dosya yolları bu köke göre çözülür
    workspace_root: str | None = None
    # Web ara kapalı olsa bile mesajdaki https URL’lerini oku
    read_message_links: bool = True
    hizir_channels: list[str] | None = Field(
        default=None,
        description="HIZIR modunda seçili pazar kanalları; boş/atlanırsa tüm kanallar",
    )
    programlama_active_file: str | None = Field(
        default=None,
        description="Programlama atölyesinde açık dosya (göreli yol)",
    )
    programlama_editor_snippet: str | None = Field(
        default=None,
        description="Editördeki kod özeti (Faz 5 oturum bağlamı)",
    )
    programlama_language: str | None = Field(
        default=None,
        description="Atölye dil seçici (python, javascript, …)",
    )


def _effective_chat_mode_raw(req: ChatRequest) -> str:
    """Kod modu açıksa her zaman programlama (UI genel kalsa bile tam indeks prefetch kapalı)."""
    if req.coding_mode:
        return "programlama"
    raw = (req.mode or "").strip() or "genel"
    try:
        from ilim_assistant.chat_core import normalize_mode
        from ilim_assistant.idrak_entegrasyon import motor_niyeti_heuristic
        from ilim_assistant.motorlar.programlama_faz10 import should_delegate_to_programlama

        mode_norm = normalize_mode(raw)
        flags = motor_niyeti_heuristic((req.message or "").strip())
        if should_delegate_to_programlama(
            req.message or "",
            mode_norm,
            coding_mode=False,
            motor_flags=flags,
        ):
            return "programlama"
    except Exception:
        pass
    return raw


class ReminderAckBody(BaseModel):
    """Dinamit hatırlatıcı onayı (masaüstü uyumluluğu)."""

    id: str = ""


class TaskCreateBody(BaseModel):
    title: str = ""
    detail: str = ""


class TaskUpdateBody(BaseModel):
    id: int = 0
    status: str = "done"
    detail: str | None = None


class VideoDownloadBody(BaseModel):
    url: str = ""


class VideoConcatBody(BaseModel):
    """İki proje içi dosyayı concat ile birleştirir (codec uyumu şarta bağlı)."""

    rel_a: str = ""
    rel_b: str = ""
    copy_streams: bool = True


class VideoEditClipItem(BaseModel):
    rel: str = ""
    start_sec: float = 0.0
    end_sec: float | None = None
    label: str = ""


class VideoEditMixBody(BaseModel):
    """Timeline kurgusu: kesilmiş parçaları sırayla birleştirir (parça ≤ 5 dk)."""

    clips: list[VideoEditClipItem] = Field(default_factory=list)
    copy_streams: bool = True
    project_name: str = ""


def _sse(obj: dict) -> str:
    return f"data: {json.dumps(obj, ensure_ascii=False)}\n\n"


class HafizaArsivAdd(BaseModel):
    raw: str = ""


class HafizaMotorRead(BaseModel):
    raw: str = ""


class HafizaImportBlok(BaseModel):
    raw: str = ""


class GenelHafizaBakBody(BaseModel):
    message: str = ""


class IdrakPretreatBody(BaseModel):
    message: str = ""
    history: list[dict[str, Any]] = Field(default_factory=list)


class TtsBody(BaseModel):
    text: str = ""
    karakter: str = "asistan"
    backend: str = "edge"
    emotion: str | None = None


class SesSettingsPatchBody(BaseModel):
    karakter: str | None = None
    hiz: float | None = None
    huzur: float | None = None


class MimarFotoModerateBody(BaseModel):
    rel: str = Field(..., description="ilim-assistant/arsiv/mimar-fotograf/…")
    op: str = Field(..., description="rotate_left, preset_auto, crop_square, …")


class MimarFotoRelBody(BaseModel):
    rel: str = Field(..., description="ilim-assistant/arsiv/mimar-fotograf/…")


class MimarFotoVoiceBody(BaseModel):
    rel: str = Field(..., description="ilim-assistant/arsiv/mimar-fotograf/…")
    mode: str = Field(default="read", description="read | speak")
    lang: str = Field(default="tur+eng", description="OCR dili (ocr uç noktası)")


class MimarSanatRelBody(BaseModel):
    rel: str = Field(..., description="ilim-assistant/arsiv/mimar-sanat/…")


class MimarSanatAnalyzeBody(BaseModel):
    rel: str = Field(..., description="ilim-assistant/arsiv/mimar-sanat/…")
    depth: str = Field(default="deep", description="quick | deep")


class MimarAtolyeParseBody(BaseModel):
    message: str = ""


class MimarSanatCopyBody(BaseModel):
    rel: str = Field(..., description="ilim-assistant/arsiv/mimar-sanat/…")
    mode: str = Field(default="trace", description="trace | poster | pencil")


class MimarSanatMetaBody(BaseModel):
    rel: str = Field(..., description="ilim-assistant/arsiv/mimar-sanat/…")
    title: str | None = Field(default=None, max_length=240)
    artist: str | None = Field(default=None, max_length=240)
    period: str | None = Field(default=None, max_length=240)
    technique: str | None = Field(default=None, max_length=240)
    notes: str | None = Field(default=None, max_length=12000)


class MimarTasarimSaveBody(BaseModel):
    project: dict[str, Any] = Field(..., description="Tuval projesi JSON")


class MimarTasarimLoadBody(BaseModel):
    project_id: str = Field(..., description="Proje kimliği")


class MimarTasarimSketchTextBody(BaseModel):
    text: str = Field(..., description="Sohbet betimlemesi veya plan notu")
    width: int = Field(default=960, ge=320, le=2400)
    height: int = Field(default=540, ge=240, le=1600)


class MimarTasarimSketchImageBody(BaseModel):
    rel: str = Field(..., description="Referans görsel yolu (mimar-tasarim arşivi)")
    width: int = Field(default=960, ge=320, le=2400)
    height: int = Field(default=540, ge=240, le=1600)


class MimarTasarimChatHandoffBody(BaseModel):
    user: str = Field(default="", description="Son kullanıcı mesajı")
    assistant: str = Field(default="", description="Son asistan yanıtı")
    notes: str = Field(default="", description="Tuval notu")
    width: int = Field(default=960, ge=320, le=2400)
    height: int = Field(default=540, ge=240, le=1600)
    project_id: str = Field(default="", description="Kayıtlı proje kimliği (opsiyonel)")


class MimarTasarimDuplicateBody(BaseModel):
    project_id: str = Field(..., description="Kaynak proje kimliği")
    name: str | None = Field(default=None, max_length=120)


class HizirPazarTaraBody(BaseModel):
    query: str = Field(default="", description="Ürün veya arama (boşsa genel tarama metni)")
    channels: list[str] | None = Field(
        default=None,
        description="Pazar kanalları (trendyol, amazon_tr, hepsiburada, amazon_us, amazon_gb, amazon_de, ebay, aliexpress); boş veya atlanırsa tümü",
    )


class HizirFirsatKaldirBody(BaseModel):
    kart_id: str = Field(default="", description="ARBITRAJ kart_id (Ticaret Avcısı)")


class ProgramlamaSessionPatchBody(BaseModel):
    workspace_root: str | None = None
    name: str | None = None
    goal: str | None = None
    stack: list[str] | str | None = None
    notes: str | None = None


class ProgramlamaScaffoldBody(BaseModel):
    template_id: str = Field(..., description="cli_python | fastapi_api | python_package | mini_ai_bot")
    project_name: str = Field(..., description="Proje adı (slug üretilir)")
    workspace_root: str | None = None
    force: bool = False


class ProgramlamaRunBody(BaseModel):
    workspace_root: str | None = None
    rel: str | None = Field(default=None, description="Dosya veya proje göreli yolu (projects/...)")
    smoke_only: bool = False
    action: str | None = Field(
        default="start",
        description="Faz 8 serve: start | stop | status",
    )


class CodeRunBody(BaseModel):
    code: str = ""
    language: str = "python"
    timeout_sec: float = 8.0
    workspace_root: str | None = Field(
        default=None,
        description="Electron proje kökü (örn. YAPAY ZEKA); sunucuda doğrulanır.",
    )
    cwd_rel: str | None = Field(
        default=None,
        description="Çalışma dizini köke göre (örn. ruzgar-desktop); kayıtlı dosyanın klasörü.",
    )


def _normalize_workspace_root(raw: str | None) -> str | None:
    wr = (raw or "").strip()
    if not wr:
        return None
    n = os.path.normpath(wr)
    return n if os.path.isdir(n) else None


def _is_inside_workspace(workspace_root: str, candidate: str) -> bool:
    wr = os.path.normpath(workspace_root)
    ca = os.path.normpath(candidate)
    try:
        return os.path.commonpath([wr, ca]) == wr
    except ValueError:
        return False


def _resolve_workspace_run_dir(workspace_root: str, cwd_rel: str | None) -> str | None:
    """Güvenli çalışma dizini: kök veya kök altı klasör (dosya yolu verilirse üst klasör)."""
    wr = _normalize_workspace_root(workspace_root)
    if not wr:
        return None
    rel = (cwd_rel or "").strip().replace("\\", "/").lstrip("/")
    if not rel:
        return wr
    cand = os.path.normpath(os.path.join(wr, rel.replace("/", os.sep)))
    if not _is_inside_workspace(wr, cand):
        return None
    if os.path.isfile(cand):
        cand = os.path.dirname(cand)
    return cand if os.path.isdir(cand) else None


@app.post("/api/hafiza/genel-bak")
def api_genel_hafiza_bak(body: GenelHafizaBakBody):
    """Ana sohbet ön kontrol: ham JSON cevabı (dogal sentez açıkken devre dışı)."""
    m = (body.message or "").strip()
    if not m:
        return {"ok": True, "hit": False, "answer": None}
    try:
        from ilim_assistant.hafiza_dogal_sentez import dogal_konus_enabled

        if dogal_konus_enabled():
            return {"ok": True, "hit": False, "answer": None, "dogal_sentez": True}
    except Exception:
        pass
    ans = genel_hafiza_lookup(m)
    if ans is None:
        return {"ok": True, "hit": False, "answer": None}
    out = finalize_assistant_reply(ans)
    return {"ok": True, "hit": True, "answer": out}


@app.post("/api/idrak/pretreat")
def api_idrak_pretreat(body: IdrakPretreatBody):
    """Ses/STT ve sohbet kutusu için hızlı yazım + devam cümlesi ön-işlemi."""
    from ilim_assistant.idrak_on_islem import pretreat_user_turn

    pt = pretreat_user_turn(body.message, body.history)
    return {
        "ok": True,
        "text": pt.text,
        "changed": pt.changed,
        "continuation": pt.continuation,
        "replacements": list(pt.replacements),
    }


@app.get("/api/hafiza/arsiv")
def api_hafiza_arsiv():
    """Öğrenme sözlüğü — arayüz sadece `data` ile motor cevabı eşlemesi yapar (analiz listesi ayrı)."""
    m = _get_hafiza_motor()
    return {"ok": True, "data": m.tum_bilgiler(motor_tipi="Hafıza"), "items": []}


@app.post("/api/hafiza/arsiv/add")
def api_hafiza_arsiv_add(body: HafizaArsivAdd):
    """Satırları `soru = cevap` formatında ayrıştırır ve kalıcı hafızaya yazar."""
    m = _get_hafiza_motor()
    raw = (body.raw or "").strip()
    if not raw:
        raise HTTPException(status_code=400, detail="Boş kayıt")
    added = 0
    for line in raw.replace("\r\n", "\n").replace("\r", "\n").split("\n"):
        txt = line.strip()
        if not txt or txt.startswith("#"):
            continue
        if "=" not in txt:
            continue
        soru, _, cevap = txt.partition("=")
        soru = soru.strip()
        cevap = cevap.strip()
        if soru and cevap:
            m.ekle_bilgi(soru, cevap, motor_tipi="Hafıza")
            added += 1

    if added == 0:
        raise HTTPException(status_code=400, detail="Geçerli 'Soru = Cevap' satırı yok")

    return {"ok": True, "added": added}


@app.post("/api/hafiza/motor-read")
def api_hafiza_motor_read(body: HafizaMotorRead):
    m = _get_hafiza_motor()
    ans = m.cevap_ver(body.raw or "")
    return {"ok": True, "answer": ans}


@app.post("/api/hafiza/import-blok")
def api_hafiza_import_blok(body: HafizaImportBlok):
    """Çok satırlı metinde yalnızca `soru = cevap` satırlarını öğrenir."""
    m = _get_hafiza_motor()
    items = m.import_metin_blok(body.raw or "", motor_tipi="Hafıza")
    return {"ok": True, "added": len(items), "items": items}


def _super_brain_health_block() -> dict[str, Any]:
    try:
        from ilim_assistant.llm_brain import brain_health_snapshot

        return brain_health_snapshot()
    except Exception as e:
        return {"ok": False, "error": str(e)}


def _connection_ui_block(super_brain: dict[str, Any] | None = None) -> dict[str, Any]:
    try:
        from ilim_assistant.ruzgar_connection_info import build_connection_info

        sb = super_brain if super_brain is not None else _super_brain_health_block()
        return build_connection_info(super_brain=sb)
    except Exception as exc:
        return {"ok": False, "error": str(exc)[:200]}


@app.get("/api/ui/connection")
def api_ui_connection():
    """Beyin zinciri + bellek dosya yolları (UI köprü paneli)."""
    sb = _super_brain_health_for_api()
    return _connection_ui_block(super_brain=sb)


def _super_brain_health_for_api() -> dict[str, Any]:
    """Health/manifest: Gemini daemon henüz kırmızıysa bir kez yeniden dene."""
    sb = _super_brain_health_block()
    try:
        gd = sb.get("gemini_daemon") if isinstance(sb.get("gemini_daemon"), dict) else {}
        if sb.get("gemini_configured") and not gd.get("ok"):
            from ilim_assistant.gemini_daemon import start_gemini_daemon

            start_gemini_daemon()
            sb = _super_brain_health_block()
    except Exception:
        pass
    return sb


def _ruzgar_mode_health_snapshot() -> dict[str, Any]:
    """Çalışan API’nin gerçek modu — UI/eski süreç teşhisi."""
    snap: dict[str, Any] = {
        "ollama_only": False,
        "gemini_disabled": False,
        "groq_disabled": False,
        "bilissel_enabled": False,
    }
    try:
        from ilim_assistant.config import gemini_disabled, groq_disabled, ollama_only_mode
        from ilim_assistant.ruzgar_bilissel_analiz import bilissel_enabled

        snap["ollama_only"] = ollama_only_mode()
        snap["gemini_disabled"] = gemini_disabled()
        snap["groq_disabled"] = groq_disabled()
        snap["bilissel_enabled"] = bilissel_enabled()
    except Exception:
        pass
    try:
        from ilim_assistant.ruzgar_public_mode import public_mode_health

        snap["public_pool"] = public_mode_health()
    except Exception:
        pass
    return snap


def _sample_cpu_percent() -> float | None:
    """İsteğe bağlı `psutil` — yoksa null (istemci yine 200 alır)."""
    try:
        import psutil  # type: ignore[import-not-found]

        return float(psutil.cpu_percent(interval=0.05))
    except Exception:
        return None


def _sample_gpu_percent() -> float | None:
    """NVIDIA GPU kullanımı; yoksa null."""
    try:
        import shutil
        import subprocess

        exe = shutil.which("nvidia-smi")
        if not exe:
            return None
        r = subprocess.run(
            [
                exe,
                "--query-gpu=utilization.gpu",
                "--format=csv,noheader,nounits",
            ],
            capture_output=True,
            text=True,
            timeout=2,
        )
        if r.returncode != 0:
            return None
        vals: list[float] = []
        for line in (r.stdout or "").splitlines():
            s = line.strip()
            if not s:
                continue
            try:
                vals.append(float(s))
            except ValueError:
                continue
        if not vals:
            return None
        return float(sum(vals) / len(vals))
    except Exception:
        return None


def _health_build_block() -> dict:
    import os as _os

    base = {
        "rev": "2026-05-27-ruzgar-faz98-v107",
        "nebula_kitap": True,
        "fast_paths": _os.environ.get("RUZGAR_FAST_PATHS", "1").strip(),
        "memory_first": True,
        "bilissel_analiz": True,
        "egitim_routing": "bilissel>gundelik>egitim",
    }
    try:
        from ilim_assistant.motorlar.programlama_faz95 import prompt_cache_metrics

        base["prompt_cache_v98"] = prompt_cache_metrics()
    except Exception:
        pass
    try:
        from ilim_assistant.motorlar.programlama_faz26 import (
            p9_strict_local_first,
            programming_brain_chain_ids,
        )

        _p9_chain = programming_brain_chain_ids()[:6]
        base["p9_v99"] = {
            "strict_local_first": p9_strict_local_first(),
            "programlama_chain": _p9_chain,
            "local_first_active": bool(_p9_chain and _p9_chain[0] in ("kod", "denge", "hizli")),
        }
    except Exception:
        pass
    try:
        from ilim_assistant.motorlar.programlama_faz60 import enrich_health_build

        return enrich_health_build(base)
    except Exception:
        return base


@app.get("/api/health")
def health():
    import os as _os

    _m = _os.environ.get("HIZIR_MOCK_MARKETPLACE", "0").strip().lower() in (
        "1",
        "true",
        "yes",
        "on",
    )
    _ak = (_os.environ.get("HIZIR_AMAZON_ACCESS_KEY") or "").strip()
    _sk = (_os.environ.get("HIZIR_AMAZON_SECRET_KEY") or "").strip()
    _tg = (_os.environ.get("HIZIR_AMAZON_PARTNER_TAG") or "").strip()
    _eb = bool(
        (_os.environ.get("HIZIR_EBAY_CLIENT_ID") or "").strip()
        and (_os.environ.get("HIZIR_EBAY_CLIENT_SECRET") or "").strip()
    )
    _ae = bool(
        (_os.environ.get("HIZIR_ALIEXPRESS_APP_KEY") or "").strip()
        and (_os.environ.get("HIZIR_ALIEXPRESS_APP_SECRET") or "").strip()
        and (_os.environ.get("HIZIR_ALIEXPRESS_ACCESS_TOKEN") or "").strip()
    )
    from ilim_assistant.chat_core import _main_chat_genel_only, _web_secondary_policy_enabled
    from ilim_assistant.defaults import DEFAULT_OLLAMA_CHAT_MODEL

    _main_only = _main_chat_genel_only()
    _sb = _super_brain_health_for_api()
    try:
        from ilim_assistant.motorlar.tercume_ocr_runtime import ocr_runtime_status

        _ocr = ocr_runtime_status()
    except Exception:
        _ocr = {"available": False, "cloud_ready": False, "hint": "OCR durumu okunamadi"}
    return {
        "ok": True,
        "service": "ruzgar-desktop-api",
        "stt": stt_runtime_available(),
        "pdf_text": pdf_text_runtime_available(),
        "docx_text": docx_text_runtime_available(),
        "ffmpeg": ffmpeg_available(),
        "ffprobe": ffprobe_available(),
        "ytdlp": __import__(
            "ilim_assistant.motorlar.video_motoru", fromlist=["ytdlp_available"]
        ).ytdlp_available(),
        "ocr": _ocr,
        "merkezi_bellek": True,
        "ana_motor": {
            "main_only_genel_hafiza": _main_only,
            "main_only_warning": (
                "Genel modda yalnızca hafıza JSON; LLM/RAG/web kapalı"
                if _main_only
                else None
            ),
            "web_secondary_strong_rag_only": _web_secondary_policy_enabled(),
            "question_plan_enabled": os.environ.get("RUZGAR_ANA_MOTOR_PLAN", "1").strip().lower()
            not in ("0", "false", "no"),
            "clarify_enabled": os.environ.get("RUZGAR_ANA_MOTOR_CLARIFY", "1").strip().lower()
            not in ("0", "false", "no"),
            "ollama_chat_model": os.environ.get("OLLAMA_CHAT_MODEL", DEFAULT_OLLAMA_CHAT_MODEL),
            "enable_web_search": os.environ.get("ENABLE_WEB_SEARCH", "1") == "1",
            "rag_top_k_default": os.environ.get("RAG_TOP_K", "2"),
            "rag_top_k_genel": os.environ.get("RUZGAR_GENEL_RAG_TOP_K", "4"),
            "archive_score_min": os.environ.get("RUZGAR_ARCHIVE_SCORE_MIN", "0.22"),
            "web_suppress_rag_min": os.environ.get("RUZGAR_WEB_SUPPRESS_RAG_MIN", "0.38"),
            "ana_motor_agent_enabled": os.environ.get("RUZGAR_ANA_MOTOR_AGENT", "1").strip().lower()
            not in ("0", "false", "no"),
        },
        "super_brain": _sb,
        "connection": _connection_ui_block(super_brain=_sb),
        "build": _health_build_block(),
        "ruzgar_mode": _ruzgar_mode_health_snapshot(),
        "hizir": {
            "mock_marketplace": _m,
            "amazon_paapi_credentials": bool(_ak and _sk and _tg),
            "ebay_browse_credentials": bool(_eb),
            "aliexpress_open_credentials": bool(_ae),
        },
    }


@app.get("/api/ui/manifest")
def api_ui_manifest():
    """Dashboard/atölye metinleri için tek kaynak; UI statik Faz 7'ye takılmaz."""
    from ilim_assistant.ruzgar_ui_manifest import build_ui_manifest

    return build_ui_manifest(health=health())


@app.get("/api/system-metrics")
def api_system_metrics():
    """Masaüstü performans pill — `ruzgar-desktop/app.js` her ~2.5 sn çağırır."""
    return {
        "cpu_percent": _sample_cpu_percent(),
        "gpu_percent": _sample_gpu_percent(),
    }


@app.get("/api/reminders/pending")
def api_reminders_pending():
    """Dinamit hatırlatıcılar — istemci 24 sn'de bir poll eder."""
    from ilim_assistant.dinamit_hatirlatici import fetch_due_reminders

    return {"ok": True, "items": fetch_due_reminders()}


@app.post("/api/reminders/ack")
def api_reminders_ack(body: ReminderAckBody):
    from ilim_assistant.dinamit_hatirlatici import mark_triggered

    raw = (body.id or "").strip()
    try:
        rid = int(raw)
    except ValueError:
        rid = 0
    if rid > 0:
        mark_triggered(rid)
    return {"ok": True, "id": raw}


@app.get("/api/self-test")
def api_self_test():
    """Rüzgar kendi sağlık ve mod davranışını hızlı kontrol eder."""
    from ilim_assistant.ruzgar_selftest import run_self_tests

    return run_self_tests()


@app.get("/api/programlama/briefing")
def api_programlama_briefing(workspace_root: str | None = None):
    """Programlama atölyesi açılış brifingi + öz-bilgi manifesti (Faz 2)."""
    from ilim_assistant.motorlar.programlama_faz2 import build_startup_briefing

    root = (workspace_root or "").strip() or None
    return build_startup_briefing(root)


@app.get("/api/programlama/self-scan")
def api_programlama_self_scan(workspace_root: str | None = None):
    """Genişletilmiş programlama öz-denetimi (JSON)."""
    from ilim_assistant.motorlar.programlama_faz2 import run_programlama_self_scan

    root = (workspace_root or "").strip() or None
    return run_programlama_self_scan(root)


@app.get("/api/programlama/environment")
def api_programlama_environment(workspace_root: str | None = None):
    """Windows / yerel geliştirme ortamı taraması (Faz 3)."""
    from ilim_assistant.motorlar.programlama_faz3 import FAZ3_VERSION, run_windows_env_scan

    root = (workspace_root or "").strip() or None
    tests = run_windows_env_scan(root)
    ok = all(t.get("ok") for t in tests)
    return {"ok": ok, "version": FAZ3_VERSION, "tests": tests}


@app.get("/api/programlama/security-audit")
def api_programlama_security_audit(workspace_root: str | None = None):
    """Güvenlik denetimi (Faz 4)."""
    from ilim_assistant.motorlar.programlama_faz4 import run_security_audit

    root = (workspace_root or "").strip() or None
    return run_security_audit(root)


@app.get("/api/programlama/session")
def api_programlama_session(workspace_root: str | None = None):
    """Faz 5 — oturum proje bağlamı (JSON + LLM metni)."""
    from ilim_assistant.motorlar.programlama_faz5 import build_session_api_payload

    root = (workspace_root or "").strip() or None
    return build_session_api_payload(root)


@app.post("/api/programlama/session")
def api_programlama_session_patch(body: ProgramlamaSessionPatchBody):
    """Faz 5 — proje adı/hedef/yığın güncelle."""
    from ilim_assistant.motorlar.programlama_faz5 import apply_project_patch, build_session_api_payload

    root = (body.workspace_root or "").strip() or None
    patch = {k: v for k, v in body.model_dump().items() if k != "workspace_root" and v is not None}
    apply_project_patch(root, patch)
    return build_session_api_payload(root)


@app.post("/api/programlama/session/clear")
def api_programlama_session_clear(workspace_root: str | None = None):
    """Faz 5 — oturum bağlamını sıfırla."""
    from ilim_assistant.motorlar.programlama_faz5 import build_session_api_payload, clear_session

    root = (workspace_root or "").strip() or None
    clear_session(root)
    return build_session_api_payload(root)


@app.get("/api/programlama/templates")
def api_programlama_templates():
    """Faz 6 — kullanılabilir proje şablonları."""
    from ilim_assistant.motorlar.programlama_faz6 import FAZ6_VERSION, list_templates

    return {"ok": True, "version": FAZ6_VERSION, "templates": list_templates()}


@app.post("/api/programlama/scaffold")
def api_programlama_scaffold(body: ProgramlamaScaffoldBody):
    """Faz 6 — şablon projeyi workspace/projects/ altına oluştur."""
    from ilim_assistant.motorlar.programlama_faz6 import format_scaffold_report, run_scaffold

    root = (body.workspace_root or "").strip() or None
    result = run_scaffold(
        body.template_id,
        body.project_name,
        root,
        force=bool(body.force),
    )
    from ilim_assistant.motorlar.programlama_faz8 import (
        FAZ8_VERSION,
        apply_scaffold_focus,
        enrich_scaffold_report,
    )

    focus = apply_scaffold_focus(root, result) if result.get("ok") else {}
    result["report"] = enrich_scaffold_report(result, focus)
    result["focus"] = focus
    result["faz8_version"] = FAZ8_VERSION
    return result


@app.get("/api/programlama/run-guide")
def api_programlama_run_guide(
    workspace_root: str | None = None,
    rel: str | None = None,
    message: str | None = None,
):
    """Faz 7 — dosya/proje açıklama ve çalıştırma rehberi (LLM yok)."""
    from ilim_assistant.motorlar.programlama_faz7 import FAZ7_VERSION, format_explain_run_report

    root = (workspace_root or "").strip() or None
    r = (rel or "").strip() or None
    msg = (message or "").strip() or "nasıl çalıştırırım"
    if r:
        msg = f"{msg} {r}"
    text = format_explain_run_report(msg, root, active_file=r)
    return {
        "ok": bool(text),
        "version": FAZ7_VERSION,
        "report": text or "Rehber üretilemedi.",
    }


@app.post("/api/programlama/run")
def api_programlama_run(body: ProgramlamaRunBody):
    """Faz 7 — şablon projeyi güvenli smoke/çalıştır (projects/ altı)."""
    from ilim_assistant.motorlar.programlama_faz7 import (
        FAZ7_VERSION,
        format_run_report,
        resolve_target_rel,
        run_project_profile,
    )

    root = (body.workspace_root or "").strip() or None
    rel = (body.rel or "").strip() or None
    if not rel:
        rel = resolve_target_rel("", workspace_root=root)
    if not rel:
        raise HTTPException(status_code=400, detail="rel veya açık oturum dosyası gerekli.")
    result = run_project_profile(root, rel, smoke_only=bool(body.smoke_only))
    result["report"] = format_run_report(result)
    result["version"] = FAZ7_VERSION
    return result


class ProgramlamaPatchBody(BaseModel):
    workspace_root: str | None = None
    text: str = ""
    rel: str | None = None
    path: str | None = None
    status: str | None = None
    mode: str = "accepted"
    run_verify: bool = True


@app.get("/api/programlama/workspace-projects")
def api_programlama_workspace_projects(workspace_root: str | None = None):
    """Faz 29 — çoklu proje listesi + aktif proje."""
    from ilim_assistant.motorlar.programlama_faz29 import (
        FAZ29_VERSION,
        get_workspace_projects_state,
    )

    root = (workspace_root or "").strip() or None
    return get_workspace_projects_state(root) | {"version": FAZ29_VERSION}


class ProgramlamaWorkspaceSwitchBody(BaseModel):
    workspace_root: str | None = None
    project_slug: str = ""


@app.post("/api/programlama/workspace/switch")
def api_programlama_workspace_switch(body: ProgramlamaWorkspaceSwitchBody):
    """Faz 29 — aktif projeyi değiştir."""
    from ilim_assistant.motorlar.programlama_faz29 import (
        FAZ29_VERSION,
        switch_to_project,
    )

    root = (body.workspace_root or "").strip() or None
    res = switch_to_project(root, body.project_slug)
    if not res.get("ok"):
        raise HTTPException(status_code=404, detail=res.get("error") or "geçiş başarısız")
    return {**res, "version": FAZ29_VERSION}


@app.get("/api/programlama/workspace-index")
def api_programlama_workspace_index(
    workspace_root: str | None = None,
    scope_rel: str | None = None,
    active_file: str | None = None,
):
    """Faz 10.1 — workspace / proje indeksi."""
    from ilim_assistant.motorlar.programlama_faz10 import (
        FAZ10_VERSION,
        build_workspace_index,
        resolve_scope_rel,
    )

    root = (workspace_root or "").strip() or None
    scope = (scope_rel or "").strip() or resolve_scope_rel(root, active_file=active_file)
    idx = build_workspace_index(root, scope_rel=scope)
    return {
        "ok": bool(idx),
        "scope_rel": scope or "",
        "index": idx,
        "version": FAZ10_VERSION,
    }


@app.get("/api/programlama/project-scan")
def api_programlama_project_scan(
    workspace_root: str | None = None,
    scope_rel: str | None = None,
    active_file: str | None = None,
):
    """Faz 13 — projects/<ad>/ dosya envanteri."""
    from ilim_assistant.motorlar.programlama_faz13 import (
        FAZ13_VERSION,
        resolve_scope_rel,
        scan_project_files,
    )

    root = (workspace_root or "").strip() or None
    scope = (scope_rel or "").strip() or resolve_scope_rel(root, active_file=active_file)
    if not scope:
        raise HTTPException(status_code=400, detail="scope_rel veya active_file gerekli")
    out = scan_project_files(root, scope)
    out["version"] = FAZ13_VERSION
    return out


@app.get("/api/programlama/project-summary")
def api_programlama_project_summary(
    workspace_root: str | None = None,
    scope_rel: str | None = None,
    active_file: str | None = None,
):
    """Faz 13 — LLM bağlam özeti (metin bloğu)."""
    from ilim_assistant.motorlar.programlama_faz13 import (
        FAZ13_VERSION,
        build_project_summary_block,
        resolve_scope_rel,
    )

    root = (workspace_root or "").strip() or None
    scope = (scope_rel or "").strip() or resolve_scope_rel(root, active_file=active_file)
    if not scope:
        raise HTTPException(status_code=400, detail="scope_rel veya active_file gerekli")
    block = build_project_summary_block(root, scope_rel=scope)
    return {
        "ok": bool(block),
        "scope_rel": scope,
        "summary": block,
        "version": FAZ13_VERSION,
    }


@app.get("/api/programlama/project-find")
def api_programlama_project_find(
    workspace_root: str | None = None,
    scope_rel: str | None = None,
    active_file: str | None = None,
    pattern: str = "",
):
    """Faz 13 — proje içi metin arama."""
    from ilim_assistant.motorlar.programlama_faz13 import (
        FAZ13_VERSION,
        format_find_report,
        resolve_scope_rel,
        search_in_project,
    )

    root = (workspace_root or "").strip() or None
    scope = (scope_rel or "").strip() or resolve_scope_rel(root, active_file=active_file)
    if not scope:
        raise HTTPException(status_code=400, detail="scope_rel veya active_file gerekli")
    pat = (pattern or "").strip()
    if not pat:
        raise HTTPException(status_code=400, detail="pattern gerekli")
    result = search_in_project(root, scope, pat)
    result["report"] = format_find_report(result)
    result["version"] = FAZ13_VERSION
    return result


@app.get("/api/programlama/agent/status")
def api_programlama_agent_status(workspace_root: str | None = None):
    """Faz 14 — otonom görev durumu."""
    from ilim_assistant.motorlar.programlama_faz14 import (
        FAZ14_VERSION,
        format_agent_status_report,
        load_agent_state,
    )

    root = (workspace_root or "").strip() or None
    st = load_agent_state(root)
    return {
        "ok": True,
        "state": st,
        "report": format_agent_status_report(root),
        "version": FAZ14_VERSION,
    }


@app.post("/api/programlama/agent/stop")
def api_programlama_agent_stop(workspace_root: str | None = None):
    """Faz 14 — görev döngüsünü durdur."""
    from ilim_assistant.motorlar.programlama_faz14 import (
        FAZ14_VERSION,
        format_stop_report,
        request_agent_stop,
    )

    root = (workspace_root or "").strip() or None
    request_agent_stop(root)
    return {
        "ok": True,
        "report": format_stop_report(),
        "version": FAZ14_VERSION,
    }


@app.get("/api/programlama/terminal/presets")
def api_programlama_terminal_presets():
    """Faz 15 — izinli terminal preset listesi."""
    try:
        from ilim_assistant.motorlar.programlama_faz43 import (
            FAZ43_VERSION,
            list_terminal_presets_v3,
        )

        presets = list_terminal_presets_v3()
        ver = FAZ43_VERSION
    except Exception:
        from ilim_assistant.motorlar.programlama_faz15 import (
            FAZ15_VERSION,
            list_terminal_presets,
        )

        presets = list_terminal_presets()
        ver = FAZ15_VERSION

    return {
        "ok": True,
        "presets": presets,
        "version": ver,
    }


@app.post("/api/programlama/terminal/run")
def api_programlama_terminal_run(body: dict):
    """Faz 43 — projects/ kapsamında onaylı terminal (npm/git/pip/pytest)."""
    from ilim_assistant.motorlar.programlama_faz15 import is_dangerous_shell

    try:
        from ilim_assistant.motorlar.programlama_faz43 import (
            FAZ43_VERSION,
            format_terminal_report_v3,
            run_terminal_v3,
        )

        fmt = format_terminal_report_v3
        run_fn = run_terminal_v3
        ver = FAZ43_VERSION
    except Exception:
        from ilim_assistant.motorlar.programlama_faz15 import (
            FAZ15_VERSION,
            format_terminal_report,
            run_terminal_preset,
        )

        fmt = format_terminal_report
        run_fn = run_terminal_preset
        ver = FAZ15_VERSION

    preset = str(body.get("preset") or body.get("preset_id") or "").strip()
    raw_cmd = str(body.get("command") or "").strip()
    if raw_cmd and is_dangerous_shell(raw_cmd):
        raise HTTPException(status_code=403, detail="Tehlikeli komut reddedildi")
    if not preset and not raw_cmd:
        raise HTTPException(status_code=400, detail="preset veya command gerekli")
    root = (body.get("workspace_root") or "").strip() or None
    scope = (body.get("scope_rel") or "").strip() or None
    active = (body.get("active_file") or "").strip() or None
    cmd_input = raw_cmd or preset
    result = run_fn(
        root,
        cmd_input,
        scope_rel=scope or None,
        active_file=active or None,
        message=str(body.get("message") or ""),
    )
    result["report"] = fmt(result)
    result["version"] = ver
    return result


@app.post("/api/programlama/shell/run")
def api_programlama_shell_run(body: dict):
    """Faz 67 — onaylı özgür shell (projects/ kapsamı)."""
    from ilim_assistant.motorlar.programlama_faz67 import (
        FAZ67_VERSION,
        approve_and_run_shell,
        format_shell_report,
        run_free_shell,
        stage_shell_request,
        validate_shell_command,
    )

    root = (body.get("workspace_root") or "").strip() or None
    scope = (body.get("scope_rel") or "").strip() or ""
    if scope and not scope.startswith("projects/"):
        scope = f"projects/{scope.lstrip('/')}"
    cmd = str(body.get("command") or "").strip()
    mode = str(body.get("mode") or body.get("action") or "onay").strip().lower()
    if not scope or not cmd:
        raise HTTPException(status_code=400, detail="scope_rel ve command gerekli")
    ok_cmd, err = validate_shell_command(cmd)
    if not ok_cmd:
        raise HTTPException(status_code=403, detail=err)
    if mode in ("istek", "request", "pending"):
        result = stage_shell_request(root, scope_rel=scope, command=cmd)
    elif mode in ("onayla", "approve"):
        token = str(body.get("token") or "").strip().lower() or None
        result = approve_and_run_shell(root, scope_rel=scope, token=token)
    else:
        result = run_free_shell(root, scope, cmd)
    result["report"] = format_shell_report(result)
    result["version"] = FAZ67_VERSION
    return result


@app.post("/api/programlama/patch/preview")
def api_programlama_patch_preview(body: ProgramlamaPatchBody):
    from ilim_assistant.motorlar.programlama_faz10 import (
        FAZ10_VERSION,
        format_patch_preview_report,
        preview_writes,
        stage_pending_from_text,
    )

    root = (body.workspace_root or "").strip() or None
    prev = preview_writes(body.text or "", root)
    try:
        from ilim_assistant.motorlar.programlama_faz16 import stage_pending_enriched

        stage_pending_enriched(body.text or "", root, source="api_preview")
    except Exception:
        from ilim_assistant.motorlar.programlama_faz10 import stage_pending_from_text

        stage_pending_from_text(body.text or "", root, source="api_preview")
    return {
        "ok": prev.get("ok"),
        "preview": prev,
        "report": format_patch_preview_report(prev),
        "version": FAZ10_VERSION,
    }


@app.get("/api/programlama/patch/pending")
def api_programlama_patch_pending(workspace_root: str | None = None):
    """Faz 16 — bekleyen patch + diff + durum."""
    from ilim_assistant.motorlar.programlama_faz16 import (
        FAZ16_VERSION,
        build_pending_bundle,
    )

    root = (workspace_root or "").strip() or None
    bundle = build_pending_bundle(root)
    items = list(bundle.get("items") or [])
    try:
        from ilim_assistant.motorlar.programlama_faz27 import enrich_pending_with_inline

        items = enrich_pending_with_inline(root, items)
    except Exception:
        pass
    try:
        from ilim_assistant.motorlar.programlama_faz45 import enrich_pending_items_v45

        items = enrich_pending_items_v45(root, items)
    except Exception:
        pass
    payload = {
        "ok": True,
        "count": bundle.get("count", 0),
        "paths": bundle.get("paths") or [],
        "items": items,
        "counts": bundle.get("counts") or {},
        "pending": bundle.get("pending") or {},
        "version": FAZ16_VERSION,
    }
    try:
        from ilim_assistant.motorlar.programlama_faz53 import patch_api_enrichments

        payload.update(patch_api_enrichments())
    except Exception:
        pass
    return payload


@app.get("/api/programlama/patch/inline-diff")
def api_programlama_patch_inline_diff(
    workspace_root: str | None = None,
    path: str | None = None,
):
    """Faz 27/45 — editör satır içi diff (v2 syntax renk)."""
    root = (workspace_root or "").strip() or None
    rel = (path or "").strip()
    if not rel:
        from ilim_assistant.motorlar.programlama_faz27 import FAZ27_VERSION

        return {"ok": False, "error": "path gerekli", "version": FAZ27_VERSION}
    try:
        from ilim_assistant.motorlar.programlama_faz45 import (
            FAZ45_VERSION,
            build_inline_diff_v2,
            editor_v2_enabled,
        )
        from ilim_assistant.motorlar.programlama_faz53 import resolve_editor_v2_for_api

        if editor_v2_enabled() or resolve_editor_v2_for_api(root):
            out = build_inline_diff_v2(root, rel)
            try:
                from ilim_assistant.motorlar.programlama_faz53 import patch_api_enrichments

                out.update(patch_api_enrichments())
            except Exception:
                pass
            out["editor_v2"] = True
            return out
    except Exception:
        pass
    from ilim_assistant.motorlar.programlama_faz27 import (
        FAZ27_VERSION,
        build_inline_diff_for_path,
    )

    return build_inline_diff_for_path(root, rel)


@app.get("/api/programlama/patch/tabs")
def api_programlama_patch_tabs(workspace_root: str | None = None):
    """Faz 45 — çok dosya patch sekmeleri."""
    from ilim_assistant.motorlar.programlama_faz45 import (
        FAZ45_VERSION,
        build_patch_tabs,
    )

    root = (workspace_root or "").strip() or None
    payload = build_patch_tabs(root)
    payload["version"] = FAZ45_VERSION
    try:
        from ilim_assistant.motorlar.programlama_faz53 import patch_api_enrichments

        payload.update(patch_api_enrichments())
    except Exception:
        pass
    return payload


@app.get("/api/programlama/patch/ux")
def api_programlama_patch_ux(workspace_root: str | None = None):
    """Faz 45 — patch + görev birleşik UX."""
    from ilim_assistant.motorlar.programlama_faz45 import (
        FAZ45_VERSION,
        build_unified_patch_ux,
    )

    root = (workspace_root or "").strip() or None
    payload = build_unified_patch_ux(root)
    payload["version"] = FAZ45_VERSION
    try:
        from ilim_assistant.motorlar.programlama_faz53 import patch_api_enrichments

        payload.update(patch_api_enrichments())
    except Exception:
        pass
    return payload


@app.post("/api/programlama/patch/unified-apply")
def api_programlama_patch_unified_apply(body: ProgramlamaPatchBody):
    """Faz 45 — tümünü kabul + kabul edilenleri uygula."""
    from ilim_assistant.motorlar.programlama_faz45 import (
        FAZ45_VERSION,
        run_unified_apply,
    )

    root = (body.workspace_root or "").strip() or None
    res = run_unified_apply(root, run_verify=bool(body.run_verify))
    return {**res, "version": FAZ45_VERSION}


@app.post("/api/programlama/patch/item")
def api_programlama_patch_item(body: ProgramlamaPatchBody):
    """Faz 16 — tek dosya kabul/red."""
    from ilim_assistant.motorlar.programlama_faz16 import (
        FAZ16_VERSION,
        STATUS_ACCEPTED,
        STATUS_REJECTED,
        build_pending_bundle,
        set_job_status,
    )

    root = (body.workspace_root or "").strip() or None
    path = (body.path or body.rel or "").strip()
    st = (body.status or "").strip().lower()
    if not path:
        raise HTTPException(status_code=400, detail="path gerekli")
    if st in ("accept", "accepted", "kabul"):
        st = STATUS_ACCEPTED
    elif st in ("reject", "rejected", "red"):
        st = STATUS_REJECTED
    else:
        raise HTTPException(status_code=400, detail="status: accepted veya rejected")
    res = set_job_status(root, path, st)
    if not res.get("ok"):
        raise HTTPException(status_code=404, detail=res.get("error") or "bulunamadı")
    bundle = build_pending_bundle(root)
    return {"ok": True, "item": res, "bundle": bundle, "version": FAZ16_VERSION}


@app.post("/api/programlama/patch/accept-all")
def api_programlama_patch_accept_all(body: ProgramlamaPatchBody):
    from ilim_assistant.motorlar.programlama_faz16 import (
        FAZ16_VERSION,
        accept_all_pending,
        build_pending_bundle,
    )

    root = (body.workspace_root or "").strip() or None
    res = accept_all_pending(root)
    bundle = build_pending_bundle(root)
    return {"ok": True, **res, "bundle": bundle, "version": FAZ16_VERSION}


@app.get("/api/programlama/quality-report")
def api_programlama_quality_report(workspace_root: str | None = None):
    """Faz 18 — SLO + orkestra birleşik kalite raporu."""
    from ilim_assistant.motorlar.programlama_faz18 import (
        FAZ18_VERSION,
        format_quality_report,
        run_offline_slo_scenarios,
    )

    root = (workspace_root or "").strip() or None
    report = run_offline_slo_scenarios(root)
    return {
        "ok": report.ok,
        "report": format_quality_report(report),
        "data": report.to_dict(),
        "version": FAZ18_VERSION,
    }


@app.get("/api/programlama/agent-compliance")
def api_programlama_agent_compliance(workspace_root: str | None = None):
    """Faz 37/48 — ajan uyum skoru (son görev oturumu)."""
    root = (workspace_root or "").strip() or None
    try:
        from ilim_assistant.motorlar.programlama_faz54 import (
            FAZ54_VERSION,
            build_compliance_report_v3,
            compliance_v3_enabled,
            format_compliance_report_v3,
        )

        if compliance_v3_enabled():
            data = build_compliance_report_v3(root)
            rep = data.get("report") or {}
            return {
                "ok": bool(data.get("ok")),
                "report": format_compliance_report_v3(root) if data.get("ok") else "",
                "data": rep,
                "target_score": rep.get("target_score", 85),
                "meets_target": bool(rep.get("meets_target")),
                "overall_kpi_ok": bool(rep.get("overall_kpi_ok")),
                "parity_passed": rep.get("parity_passed"),
                "parity_total": rep.get("parity_total"),
                "version": FAZ54_VERSION,
            }
    except Exception:
        pass
    try:
        from ilim_assistant.motorlar.programlama_faz48 import (
            FAZ48_VERSION,
            build_compliance_report_v2,
            compliance_v2_enabled,
            format_compliance_report_v2,
            target_compliance_score,
        )

        if compliance_v2_enabled():
            data = build_compliance_report_v2(root)
            rep = data.get("report") or {}
            return {
                "ok": bool(data.get("ok")),
                "report": format_compliance_report_v2(root) if data.get("ok") else "",
                "data": rep,
                "target_score": target_compliance_score(),
                "meets_target": bool(rep.get("meets_target")),
                "version": FAZ48_VERSION,
            }
    except Exception:
        pass
    from ilim_assistant.motorlar.programlama_faz37 import (
        FAZ37_VERSION,
        build_compliance_report,
        format_compliance_report,
    )

    data = build_compliance_report(root)
    rep = data.get("report") or {}
    return {
        "ok": bool(data.get("ok")),
        "report": format_compliance_report(root) if data.get("ok") else "",
        "data": rep,
        "version": FAZ37_VERSION,
    }


@app.get("/api/programlama/live-task-battery")
def api_programlama_live_task_battery(workspace_root: str | None = None):
    """Faz 86 — canlı görev pili (LLM'siz senaryo ölçümü)."""
    from ilim_assistant.motorlar.programlama_faz86 import (
        FAZ86_VERSION,
        format_live_battery_report,
        run_live_task_battery,
    )

    root = (workspace_root or "").strip() or None
    report = run_live_task_battery(root)
    return {
        "ok": bool(report.get("ok")),
        "report": report,
        "text": format_live_battery_report(report),
        "version": FAZ86_VERSION,
    }


@app.get("/api/programlama/agent-task-battery")
def api_programlama_agent_task_battery(
    workspace_root: str | None = None,
    live: int = 0,
):
    """Faz 88 — ajan görev pili (E1 verify→heal + isteğe canlı LLM)."""
    from ilim_assistant.motorlar.programlama_faz88 import (
        FAZ88_VERSION,
        format_agent_battery_report,
        run_agent_task_battery,
    )

    root = (workspace_root or "").strip() or None
    report = run_agent_task_battery(root, live_llm=bool(live))
    return {
        "ok": bool(report.get("ok")),
        "report": report,
        "text": format_agent_battery_report(report),
        "version": FAZ88_VERSION,
    }


@app.get("/api/programlama/combined-e1-battery")
def api_programlama_combined_e1_battery(
    workspace_root: str | None = None,
    live: int = 0,
):
    """Faz 86+88 — birleşik E1 pili."""
    from ilim_assistant.motorlar.programlama_faz88 import (
        FAZ88_VERSION,
        format_combined_e1_report,
        run_combined_e1_battery,
    )

    root = (workspace_root or "").strip() or None
    bundle = run_combined_e1_battery(root, live_llm=bool(live))
    return {
        "ok": bool(bundle.get("ok")),
        "report": bundle,
        "text": format_combined_e1_report(bundle),
        "version": FAZ88_VERSION,
    }


@app.post("/api/programlama/e1-maintenance")
def api_programlama_e1_maintenance(workspace_root: str | None = None):
    """Faz 91 — E1 bakım pili (birleşik pil + temiz KPI)."""
    from ilim_assistant.motorlar.programlama_faz91 import (
        FAZ91_VERSION,
        format_e1_maintenance_report,
        run_e1_maintenance,
    )

    root = (workspace_root or "").strip() or None
    report = run_e1_maintenance(root)
    return {
        "ok": bool(report.get("ok")),
        "report": report,
        "text": format_e1_maintenance_report(report),
        "version": FAZ91_VERSION,
    }


@app.post("/api/programlama/weekly-parity-full")
def api_programlama_weekly_parity_full(
    workspace_root: str | None = None,
    force: int = 0,
):
    """Faz 89 — haftalık parity full (E6)."""
    from ilim_assistant.motorlar.programlama_faz89 import (
        FAZ89_VERSION,
        format_weekly_parity_report,
        run_weekly_parity_battery,
    )

    root = (workspace_root or "").strip() or None
    report = run_weekly_parity_battery(root, force=bool(force))
    return {
        "ok": bool(report.get("ok")),
        "report": report,
        "text": format_weekly_parity_report(report),
        "version": FAZ89_VERSION,
    }


@app.get("/api/programlama/weakness-report")
def api_programlama_weakness_report(workspace_root: str | None = None):
    """Faz 82 — programlama zayıflık raporu (E1/E3/E6/E7)."""
    from ilim_assistant.motorlar.programlama_faz82 import (
        FAZ82_VERSION,
        build_weakness_report,
        format_weakness_report,
    )

    root = (workspace_root or "").strip() or None
    report = build_weakness_report(root)
    return {
        "ok": True,
        "report": report,
        "text": format_weakness_report(report),
        "version": FAZ82_VERSION,
    }


@app.get("/api/programlama/system-analysis")
def api_programlama_system_analysis(
    workspace_root: str | None = None,
    include_parity: int = 1,
):
    """Faz 96 (P11) — otonom sistem analizi."""
    from ilim_assistant.motorlar.programlama_faz96 import (
        FAZ96_VERSION,
        format_system_analysis_report,
        run_autonomous_system_analysis,
    )

    root = (workspace_root or "").strip() or None
    report = run_autonomous_system_analysis(
        root,
        include_parity=bool(int(include_parity or 0)),
    )
    return {
        "ok": True,
        "report": report,
        "text": format_system_analysis_report(report),
        "version": FAZ96_VERSION,
    }


@app.get("/api/programlama/local-workbench")
def api_programlama_local_workbench(workspace_root: str | None = None):
    """Blok I — yerel öncelik zinciri (E3, ollama-only, metin-only)."""
    from ilim_assistant.motorlar.programlama_local_workbench import (
        LOCAL_WORKBENCH_VERSION,
        build_local_workbench_payload,
    )

    root = (workspace_root or "").strip() or None
    payload = build_local_workbench_payload(root)
    payload["version"] = LOCAL_WORKBENCH_VERSION
    return payload


@app.get("/api/programlama/handoff-workbench")
def api_programlama_handoff_workbench(
    workspace_root: str | None = None,
    message: str | None = None,
    active_file: str | None = None,
):
    """Blok H — Ana Motor ↔ Programlama handoff v4 özeti."""
    from ilim_assistant.motorlar.programlama_handoff_workbench import (
        HANDOFF_WORKBENCH_VERSION,
        build_handoff_workbench_payload,
    )

    root = (workspace_root or "").strip() or None
    payload = build_handoff_workbench_payload(
        root,
        message=(message or "").strip(),
        active_file=(active_file or "").strip() or None,
    )
    payload["version"] = HANDOFF_WORKBENCH_VERSION
    return payload


@app.get("/api/programlama/mega-workbench")
def api_programlama_mega_workbench(workspace_root: str | None = None):
    """Blok G — mega görev özeti (tur, touched_files, patch, verify, E2)."""
    from ilim_assistant.motorlar.programlama_mega_workbench import (
        MEGA_WORKBENCH_VERSION,
        build_mega_workbench_payload,
    )

    root = (workspace_root or "").strip() or None
    payload = build_mega_workbench_payload(root)
    payload["version"] = MEGA_WORKBENCH_VERSION
    return payload


@app.get("/api/programlama/task-stats")
def api_programlama_task_stats(workspace_root: str | None = None):
    """Faz 55 — canlı görev başarı istatistikleri."""
    from ilim_assistant.motorlar.programlama_faz55 import (
        FAZ55_VERSION,
        compute_task_stats,
        format_task_stats_report,
    )

    root = (workspace_root or "").strip() or None
    stats = compute_task_stats(root, window_days=30)
    ver = FAZ55_VERSION
    try:
        from ilim_assistant.motorlar.programlama_faz102_e1_live import FAZ102_VERSION

        ver = f"{FAZ55_VERSION}+{FAZ102_VERSION}"
    except Exception:
        pass
    return {
        "ok": True,
        "stats": stats,
        "report": format_task_stats_report(stats),
        "version": ver,
    }


@app.get("/api/programlama/pr-plan")
def api_programlama_pr_plan(
    workspace_root: str | None = None,
    scope_rel: str | None = None,
    title: str | None = None,
):
    """Faz 83 — PR hazırlık (onaylı öneri; otomatik PR açmaz)."""
    from ilim_assistant.motorlar.programlama_faz83 import (
        FAZ83_VERSION,
        build_pr_plan,
        format_pr_plan,
    )

    root = (workspace_root or "").strip() or None
    scope = (scope_rel or "").strip() or None
    plan = build_pr_plan(root, title_hint=(title or "").strip(), scope_rel=scope)
    return {
        "ok": bool(plan.get("ok")),
        "plan": plan,
        "report": format_pr_plan(plan),
        "version": FAZ83_VERSION,
    }


@app.get("/api/programlama/weekly-kpi")
def api_programlama_weekly_kpi(workspace_root: str | None = None):
    """Faz 60 — haftalık KPI raporu (JSON)."""
    from ilim_assistant.motorlar.programlama_faz60 import (
        FAZ60_VERSION,
        format_weekly_kpi_report_text,
        generate_weekly_kpi_report,
    )

    root = (workspace_root or "").strip() or None
    report = generate_weekly_kpi_report(root)
    report["version"] = FAZ60_VERSION
    report["report_text"] = format_weekly_kpi_report_text(report)
    return report


@app.post("/api/programlama/ci-run")
def api_programlama_ci_run(
    workspace_root: str | None = None,
    force_full_parity: bool = False,
):
    """Faz 60 — CI otomasyon (parity quick + haftalık full)."""
    from ilim_assistant.motorlar.programlama_faz60 import (
        FAZ60_VERSION,
        run_ci_automation,
    )

    root = (workspace_root or "").strip() or None
    out = run_ci_automation(root, force_full_parity=bool(force_full_parity))
    out["version"] = FAZ60_VERSION
    return out


@app.get("/api/rok/kpi")
def api_rok_kpi():
    """Faz 77 — ROK motor KPI özeti (offline bayraklar)."""
    from ilim_assistant.motorlar.ruzgar_cila_faz77 import FAZ77_VERSION, collect_rok_kpi

    return {"ok": True, "kpi": collect_rok_kpi(), "version": FAZ77_VERSION}


@app.get("/api/ana-motor/hub-route")
def api_ana_motor_hub_route(message: str = "", workspace_root: str | None = None):
    """Faz 76 — Ana Motor hub: mesaj için hedef motor önizlemesi."""
    from ilim_assistant.idrak_entegrasyon import motor_niyeti_heuristic
    from ilim_assistant.motorlar.ana_motor_hub_faz76 import (
        FAZ76_VERSION,
        motor_label,
        resolve_hub_target,
    )

    msg = (message or "").strip()
    root = (workspace_root or "").strip() or None
    flags = motor_niyeti_heuristic(msg) if msg else {}
    target, meta = resolve_hub_target(msg, flags, workspace_root=root)
    return {
        "ok": True,
        "target": target,
        "target_label": motor_label(target),
        "meta": meta,
        "version": FAZ76_VERSION,
    }


@app.get("/api/ana-motor/learned-instant")
def api_ana_motor_learned_instant(message: str = "", workspace_root: str | None = None):
    """Faz C — eylem öğret / listele / sil anında yanıt."""
    from ilim_assistant.motorlar.motor_ogrenilen_eylemler import (
        FAZ_C_VERSION,
        try_instant_learned_commands,
    )

    msg = (message or "").strip()
    root = (workspace_root or "").strip() or None
    reply = try_instant_learned_commands(msg, root) if msg else None
    return {
        "ok": True,
        "handled": bool(reply),
        "reply": reply,
        "version": FAZ_C_VERSION,
    }


@app.get("/api/ana-motor/learned-actions")
def api_ana_motor_learned_actions(workspace_root: str | None = None):
    """Faz C — öğrenilen eylem listesi."""
    from ilim_assistant.motorlar.motor_ogrenilen_eylemler import learned_actions_snapshot

    root = (workspace_root or "").strip() or None
    return learned_actions_snapshot(root)


class LearnedActionBody(BaseModel):
    trigger: str = ""
    motor: str = ""
    sub_intent: str | None = None
    workspace_root: str | None = None
    approved: bool | None = None
    note: str = ""


@app.post("/api/ana-motor/learned-action")
def api_ana_motor_learned_action_post(body: LearnedActionBody):
    """Faz C — eylem ekle (API / onaylı otomasyon)."""
    from ilim_assistant.motorlar.motor_ogrenilen_eylemler import (
        FAZ_C_VERSION,
        get_learned_store,
    )

    root = (body.workspace_root or "").strip() or None
    try:
        row = get_learned_store(root).add_action(
            body.trigger,
            body.motor,
            sub_intent=body.sub_intent,
            approved=body.approved,
            note=(body.note or "").strip(),
        )
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    return {"ok": True, "action": row, "version": FAZ_C_VERSION}


@app.delete("/api/ana-motor/learned-action/{action_id}")
def api_ana_motor_learned_action_delete(
    action_id: str,
    workspace_root: str | None = None,
):
    """Faz C — eylem sil (panel / API)."""
    from ilim_assistant.motorlar.motor_ogrenilen_eylemler import (
        FAZ_C_VERSION,
        get_learned_store,
    )

    root = (workspace_root or "").strip() or None
    ok = get_learned_store(root).delete_by_id(action_id)
    if not ok:
        raise HTTPException(status_code=404, detail="Eylem bulunamadı.")
    return {"ok": True, "deleted": action_id, "version": FAZ_C_VERSION}


@app.get("/api/ana-motor/sub-intents")
def api_ana_motor_sub_intents():
    """Faz C — geçerli motor/alt-intent kataloğu."""
    from ilim_assistant.motorlar.motor_ogrenilen_eylemler import sub_intents_catalog

    return sub_intents_catalog()


@app.get("/api/ana-motor/motor-dispatch")
def api_ana_motor_motor_dispatch(
    message: str = "",
    target: str = "",
    workspace_root: str | None = None,
):
    """Faz A — hedef motorda anında komut/yanıt (Ana Motor delegasyon)."""
    from ilim_assistant.motorlar.ana_motor_hub_faz76 import build_motor_dispatch_payload

    msg = (message or "").strip()
    tgt = (target or "").strip().lower()
    if not msg or not tgt:
        return {
            "ok": False,
            "handled": False,
            "error": "message ve target gerekli.",
        }
    root = (workspace_root or "").strip() or None
    return build_motor_dispatch_payload(msg, tgt, workspace_root=root)


@app.get("/api/ana-motor/capabilities")
def api_ana_motor_capabilities():
    """Faz B — merkezi motor kabiliyet kaydı özeti."""
    from ilim_assistant.motorlar.motor_kabiliyetleri import registry_snapshot

    return registry_snapshot()


@app.get("/api/ana-motor/delegation-summary")
def api_ana_motor_delegation_summary(workspace_root: str | None = None):
    """Faz 59 — son delege özeti."""
    from ilim_assistant.ana_motor_faz59 import (
        FAZ59_VERSION,
        format_delegation_summary_text,
        load_last_delegation_summary,
    )

    root = (workspace_root or "").strip() or None
    summ = load_last_delegation_summary(root)
    return {
        "ok": summ is not None,
        "summary": summ,
        "report": format_delegation_summary_text(summ or {}),
        "version": FAZ59_VERSION,
    }


@app.get("/api/programlama/git-changes")
def api_programlama_git_changes(
    workspace_root: str | None = None,
    scope_rel: str | None = None,
    active_file: str | None = None,
):
    """Faz 58 — git status/diff özeti (read-only)."""
    from ilim_assistant.motorlar.programlama_faz58 import (
        FAZ58_VERSION,
        build_git_changes_api_payload,
    )

    root = (workspace_root or "").strip() or None
    payload = build_git_changes_api_payload(
        root,
        scope_rel=(scope_rel or "").strip() or None,
        active_file=(active_file or "").strip() or None,
    )
    payload["version"] = FAZ58_VERSION
    try:
        from ilim_assistant.motorlar.programlama_faz62 import enrich_git_changes_api_payload

        payload = enrich_git_changes_api_payload(
            payload,
            root,
            scope_rel=payload.get("scope_rel") or (scope_rel or "").strip() or None,
        )
    except Exception:
        pass
    return payload


@app.get("/api/programlama/kpi-dashboard")
def api_programlama_kpi_dashboard(workspace_root: str | None = None):
    """Faz 54 — KPI dashboard (compliance v3 + parity 8/8)."""
    from ilim_assistant.motorlar.programlama_faz54 import (
        FAZ54_VERSION,
        build_kpi_dashboard,
    )

    root = (workspace_root or "").strip() or None
    payload = build_kpi_dashboard(root)
    payload["version"] = FAZ54_VERSION
    try:
        from ilim_assistant.motorlar.programlama_faz55 import compute_task_stats

        payload["task_stats"] = compute_task_stats(root, window_days=30)
    except Exception:
        pass
    try:
        from ilim_assistant.motorlar.programlama_faz57 import compute_text_only_stats

        payload["text_only_stats"] = compute_text_only_stats(root, window_days=7)
    except Exception:
        pass
    try:
        from ilim_assistant.motorlar.programlama_faz63 import enrich_kpi_dashboard

        payload = enrich_kpi_dashboard(payload, root)
    except Exception:
        pass
    return payload


@app.get("/api/programlama/umit-onay")
def api_programlama_umit_onay(workspace_root: str | None = None):
    """Faz 98 — Ümit onay kapısı bekleyen işlem özeti."""
    from ilim_assistant.motorlar.programlama_faz98 import FAZ98_VERSION, load_pending

    root = (workspace_root or "").strip() or None
    pending = load_pending(root)
    if not pending:
        return {"ok": True, "has_pending": False, "version": FAZ98_VERSION}

    op = pending.get("operation") or pending
    preview = op.get("preview") or {}
    risks = list(preview.get("risks") or [])
    return {
        "ok": True,
        "has_pending": True,
        "token": pending.get("token"),
        "created_at": pending.get("created_at"),
        "expires_at": pending.get("expires_at"),
        "operation": {
            "kind": op.get("kind"),
            "label": op.get("label"),
            "src": op.get("src_abs") or op.get("src"),
            "dst": op.get("dst_abs") or op.get("dst"),
            "command": op.get("command"),
            "package": op.get("package"),
            "scope_rel": op.get("scope_rel"),
            "success": preview.get("success"),
            "risk_count": len(risks),
            "risk_head": risks[0] if risks else "",
        },
        "version": FAZ98_VERSION,
    }


@app.post("/api/programlama/best-of-n/plan")
def api_programlama_best_of_n_plan(body: ProgramlamaPatchBody):
    """Faz 64 — Best-of-N aday planı + pytest skoru."""
    from ilim_assistant.motorlar.programlama_faz64 import (
        FAZ64_VERSION,
        format_best_of_n_report,
        plan_best_of_n_run,
        parse_best_of_n_command,
    )

    root = (body.workspace_root or "").strip() or None
    text = (body.text or "").strip()
    parsed = parse_best_of_n_command(text) if text else None
    n = 2
    scope = (body.rel or "").strip()
    goal = text
    if parsed:
        n = parsed["n"]
        scope = parsed["scope_rel"]
        goal = parsed["goal"]
    if scope and not scope.startswith("projects/"):
        scope = f"projects/{scope}"
    manifest = plan_best_of_n_run(
        root,
        scope_rel=scope or "projects/unknown",
        goal=goal or "best-of-n",
        n=n,
    )
    manifest["version"] = FAZ64_VERSION
    manifest["report"] = format_best_of_n_report(manifest)
    return manifest


@app.get("/api/programlama/best-of-n/runs")
def api_programlama_best_of_n_runs(workspace_root: str | None = None):
    """Faz 64 — son Best-of-N koşuları."""
    from ilim_assistant.motorlar.programlama_faz64 import (
        FAZ64_VERSION,
        list_best_of_n_runs,
    )

    root = (workspace_root or "").strip() or None
    out = list_best_of_n_runs(root)
    out["version"] = FAZ64_VERSION
    return out


@app.post("/api/programlama/best-of-n/agents")
def api_programlama_best_of_n_agents(body: ProgramlamaPatchBody):
    """Faz 65 — mevcut Best-of-N koşusunda her adayda kod ajanı + yeniden skor."""
    from ilim_assistant.motorlar.programlama_faz65 import (
        FAZ65_VERSION,
        execute_best_of_n_agents,
        parse_best_of_n_agent_command,
        run_best_of_n_plus,
        parse_best_of_n_plus_command,
    )

    root = (body.workspace_root or "").strip() or None
    text = (body.text or "").strip()
    merge = os.environ.get("RUZGAR_FAZ65_MERGE", "").strip().lower() in (
        "1",
        "true",
        "yes",
    )
    plus = parse_best_of_n_plus_command(text) if text else None
    if plus:
        out = run_best_of_n_plus(
            root,
            scope_rel=plus["scope_rel"],
            goal=plus["goal"],
            n=plus["n"],
            merge=merge,
        )
    else:
        parsed = parse_best_of_n_agent_command(text) if text else None
        run_id = (parsed or {}).get("run_id") or (body.rel or "").strip()
        if not run_id:
            return {"ok": False, "error": "run_id gerekli (best-of-n-agent: bon-...)"}
        out = execute_best_of_n_agents(root, run_id, merge=merge)
    out["version"] = FAZ65_VERSION
    return out


@app.get("/api/programlama/live-kpi")
def api_programlama_live_kpi(workspace_root: str | None = None):
    """Faz 63 — canlı görev KPI (7/30 gün trend)."""
    from ilim_assistant.motorlar.programlama_faz63 import (
        FAZ63_VERSION,
        compute_live_kpi,
        format_live_kpi_report,
    )

    root = (workspace_root or "").strip() or None
    live = compute_live_kpi(root)
    return {
        "ok": bool(live.get("ok")),
        "live_kpi": live,
        "report": format_live_kpi_report(live),
        "version": FAZ63_VERSION,
    }


@app.get("/api/programlama/parity-report")
def api_programlama_parity_report(workspace_root: str | None = None):
    """Faz 25 — Cursor parity smoke raporu."""
    from ilim_assistant.motorlar.programlama_faz25 import (
        FAZ25_VERSION,
        format_parity_report,
        run_offline_parity_scenario,
    )

    root = (workspace_root or "").strip() or None
    report = run_offline_parity_scenario(root)
    return {
        "ok": report.ok,
        "report": format_parity_report(report),
        "data": report.to_dict(),
        "version": FAZ25_VERSION,
    }


@app.get("/api/programlama/proje-uret")
def api_programlama_proje_uret(
    template_id: str = "fastapi_api",
    project_name: str = "",
    goal: str = "",
    features: str = "",
    workspace_root: str | None = None,
):
    """Faz 47–50 — bağımsız proje üret (scaffold + offline bootstrap + verify)."""
    from ilim_assistant.motorlar.programlama_faz47 import (
        FAZ47_VERSION,
        ProjeUretSpec,
        format_proje_uret_report,
        run_proje_uret_prepare,
    )
    from ilim_assistant.motorlar.programlama_faz50 import (
        extract_features_list,
        merge_goal_with_features,
    )

    name = (project_name or "").strip() or f"api-{int(__import__('time').time()) % 100000}"
    feat_list = extract_features_list(features or goal or "")
    goal_merged = merge_goal_with_features(
        (goal or "").strip() or "health version pytest geçir",
        feat_list,
    )
    spec = ProjeUretSpec(
        template_id=(template_id or "fastapi_api").strip(),
        project_name=name,
        goal=goal_merged,
        features=feat_list,
        source="api",
    )
    root = (workspace_root or "").strip() or None
    rep = run_proje_uret_prepare(root, spec)
    agent_message = None
    if rep.agent_required and rep.scaffold_ok:
        from ilim_assistant.motorlar.programlama_faz47 import build_agent_message_for_spec

        agent_message = build_agent_message_for_spec(spec)
    return {
        "ok": rep.ok,
        "ready_without_agent": rep.ready_without_agent,
        "agent_required": rep.agent_required,
        "agent_message": agent_message,
        "report": format_proje_uret_report(rep),
        "data": rep.to_dict(),
        "version": FAZ47_VERSION,
    }


@app.get("/api/programlama/proje-uret/templates")
def api_programlama_proje_uret_templates():
    """Faz 47 UI — şablon listesi."""
    from ilim_assistant.motorlar.programlama_faz6 import list_templates
    from ilim_assistant.motorlar.programlama_faz47 import FAZ47_VERSION, proje_uret_enabled

    return {
        "ok": True,
        "enabled": proje_uret_enabled(),
        "templates": list_templates(),
        "version": FAZ47_VERSION,
    }


@app.get("/api/programlama/cursor-seviye")
def api_programlama_cursor_seviye(workspace_root: str | None = None):
    """Faz 46 — Cursor seviye kilidi (3 senaryo + skor)."""
    from ilim_assistant.motorlar.programlama_cursor_parity import (
        CURSOR_PARITY_VERSION,
        format_cursor_seviye_report,
        run_cursor_seviye_assessment,
    )

    root = (workspace_root or "").strip() or None
    report = run_cursor_seviye_assessment(root)
    return {
        "ok": report.ok,
        "score": report.score,
        "target_score": report.target_score,
        "meets_target": report.score >= report.target_score,
        "report": format_cursor_seviye_report(report),
        "data": report.to_dict(),
        "warnings": report.warnings,
        "version": CURSOR_PARITY_VERSION,
    }


@app.get("/api/programlama/git/status")
def api_programlama_git_status(
    workspace_root: str | None = None,
    scope_rel: str | None = None,
    active_file: str | None = None,
):
    """Faz 17 — git status -sb + diff --stat özeti."""
    from ilim_assistant.motorlar.programlama_faz17 import (
        FAZ17_VERSION,
        format_git_status_report,
        gather_git_snapshot,
        resolve_scope_rel,
    )

    root = (workspace_root or "").strip() or None
    scope = (scope_rel or "").strip() or resolve_scope_rel(
        root, active_file=active_file
    )
    snap = gather_git_snapshot(root, scope_rel=scope)
    return {
        "ok": bool(snap.get("ok")),
        "snapshot": snap,
        "report": format_git_status_report(snap),
        "version": FAZ17_VERSION,
    }


@app.post("/api/programlama/git/suggest-commit")
def api_programlama_git_suggest_commit(body: ProgramlamaPatchBody):
    """Faz 17 — commit mesajı öner (heuristic + isteğe bağlı LLM)."""
    from ilim_assistant.motorlar.programlama_faz17 import (
        FAZ17_VERSION,
        format_commit_suggest_report,
        resolve_scope_rel,
        suggest_commit_message,
    )

    root = (body.workspace_root or "").strip() or None
    scope = resolve_scope_rel(root, active_file=body.rel)
    res = suggest_commit_message(
        root,
        scope_rel=scope,
        active_file=body.rel,
        message=body.text or "",
        user_hint=body.text or "",
    )
    res["report"] = format_commit_suggest_report(res)
    res["version"] = FAZ17_VERSION
    return res


@app.get("/api/programlama/git/pending-commit")
def api_programlama_git_pending_commit(workspace_root: str | None = None):
    from ilim_assistant.motorlar.programlama_faz17 import FAZ17_VERSION, load_pending_commit

    root = (workspace_root or "").strip() or None
    pending = load_pending_commit(root)
    return {
        "ok": bool(pending.get("message")),
        "pending": pending,
        "version": FAZ17_VERSION,
    }


@app.get("/api/programlama/git/pr-status")
def api_programlama_git_pr_status(
    workspace_root: str | None = None,
    scope_rel: str | None = None,
    active_file: str | None = None,
):
    """Faz 31 — PR özeti (dal, ahead/behind, açık PR)."""
    from ilim_assistant.motorlar.programlama_faz31 import (
        FAZ31_VERSION,
        format_pr_status_report,
        gather_pr_snapshot,
    )

    root = (workspace_root or "").strip() or None
    snap = gather_pr_snapshot(
        root,
        scope_rel=(scope_rel or "").strip() or None,
        active_file=active_file,
    )
    return {
        "ok": bool(snap.get("ok")),
        "snapshot": snap,
        "report": format_pr_status_report(snap),
        "version": FAZ31_VERSION,
    }


class ProgramlamaPrBody(BaseModel):
    workspace_root: str | None = None
    scope_rel: str | None = None
    rel: str | None = None
    title: str = ""
    body: str = ""
    push_first: bool = True


@app.post("/api/programlama/git/push-branch")
def api_programlama_git_push_branch(body: ProgramlamaPrBody):
    """Faz 31 — güvenli git push -u origin HEAD."""
    from ilim_assistant.motorlar.programlama_faz31 import FAZ31_VERSION, push_current_branch

    root = (body.workspace_root or "").strip() or None
    res = push_current_branch(
        root,
        scope_rel=body.scope_rel,
        active_file=body.rel,
    )
    if not res.get("ok"):
        raise HTTPException(status_code=400, detail=res.get("error") or "push başarısız")
    return {**res, "version": FAZ31_VERSION}


@app.post("/api/programlama/task/save-pipeline")
def api_programlama_task_save_pipeline(body: ProgramlamaPrBody):
    """Faz 32 — commit + push + PR (görev kaydet)."""
    from ilim_assistant.motorlar.programlama_faz32 import (
        FAZ32_VERSION,
        format_pipeline_report,
        run_task_save_pipeline,
    )

    root = (body.workspace_root or "").strip() or None
    res = run_task_save_pipeline(
        root,
        scope_rel=body.scope_rel,
        active_file=body.rel,
        pr_title=body.title,
        open_pr=True,
    )
    res["report"] = format_pipeline_report(res)
    res["version"] = FAZ32_VERSION
    if not res.get("ok"):
        raise HTTPException(status_code=400, detail=res.get("error") or "pipeline başarısız")
    return res


@app.post("/api/programlama/git/pr-create")
def api_programlama_git_pr_create(body: ProgramlamaPrBody):
    """Faz 31 — gh pr create (önce push)."""
    from ilim_assistant.motorlar.programlama_faz31 import (
        FAZ31_VERSION,
        create_pull_request,
        format_pr_create_report,
    )

    root = (body.workspace_root or "").strip() or None
    res = create_pull_request(
        root,
        title=body.title,
        body=body.body,
        push_first=body.push_first,
        scope_rel=body.scope_rel,
        active_file=body.rel,
    )
    res["report"] = format_pr_create_report(res)
    res["version"] = FAZ31_VERSION
    if not res.get("ok"):
        raise HTTPException(status_code=400, detail=res.get("error") or "PR oluşturulamadı")
    return res


@app.post("/api/programlama/git/commit")
def api_programlama_git_commit(body: ProgramlamaPatchBody):
    """Faz 17 — onaylı git commit (git add -A + commit -m; --no-verify yasak)."""
    from ilim_assistant.motorlar.programlama_faz17 import (
        FAZ17_VERSION,
        execute_pending_commit,
        format_commit_done_report,
        stage_commit_message,
    )

    root = (body.workspace_root or "").strip() or None
    custom = (body.text or body.status or "").strip()
    if custom and not body.run_verify:
        stage_commit_message(
            root,
            custom,
            active_file=body.rel,
        )
    res = execute_pending_commit(root, commit_message=custom or None)
    res["report"] = format_commit_done_report(res)
    res["version"] = FAZ17_VERSION
    return res


@app.post("/api/programlama/patch/rollback")
def api_programlama_patch_rollback(body: ProgramlamaPatchBody):
    from ilim_assistant.motorlar.programlama_faz16 import (
        FAZ16_VERSION,
        rollback_last_applied,
    )

    root = (body.workspace_root or "").strip() or None
    res = rollback_last_applied(root)
    res["version"] = FAZ16_VERSION
    return res


@app.post("/api/programlama/patch/apply")
def api_programlama_patch_apply(body: ProgramlamaPatchBody):
    from ilim_assistant.motorlar.programlama_faz10 import resolve_scope_rel
    from ilim_assistant.motorlar.programlama_faz16 import (
        FAZ16_VERSION,
        apply_pending_selective,
        format_apply_report_v16,
    )

    root = (body.workspace_root or "").strip() or None
    scope = resolve_scope_rel(root, active_file=body.rel)
    mode = (body.mode or "accepted").strip().lower() or "accepted"
    res = apply_pending_selective(
        root,
        mode=mode,
        run_verify=bool(body.run_verify),
        scope_rel=scope,
    )
    res["report"] = format_apply_report_v16(res)
    res["version"] = FAZ16_VERSION
    return res


@app.get("/api/programlama/project-tree")
def api_programlama_project_tree(
    workspace_root: str | None = None,
    project_rel: str | None = None,
):
    """Faz 8 — projects/<ad>/ altı güvenli dosya listesi."""
    from ilim_assistant.motorlar.programlama_faz8 import FAZ8_VERSION, list_project_tree

    root = (workspace_root or "").strip() or None
    proj = (project_rel or "").strip()
    if not proj:
        raise HTTPException(status_code=400, detail="project_rel gerekli")
    out = list_project_tree(root, proj)
    out["version"] = FAZ8_VERSION
    return out


@app.post("/api/programlama/serve")
def api_programlama_serve(body: ProgramlamaRunBody):
    """Faz 8 — FastAPI arka plan (başlat/durdur). action=start|stop|status."""
    from ilim_assistant.motorlar.programlama_faz7 import resolve_target_rel
    from ilim_assistant.motorlar.programlama_faz8 import (
        FAZ8_VERSION,
        format_serve_status_report,
        serve_status,
        start_background_api,
        stop_background_api,
    )

    root = (body.workspace_root or "").strip() or None
    action = (body.action or "start").strip().lower()
    rel = (body.rel or "").strip() or None
    if not rel:
        rel = resolve_target_rel("", workspace_root=root)
    if action == "status":
        st = serve_status(root, rel or None)
        st["report"] = format_serve_status_report(root) or ""
        st["version"] = FAZ8_VERSION
        return st
    if not rel:
        raise HTTPException(status_code=400, detail="rel veya açık oturum dosyası gerekli.")
    if action in ("stop", "durdur"):
        result = stop_background_api(root, rel)
    else:
        result = start_background_api(root, rel)
    result["version"] = FAZ8_VERSION
    return result


@app.get("/api/system-health-report")
def api_system_health_report(mode: str = "genel"):
    """Kısa kişisel asistan sağlık raporu: aktif mod + hafıza kapasitesi."""
    from ilim_assistant.ruzgar_session_context import memory_capacity_snapshot
    from ilim_assistant.ruzgar_ui_manifest import build_ui_manifest

    mode_norm = normalize_mode(mode)
    h = health()
    manifest = build_ui_manifest(health=h)
    memory = memory_capacity_snapshot()
    super_brain = h.get("super_brain") or {}
    ana_motor = h.get("ana_motor") or {}
    lines = [
        "Sistem Sağlık Raporu",
        f"- Aktif mod: {mode_norm}",
        f"- Faz: {manifest.get('current_phase_label')}",
        f"- Gemini: {'hazır' if super_brain.get('gemini_configured') else 'kapalı/eksik'}",
        f"- Ana Motor ajan: {'aktif' if ana_motor.get('ana_motor_agent_enabled') else 'kapalı'}",
        (
            "- Hafıza kapasitesi: "
            f"{memory.get('hafiza_total', 0)} kalıcı kayıt, "
            f"{memory.get('personal_notes', 0)} kişisel not, "
            f"{memory.get('shared_sqlite_context', 0)} SQLite bağlam, "
            f"{memory.get('active_tasks', 0)} aktif görev, "
            f"{memory.get('pending_reminders', 0)} bekleyen hatırlatıcı"
        ),
        f"- Oturum hafıza context'i: {'aktif' if memory.get('session_memory_context_enabled') else 'kapalı'}",
    ]
    try:
        from ilim_assistant.ruzgar_debug_report import last_debug_report

        dbg = last_debug_report()
    except Exception:
        dbg = None
    if dbg:
        lines.append(f"- Son debug: exit={dbg.get('exit_code')} · {dbg.get('likely_cause')}")
    else:
        lines.append("- Son debug: kayıt yok")

    return {
        "ok": True,
        "mode": mode_norm,
        "phase": manifest.get("current_phase_label"),
        "health": h,
        "memory": memory,
        "report": "\n".join(lines),
    }


@app.get("/api/tasks")
def api_tasks(limit: int = 50):
    from ilim_assistant.gorev_yoneticisi import list_tasks

    return {"ok": True, "items": list_tasks(limit)}


@app.post("/api/tasks")
def api_tasks_create(body: TaskCreateBody):
    from ilim_assistant.gorev_yoneticisi import create_task

    return {"ok": True, "task": create_task(body.title, body.detail)}


@app.post("/api/tasks/update")
def api_tasks_update(body: TaskUpdateBody):
    from ilim_assistant.gorev_yoneticisi import update_task

    ok = update_task(body.id, body.status, body.detail)
    return {"ok": ok, "id": body.id}


@app.get("/api/merkezi-bellek")
@app.get("/api/merkezi-bellek/")
def api_merkezi_bellek():
    """Merkezi Bellek v3 — Ticaret Avcısı senkronu + HIZIR ticaret + genel önbellek."""
    try:
        from ilim_assistant.hizir.bellek import load_merkezi_bellek, merkezi_bellek_path
        from ilim_assistant.hizir.ticaret_avci import reconcile_ticaret_avci_firsatlar

        data = reconcile_ticaret_avci_firsatlar()
        return {
            "ok": True,
            "path": str(merkezi_bellek_path()),
            "schema": data.get("schema"),
            "version": data.get("version"),
            "data": data,
        }
    except Exception as exc:
        raise HTTPException(status_code=500, detail=str(exc)) from exc


@app.post("/api/hizir/pazar-tara")
def api_hizir_pazar_tara(body: HizirPazarTaraBody):
    """Panelden canlı keşif: Tool-Bridge + UniversalScraper; sonuç `merkezi_bellek`e yazılır."""
    try:
        from ilim_assistant.hizir.tool_bridge import build_dynamic_operasyon_context

        q = (body.query or "").strip()
        msg = f"Pazar yerini tara: {q}" if q else "Pazar yerini tara"
        ctx = build_dynamic_operasyon_context(
            msg,
            weather_q=False,
            has_live_weather_block=False,
            mode_norm="hizir",
            pazar_kanallari=body.channels,
        )
        from ilim_assistant.hizir.ticaret_avci import reconcile_ticaret_avci_firsatlar

        data = reconcile_ticaret_avci_firsatlar()
        return {
            "ok": True,
            "tool_context_chars": len(ctx),
            "data": data,
        }
    except Exception as exc:
        raise HTTPException(status_code=500, detail=str(exc)) from exc


@app.post("/api/hizir/temizle")
def api_hizir_temizle():
    """HIZIR vitrin: otomatik fırsatlar + pazar önbelleği + pas geç listesi temizlenir."""
    try:
        from ilim_assistant.hizir.bellek import clear_hizir_vitrin_state, merkezi_bellek_path

        data = clear_hizir_vitrin_state()
        return {"ok": True, "path": str(merkezi_bellek_path()), "data": data}
    except Exception as exc:
        raise HTTPException(status_code=500, detail=str(exc)) from exc


@app.post("/api/hizir/firsat-kaldir")
def api_hizir_firsat_kaldir(body: HizirFirsatKaldirBody):
    """PAS GEÇ: kartı pas_gecildi listesine alır ve listeyi yeniden üretir."""
    try:
        from ilim_assistant.hizir.bellek import merkezi_bellek_path, pas_gec_hizir_kart

        data = pas_gec_hizir_kart((body.kart_id or "").strip())
        return {"ok": True, "path": str(merkezi_bellek_path()), "data": data}
    except Exception as exc:
        raise HTTPException(status_code=500, detail=str(exc)) from exc


class VideoSearchBody(BaseModel):
    message: str = ""
    query: str = ""
    workspace_root: str = ""


@app.post("/api/video/search")
async def api_video_search(body: VideoSearchBody):
    """Faz 84 — YouTube isimle arama veya «N numarayı indir»."""
    msg = (body.message or body.query or "").strip()
    if not msg:
        raise HTTPException(status_code=400, detail="Arama metni boş.")
    ws = (body.workspace_root or "").strip() or None

    def _run() -> dict[str, Any]:
        from ilim_assistant.motorlar.video_faz84 import (
            FAZ84_VERSION,
            extract_search_query,
            format_search_results,
            maybe_instant_faz84,
            maybe_search_and_pick,
            search_youtube,
        )

        pick = maybe_search_and_pick(msg, ws)
        if pick:
            return {"ok": True, "mode": "download", "text": pick}
        instant = maybe_instant_faz84(msg, ws)
        if instant:
            return {"ok": True, "mode": "instant", "text": instant}
        q = extract_search_query(msg)
        data = search_youtube(q)
        if data.get("ok"):
            from ilim_assistant.motorlar.video_faz84 import _save_last_search

            _save_last_search(data, ws)
        return {
            "ok": bool(data.get("ok")),
            "mode": "search",
            "data": data,
            "text": format_search_results(data),
            "version": FAZ84_VERSION,
        }

    try:
        data = await run_in_threadpool(_run)
        return data
    except HTTPException:
        raise
    except Exception as exc:
        raise HTTPException(status_code=500, detail=str(exc)) from exc


@app.post("/api/video/download")
async def api_video_download(body: VideoDownloadBody):
    """YouTube/web video — doğrudan indirme kapalı; virus-guard/preflight (kind=video) kullanın."""
    from ilim_assistant.motorlar.ruzgar_virus_guard import guard_enabled

    url = (body.url or "").strip()
    if not url:
        raise HTTPException(status_code=400, detail="URL boş.")
    if guard_enabled():
        raise HTTPException(
            status_code=403,
            detail=(
                "Video indirme için Rüzgar virüs koruması zorunlu: "
                "/api/ruzgar/virus-guard/preflight (kind=video) sonra commit."
            ),
        )
    from urllib.parse import urlparse

    parsed = urlparse(url)
    if parsed.scheme not in ("http", "https") or not parsed.netloc:
        raise HTTPException(status_code=400, detail="Yalnızca geçerli http/https URL kabul edilir.")

    def _run():
        from ilim_assistant.motorlar.video_motoru import download_video_with_yt_dlp

        result = download_video_with_yt_dlp(url)
        payload = result.to_metadata()
        if result.error:
            payload["error"] = result.error
        return payload

    try:
        data = await run_in_threadpool(_run)
        return {"ok": bool(data.get("ok")), "result": data}
    except HTTPException:
        raise
    except Exception as exc:
        raise HTTPException(status_code=500, detail=str(exc)) from exc


@app.post("/api/video/probe")
async def api_video_probe(
    rel: Annotated[str | None, Form()] = None,
    file: UploadFile | None = File(None),
):
    """ffprobe özeti — dosya yükleme veya proje içi göreli yol (rel)."""
    if not ffprobe_available():
        raise HTTPException(
            status_code=503,
            detail="ffprobe bulunamadı: FFmpeg kurun ve PATH'e ekleyin.",
        )
    rel_s = (rel or "").strip()
    if rel_s:
        try:
            input_path = _repo_resolve_under_root(rel_s)
        except HTTPException:
            raise
        except Exception as exc:
            raise HTTPException(status_code=400, detail=str(exc)) from exc
        if not input_path.is_file():
            raise HTTPException(status_code=404, detail=f"Dosya bulunamadı: {rel_s}")

        def _run_rel():
            data = ffprobe_json(str(input_path))
            return summarize_probe(data)

        summary = await run_in_threadpool(_run_rel)
        return {"ok": True, "summary": summary, "rel": rel_s}

    if file is None:
        raise HTTPException(status_code=400, detail="Dosya yükleyin veya rel yolu gönderin.")

    raw = await file.read()
    ln = len(raw)
    if ln < 32:
        raise HTTPException(status_code=400, detail="Dosya çok kısa veya boş.")
    if ln > MAX_VIDEO_PROBE_BYTES:
        raise HTTPException(
            status_code=400,
            detail=f"Dosya çok büyük (>{MAX_VIDEO_PROBE_BYTES // (1024 * 1024)} MB).",
        )

    fn = (file.filename or "").lower()
    sfx = ".bin"
    for ext in (
        ".mkv",
        ".webm",
        ".mov",
        ".avi",
        ".mp4",
        ".m4v",
        ".wav",
        ".mp3",
        ".flac",
        ".ogg",
        ".opus",
    ):
        if fn.endswith(ext):
            sfx = ext
            break

    tmp_path = None
    try:
        fd, tmp_path = tempfile.mkstemp(suffix=sfx)
        os.write(fd, raw)
        os.close(fd)

        def _run():
            data = ffprobe_json(tmp_path)
            return summarize_probe(data)

        summary = await run_in_threadpool(_run)
        return {"ok": True, "summary": summary}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=400, detail=str(e)) from e
    finally:
        if tmp_path and os.path.isfile(tmp_path):
            try:
                os.unlink(tmp_path)
            except OSError:
                pass


@app.post("/api/video/trim")
async def api_video_trim(
    start_sec: Annotated[float, Form()] = 0.0,
    duration_sec: Annotated[float | None, Form()] = None,
    end_sec: Annotated[float | None, Form()] = None,
    copy_streams: Annotated[str, Form()] = "true",
    rel: Annotated[str | None, Form()] = None,
    file: UploadFile | None = File(None),
):
    """Zaman aralığı keser; çıktı `.ruzgar-video-export/` altına yazılır."""
    if not ffmpeg_available():
        raise HTTPException(
            status_code=503,
            detail="ffmpeg bulunamadı. https://ffmpeg.org/download.html — kurun ve PATH'e ekleyin; ardından sunucuyu yeniden başlatın.",
        )
    rel_s = (rel or "").strip()
    if rel_s and file is not None:
        raise HTTPException(status_code=400, detail="rel ve file aynı istekte kullanılamaz.")
    if not rel_s and file is None:
        raise HTTPException(status_code=400, detail="rel veya file gerekli.")

    copy_flag = _parse_bool_form(copy_streams, True)
    sp = max(0.0, float(start_sec))
    dur: float | None = None
    if duration_sec is not None and float(duration_sec) > 0:
        dur = float(duration_sec)
    elif end_sec is not None:
        dur = float(end_sec) - sp
    else:
        raise HTTPException(status_code=400, detail="duration_sec veya end_sec gerekli.")
    if dur is None or dur <= 0:
        raise HTTPException(status_code=400, detail="Geçersiz süre (duration veya end).")
    if dur > MAX_SEGMENT_SECONDS:
        raise HTTPException(
            status_code=400,
            detail=f"Kesim süresi en fazla {MAX_SEGMENT_SECONDS // 3600} saat olabilir.",
        )

    tmp_input: str | None = None
    try:
        if rel_s:
            input_path = _repo_resolve_under_root(rel_s)
        else:
            raw = await file.read()
            if len(raw) < 32:
                raise HTTPException(status_code=400, detail="Dosya çok kısa veya boş.")
            if len(raw) > MAX_VIDEO_PROBE_BYTES:
                raise HTTPException(
                    status_code=400,
                    detail=f"Dosya çok büyük (>{MAX_VIDEO_PROBE_BYTES // (1024 * 1024)} MB); proje içi rel kullanın.",
                )
            sfx = ".bin"
            low = ((file.filename if file else "") or "").lower()
            for ext in (
                ".mkv",
                ".webm",
                ".mov",
                ".avi",
                ".mp4",
                ".m4v",
                ".wav",
                ".mp3",
            ):
                if low.endswith(ext):
                    sfx = ext
                    break
            fd, tmp_input = tempfile.mkstemp(suffix=sfx)
            os.write(fd, raw)
            os.close(fd)
            input_path = Path(tmp_input)

        out_dir = export_directory(REPO_ROOT)
        stem = input_path.stem or "media"
        suf = ".mp4" if not copy_flag else (input_path.suffix or ".mp4")
        out_name = f"ruzgar_trim_{uuid.uuid4().hex[:10]}_{stem}{suf}"
        out_path = out_dir / out_name

        def _run():
            trim_media(
                input_path,
                out_path,
                start_sec=sp,
                duration_sec=dur,
                copy_streams=copy_flag,
            )

        await run_in_threadpool(_run)
        rel_out = out_path.relative_to(REPO_ROOT.resolve()).as_posix()
        return {"ok": True, "output_rel": rel_out}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=400, detail=str(e)) from e
    finally:
        if tmp_input and os.path.isfile(tmp_input):
            try:
                os.unlink(tmp_input)
            except OSError:
                pass


@app.post("/api/video/transcode")
async def api_video_transcode(
    rel: Annotated[str | None, Form()] = None,
    file: UploadFile | None = File(None),
):
    """Tam dosyayı H.264 + AAC MP4 olarak yeniden kodlar."""
    if not ffmpeg_available():
        raise HTTPException(
            status_code=503,
            detail="ffmpeg bulunamadı. https://ffmpeg.org/download.html — kurun ve PATH'e ekleyin; ardından sunucuyu yeniden başlatın.",
        )
    rel_s = (rel or "").strip()
    if rel_s and file is not None:
        raise HTTPException(status_code=400, detail="rel ve file aynı istekte kullanılamaz.")
    if not rel_s and file is None:
        raise HTTPException(status_code=400, detail="rel veya file gerekli.")

    tmp_input: str | None = None
    try:
        if rel_s:
            input_path = _repo_resolve_under_root(rel_s)
        else:
            raw = await file.read()
            if len(raw) < 64:
                raise HTTPException(status_code=400, detail="Dosya çok kısa veya boş.")
            if len(raw) > MAX_VIDEO_PROBE_BYTES:
                raise HTTPException(
                    status_code=400,
                    detail=f"Dosya çok büyük (>{MAX_VIDEO_PROBE_BYTES // (1024 * 1024)} MB); proje içi rel kullanın.",
                )
            low = ((file.filename if file else "") or "").lower()
            sfx = ".bin"
            for ext in (".mkv", ".webm", ".mov", ".avi", ".mp4", ".m4v", ".mp3", ".wav"):
                if low.endswith(ext):
                    sfx = ext
                    break
            fd, tmp_input = tempfile.mkstemp(suffix=sfx)
            os.write(fd, raw)
            os.close(fd)
            input_path = Path(tmp_input)

        out_dir = export_directory(REPO_ROOT)
        out_path = out_dir / f"ruzgar_x264_{uuid.uuid4().hex[:10]}.mp4"

        def _run():
            transcode_to_mp4(input_path, out_path)

        await run_in_threadpool(_run)
        rel_out = out_path.relative_to(REPO_ROOT.resolve()).as_posix()
        return {"ok": True, "output_rel": rel_out}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=400, detail=str(e)) from e
    finally:
        if tmp_input and os.path.isfile(tmp_input):
            try:
                os.unlink(tmp_input)
            except OSError:
                pass


@app.post("/api/video/concat")
async def api_video_concat(body: VideoConcatBody):
    """İki proje içi dosyayı birleştirir (ffmpeg concat demuxer)."""
    if not ffmpeg_available():
        raise HTTPException(
            status_code=503,
            detail="ffmpeg bulunamadı. https://ffmpeg.org/download.html — kurun ve PATH'e ekleyin; ardından sunucuyu yeniden başlatın.",
        )
    a = (body.rel_a or "").strip()
    b = (body.rel_b or "").strip()
    if not a or not b:
        raise HTTPException(status_code=400, detail="rel_a ve rel_b gerekli.")
    pa = _repo_resolve_under_root(a)
    pb = _repo_resolve_under_root(b)
    out_dir = export_directory(REPO_ROOT)
    out_path = out_dir / f"ruzgar_concat_{uuid.uuid4().hex[:12]}.mp4"

    def _run():
        concat_two_files(pa, pb, out_path, copy_streams=bool(body.copy_streams))

    try:
        await run_in_threadpool(_run)
        rel_out = out_path.relative_to(REPO_ROOT.resolve()).as_posix()
        return {"ok": True, "output_rel": rel_out}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=400, detail=str(e)) from e


@app.post("/api/video/edit/mix")
async def api_video_edit_mix(body: VideoEditMixBody):
    """Timeline parçalarını kesip birleştirir; metadata Merkezi Zihin Havuzu'na yazılır."""
    if not ffmpeg_available():
        raise HTTPException(
            status_code=503,
            detail="ffmpeg bulunamadı. https://ffmpeg.org/download.html — kurun ve PATH'e ekleyin; ardından sunucuyu yeniden başlatın.",
        )
    if not body.clips:
        raise HTTPException(status_code=400, detail="En az bir parça gerekli.")

    from ilim_assistant.motorlar.video_motoru import EDIT_CLIP_MAX_SEC, mix_timeline_clips

    payload = [
        {
            "rel": (c.rel or "").strip(),
            "start_sec": float(c.start_sec or 0),
            "end_sec": c.end_sec,
            "label": (c.label or "").strip(),
        }
        for c in body.clips
    ]
    for i, row in enumerate(payload):
        if not row["rel"]:
            raise HTTPException(status_code=400, detail=f"Parça {i + 1}: rel boş.")
        if row["end_sec"] is not None:
            span = float(row["end_sec"]) - row["start_sec"]
            if span > EDIT_CLIP_MAX_SEC + 0.05:
                raise HTTPException(
                    status_code=400,
                    detail=f"Parça {i + 1}: en fazla {EDIT_CLIP_MAX_SEC // 60} dakika.",
                )

    def _run():
        return mix_timeline_clips(
            payload,
            workspace_root=str(REPO_ROOT.resolve()),
            copy_streams=bool(body.copy_streams),
            project_name=(body.project_name or "").strip(),
        )

    try:
        result = await run_in_threadpool(_run)
        meta = result.to_metadata()
        return {
            "ok": result.ok,
            "output_rel": result.output_rel,
            "project_id": result.project_id,
            "parts": result.parts,
            "total_duration_sec": result.total_duration_sec,
            "error": result.error,
            "metadata": meta,
        }
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=400, detail=str(e)) from e


MAX_SUBTITLE_BYTES = 12 * 1024 * 1024


@app.post("/api/video/burn-subtitles")
async def api_video_burn_subtitles(
    rel_video: Annotated[str, Form()],
    rel_sub: Annotated[str, Form()],
):
    """Altyazıyı (SRT/ASS/VTT) videoya gömer; çıktı `.ruzgar-video-export/` altına yazılır."""
    if not ffmpeg_available():
        raise HTTPException(
            status_code=503,
            detail="ffmpeg bulunamadı. https://ffmpeg.org/download.html — kurun ve PATH'e ekleyin; ardından sunucuyu yeniden başlatın.",
        )
    rv = (rel_video or "").strip()
    rs = (rel_sub or "").strip()
    if not rv or not rs:
        raise HTTPException(status_code=400, detail="rel_video ve rel_sub gerekli.")
    video_path = _repo_resolve_under_root(rv)
    sub_path = _repo_resolve_under_root(rs)
    if sub_path.stat().st_size > MAX_SUBTITLE_BYTES:
        raise HTTPException(status_code=400, detail="Altyazı dosyası çok büyük.")
    low = sub_path.suffix.lower()
    if low not in (".srt", ".ass", ".ssa", ".vtt"):
        raise HTTPException(
            status_code=400,
            detail="Altyazı uzantısı desteklenmiyor (.srt, .ass, .ssa, .vtt).",
        )
    out_dir = export_directory(REPO_ROOT)
    stem = video_path.stem or "video"
    out_path = out_dir / f"ruzgar_subburn_{uuid.uuid4().hex[:10]}_{stem}.mp4"

    def _run():
        burn_subtitles_into_mp4(video_path, sub_path, out_path)

    try:
        await run_in_threadpool(_run)
        rel_out = out_path.relative_to(REPO_ROOT.resolve()).as_posix()
        return {"ok": True, "output_rel": rel_out}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=400, detail=str(e)) from e


@app.post("/api/video/mux-audio")
async def api_video_mux_audio(
    rel_video: Annotated[str, Form()],
    rel_audio: Annotated[str, Form()],
    copy_video: Annotated[str, Form()] = "true",
    shortest: Annotated[str, Form()] = "true",
):
    """Harici ses dosyasını videonun ses izine bağlar (çoğu durumda eski ses yerine geçer)."""
    if not ffmpeg_available():
        raise HTTPException(
            status_code=503,
            detail="ffmpeg bulunamadı. https://ffmpeg.org/download.html — kurun ve PATH'e ekleyin; ardından sunucuyu yeniden başlatın.",
        )
    rv = (rel_video or "").strip()
    ra = (rel_audio or "").strip()
    if not rv or not ra:
        raise HTTPException(status_code=400, detail="rel_video ve rel_audio gerekli.")
    video_path = _repo_resolve_under_root(rv)
    audio_path = _repo_resolve_under_root(ra)
    copy_v = _parse_bool_form(copy_video, True)
    short = _parse_bool_form(shortest, True)
    out_dir = export_directory(REPO_ROOT)
    stem = video_path.stem or "video"
    out_path = out_dir / f"ruzgar_muxaud_{uuid.uuid4().hex[:10]}_{stem}.mp4"

    def _run():
        mux_replace_audio(
            video_path,
            audio_path,
            out_path,
            copy_video=copy_v,
            shortest=short,
        )

    try:
        await run_in_threadpool(_run)
        rel_out = out_path.relative_to(REPO_ROOT.resolve()).as_posix()
        return {"ok": True, "output_rel": rel_out}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=400, detail=str(e)) from e


@app.get("/api/workspace/list")
def api_workspace_list(rel: str = Query("")) -> dict[str, Any]:
    """Proje köküne göre klasör içeriği (Okuma / Programlama ağaçları; Electron olmadan da çalışır)."""
    items = _repo_list_children(rel)
    return {"ok": True, "items": items}


@app.get("/api/workspace/read-text")
def api_workspace_read_text(rel: str = Query("")) -> dict[str, Any]:
    """UTF-8 metin dosyası okuma (PDF vb. ikilik içerik kullanıcıya ham görünebilir)."""
    raw = (rel or "").strip().replace("\\", "/").lstrip("/")
    if not raw:
        raise HTTPException(status_code=400, detail="rel gerekli.")
    target = _repo_resolve_under_root(raw, must_be_dir=False)
    if target.stat().st_size > 2_000_000:
        raise HTTPException(status_code=400, detail="Dosya çok büyük (2 MB).")
    text = target.read_text(encoding="utf-8", errors="replace")
    return {"ok": True, "text": text}


_WORKSPACE_MEDIA_EXT = {
    ".mp4",
    ".mkv",
    ".webm",
    ".mov",
    ".avi",
    ".m4v",
    ".mp3",
    ".wav",
    ".opus",
    ".aac",
    ".flac",
    ".ogg",
    ".m4a",
}
_WORKSPACE_MEDIA_TYPES = {
    ".mp4": "video/mp4",
    ".mkv": "video/x-matroska",
    ".webm": "video/webm",
    ".mov": "video/quicktime",
    ".avi": "video/x-msvideo",
    ".m4v": "video/x-m4v",
    ".mp3": "audio/mpeg",
    ".wav": "audio/wav",
    ".opus": "audio/opus",
    ".aac": "audio/aac",
    ".flac": "audio/flac",
    ".ogg": "audio/ogg",
    ".m4a": "audio/mp4",
}


@app.get("/api/workspace/media")
def api_workspace_media(rel: str = Query("")):
    """Video/ses önizleme akışı (sinema oynatıcısı)."""
    raw = (rel or "").strip().replace("\\", "/").lstrip("/")
    if not raw:
        raise HTTPException(status_code=400, detail="rel gerekli.")
    target = _repo_resolve_under_root(raw, must_be_dir=False)
    ext = target.suffix.lower()
    if ext not in _WORKSPACE_MEDIA_EXT:
        raise HTTPException(
            status_code=400,
            detail=f"Desteklenmeyen medya uzantısı: {ext or '(yok)'}",
        )
    size = target.stat().st_size
    if size > 900_000_000:
        raise HTTPException(status_code=400, detail="Dosya çok büyük (900 MB önizleme sınırı).")
    media = _WORKSPACE_MEDIA_TYPES.get(ext, "application/octet-stream")
    return FileResponse(str(target), media_type=media, filename=target.name)


@app.get("/api/workspace/read-pdf")
def api_workspace_read_pdf(rel: str = Query("")) -> dict[str, Any]:
    """PDF'den düz metin çıkarır (Okuma atölyesi). pypdf yoksa 503."""
    if not pdf_text_runtime_available():
        raise HTTPException(
            status_code=503,
            detail="PDF metin çıkarma için: pip install pypdf",
        )
    raw = (rel or "").strip().replace("\\", "/").lstrip("/")
    if not raw:
        raise HTTPException(status_code=400, detail="rel gerekli.")
    target = _repo_resolve_under_root(raw, must_be_dir=False)
    if target.suffix.lower() != ".pdf":
        raise HTTPException(status_code=400, detail="Dosya PDF değil.")
    file_size = target.stat().st_size
    # Dosya boyutu sınırı yok; önizleme için sayfa sayısı kademeli
    if file_size > 120_000_000:
        page_cap = 25
    elif file_size > 40_000_000:
        page_cap = 50
    else:
        page_cap = 120

    from pypdf import PdfReader

    reader = PdfReader(str(target))
    n_pages = len(reader.pages)
    max_pages = min(n_pages, page_cap)
    parts: list[str] = []
    for i in range(max_pages):
        try:
            t = reader.pages[i].extract_text() or ""
        except Exception:
            t = ""
        parts.append(t)
    full = "\n\n".join(parts).strip()
    hard_cap = 900_000
    truncated_len = False
    if len(full) > hard_cap:
        full = full[:hard_cap] + "\n\n… [metin uzun olduğu için kesildi]"
        truncated_len = True
    truncated_pages = max_pages < n_pages
    if truncated_pages and full:
        full = (
            f"[PDF: {n_pages} sayfa — önizleme: ilk {max_pages} sayfa. "
            f"Tam çeviri için «Sayfa sayfa» modunu kullanın.]\n\n{full}"
        )
    return {
        "ok": True,
        "text": full,
        "pages_total": n_pages,
        "pages_read": max_pages,
        "truncated_pages": truncated_pages,
        "file_size": file_size,
        "truncated_length": truncated_len,
    }


@app.get("/api/tercume/pdf-page-preview")
def api_tercume_pdf_page_preview(
    rel: str = Query(""),
    page: int = Query(1, ge=1),
    dpi: int = Query(110, ge=72, le=200),
) -> dict[str, Any]:
    """PDF tek sayfa PNG önizleme (base64). PyMuPDF yoksa 503."""
    from ilim_assistant.motorlar.tercume_pdf_preview import pymupdf_available, render_pdf_page_png_b64

    if not pymupdf_available():
        raise HTTPException(
            status_code=503,
            detail="PDF önizleme için: pip install pymupdf",
        )
    raw = (rel or "").strip().replace("\\", "/").lstrip("/")
    if not raw:
        raise HTTPException(status_code=400, detail="rel gerekli.")
    target = _repo_resolve_under_root(raw, must_be_dir=False)
    hit = render_pdf_page_png_b64(target, page=page, dpi=dpi)
    if not hit.get("ok"):
        raise HTTPException(status_code=400, detail=str(hit.get("error") or "Önizleme hatası"))
    hit["rel"] = raw
    return hit


@app.get("/api/tercume/pdf-page-ocr")
async def api_tercume_pdf_page_ocr(
    rel: str = Query(""),
    page: int = Query(1, ge=1),
    src_lang: str = Query("auto"),
    ocr_preset: str = Query("auto"),
) -> dict[str, Any]:
    """PDF tek sayfa OCR (PyMuPDF + Tesseract) — taranmış PDF için."""
    from ilim_assistant.motorlar.tercume_ocr_cascade import ocr_pdf_page_via_pymupdf
    from ilim_assistant.motorlar.tercume_ocr_runtime import ocr_runtime_status

    raw = (rel or "").strip().replace("\\", "/").lstrip("/")
    if not raw:
        raise HTTPException(status_code=400, detail="rel gerekli.")
    target = _repo_resolve_under_root(raw, must_be_dir=False)
    if target.suffix.lower() != ".pdf":
        raise HTTPException(status_code=400, detail="Dosya PDF değil.")

    st = ocr_runtime_status()
    if not st.get("available"):
        raise HTTPException(
            status_code=503,
            detail=st.get("hint") or "Tesseract kurulu değil — Ruzgar_OCR_Kur.bat",
        )

    def _run():
        return ocr_pdf_page_via_pymupdf(
            str(target),
            max(0, int(page) - 1),
            src_lang=(src_lang or "auto").strip(),
        )

    hit = await run_in_threadpool(_run)
    if not hit.get("ok"):
        msg = str(hit.get("error") or hit.get("reason") or "OCR başarısız")
        raise HTTPException(status_code=400, detail=msg)
    return {
        "ok": True,
        "rel": raw,
        "page": hit.get("page") or page,
        "text": str(hit.get("text") or ""),
        "quality": hit.get("quality"),
        "quality_score": hit.get("quality_score"),
        "preset": hit.get("preset"),
        "engine": hit.get("engine"),
    }


def _docx_file_to_plain(target: Path) -> str:
    from docx import Document

    doc = Document(str(target))
    chunks: list[str] = []
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            chunks.append(t)
    for tbl in doc.tables:
        for row in tbl.rows:
            cells = [(c.text or "").strip() for c in row.cells]
            line = " | ".join(x for x in cells if x)
            if line:
                chunks.append(line)
    return "\n\n".join(chunks).strip()


@app.get("/api/workspace/read-docx")
def api_workspace_read_docx(rel: str = Query("")) -> dict[str, Any]:
    """DOCX (Word) gövdesinden düz metin çıkarır. python-docx yoksa 503."""
    if not docx_text_runtime_available():
        raise HTTPException(
            status_code=503,
            detail="DOCX okuma için: pip install python-docx",
        )
    raw = (rel or "").strip().replace("\\", "/").lstrip("/")
    if not raw:
        raise HTTPException(status_code=400, detail="rel gerekli.")
    target = _repo_resolve_under_root(raw, must_be_dir=False)
    if target.suffix.lower() != ".docx":
        raise HTTPException(status_code=400, detail="Dosya DOCX değil (.doc eski format desteklenmez).")
    max_bytes = 15_000_000
    if target.stat().st_size > max_bytes:
        raise HTTPException(status_code=400, detail=f"Dosya çok büyük ({max_bytes // 1_000_000} MB sınır).")
    full = _docx_file_to_plain(target)
    hard_cap = 450_000
    truncated = False
    if len(full) > hard_cap:
        full = full[:hard_cap] + "\n\n… [metin uzun olduğu için kesildi]"
        truncated = True
    return {"ok": True, "text": full, "truncated_length": truncated}


@app.get("/api/workspace/read-ebook")
def api_workspace_read_ebook(rel: str = Query("")) -> dict[str, Any]:
    """
    E-kitap/kitap metin çıkarımı:
    epub, fb2, rtf, html/htm, md/markdown, txt.
    """
    raw = (rel or "").strip().replace("\\", "/").lstrip("/")
    if not raw:
        raise HTTPException(status_code=400, detail="rel gerekli.")
    target = _repo_resolve_under_root(raw, must_be_dir=False)
    ext = target.suffix.lower()
    max_bytes = 25_000_000
    if target.stat().st_size > max_bytes:
        raise HTTPException(status_code=400, detail=f"Dosya çok büyük ({max_bytes // 1_000_000} MB sınır).")

    meta: dict[str, Any] = {"ext": ext}
    if ext == ".epub":
        text, m = _read_epub_text(target)
        meta.update(m)
        if m.get("title"):
            meta["ebook_title"] = m["title"]
        if m.get("author"):
            meta["ebook_author"] = m["author"]
    elif ext == ".fb2":
        text, m = _read_fb2_text(target)
        meta.update(m)
        if m.get("title"):
            meta["ebook_title"] = m["title"]
        if m.get("author"):
            meta["ebook_author"] = m["author"]
    elif ext in {".rtf"}:
        text = _rtf_to_text(target.read_text(encoding="utf-8", errors="replace"))
    elif ext in {".html", ".htm"}:
        text = _simple_html_to_text(target.read_text(encoding="utf-8", errors="replace"))
    elif ext in {".md", ".markdown", ".txt"}:
        text = target.read_text(encoding="utf-8", errors="replace")
    elif ext in {".mobi", ".azw", ".azw3", ".kfx", ".djvu", ".djv"}:
        from ilim_assistant.motorlar.tercume_ebook_read import read_ebook_auto

        hit = read_ebook_auto(target)
        if not hit.get("ok"):
            raise HTTPException(
                status_code=400,
                detail=str(hit.get("error") or "E-kitap okunamadı"),
            )
        text = str(hit.get("text") or "")
        meta.update(hit.get("meta") or {})
        if hit.get("title"):
            meta["title"] = hit["title"]
        if hit.get("chapters_read") is not None:
            meta["chapters_read"] = hit["chapters_read"]
    else:
        raise HTTPException(status_code=400, detail="Bu uzantı read-ebook ile desteklenmiyor.")

    hard_cap = 900_000
    truncated = False
    if len(text) > hard_cap:
        text = text[:hard_cap] + "\n\n… [metin uzun olduğu için kesildi]"
        truncated = True
    out: dict[str, Any] = {
        "ok": True,
        "text": text,
        "truncated_length": truncated,
        "meta": meta,
    }
    if meta.get("title") or meta.get("ebook_title"):
        out["title"] = meta.get("title") or meta.get("ebook_title")
    if meta.get("author") or meta.get("ebook_author"):
        out["author"] = meta.get("author") or meta.get("ebook_author")
    if meta.get("chapters_read") is not None:
        out["chapters_read"] = meta.get("chapters_read")
    return out


@app.get("/api/workspace/read-image-ocr")
def api_workspace_read_image_ocr(
    rel: str = Query(""),
    lang: str = Query("tur+eng"),
    ocr_preset: str = Query(""),
    src_lang: str = Query("auto"),
    tercume: int = Query(0),
) -> dict[str, Any]:
    """
    Görselden metin çıkarımı (opsiyonel):
    - pytesseract + pillow kuruluysa OCR yapar.
    - yoksa net kurulum mesajı döner.
    """
    raw = (rel or "").strip().replace("\\", "/").lstrip("/")
    if not raw:
        raise HTTPException(status_code=400, detail="rel gerekli.")
    target = _repo_resolve_under_root(raw, must_be_dir=False)
    ext = target.suffix.lower()
    if ext not in {".png", ".jpg", ".jpeg", ".webp", ".bmp", ".tif", ".tiff"}:
        raise HTTPException(status_code=400, detail="Dosya görsel değil.")
    max_bytes = 25_000_000
    if target.stat().st_size > max_bytes:
        raise HTTPException(status_code=400, detail=f"Görsel çok büyük ({max_bytes // 1_000_000} MB sınır).")
    try:
        from PIL import Image
        import pytesseract
    except Exception as exc:
        raise HTTPException(
            status_code=503,
            detail="OCR için: pip install pillow pytesseract ve sistemde tesseract kurulumu gerekli.",
        ) from exc
    tess_lang = (lang or "tur+eng").strip()
    if int(tercume or 0) == 1 or (ocr_preset or "").strip():
        from ilim_assistant.motorlar.tercume_ocr_lang import resolve_ocr_lang

        tess_lang = resolve_ocr_lang((ocr_preset or lang or "auto").strip(), src_lang=(src_lang or "auto"))
    try:
        img = Image.open(target)
        txt = pytesseract.image_to_string(img, lang=tess_lang)
    except Exception as exc:
        raise HTTPException(status_code=400, detail=f"OCR başarısız: {str(exc)[:220]}") from exc
    text = (txt or "").strip()
    if int(tercume or 0) == 1:
        from ilim_assistant.motorlar.tercume_ocr_clean import clean_ocr_text

        text = clean_ocr_text(text)
    if not text:
        text = "[OCR boş sonuç verdi]"
    hard_cap = 700_000
    truncated = False
    if len(text) > hard_cap:
        text = text[:hard_cap] + "\n\n… [OCR metni uzun olduğu için kesildi]"
        truncated = True
    return {"ok": True, "text": text, "truncated_length": truncated, "lang": tess_lang}


def _tercume_download_target_dir(
    target_dir_rel: str = "",
    target_abs: str = "",
) -> Path:
    """Tercüme indirme hedefi: proje içi göreli yol veya Windows «farklı kaydet» mutlak klasör."""
    abs_raw = (target_abs or "").strip()
    if abs_raw:
        p = Path(abs_raw).expanduser().resolve()
        if not p.is_dir():
            try:
                p.mkdir(parents=True, exist_ok=True)
            except OSError as exc:
                raise HTTPException(status_code=400, detail=f"Klasör oluşturulamadı: {str(exc)[:120]}") from exc
        return p
    rel = (target_dir_rel or "").strip().replace("\\", "/").lstrip("/")
    if not rel:
        rel = "ilim-assistant/arsiv/tercume-imports"
    from ilim_assistant.motorlar.arsiv_indirme import resolve_arsiv_dir

    try:
        return resolve_arsiv_dir(rel if rel.startswith("ilim-assistant/") else f"ilim-assistant/arsiv/{rel}")
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc


@app.get("/api/ruzgar/virus-guard/capabilities")
def api_virus_guard_capabilities() -> dict[str, Any]:
    from ilim_assistant.motorlar.ruzgar_virus_guard import capabilities

    return capabilities()


@app.post("/api/ruzgar/virus-guard/preflight")
async def api_virus_guard_preflight(
    url: str = Form(""),
    filename_hint: str = Form(""),
    kind: str = Form("url"),
    scan_mode: str = Form("deep"),
) -> dict[str, Any]:
    """Karantinaya indir + tara; kullanıcı sesli onayı beklenir."""
    from ilim_assistant.motorlar.ruzgar_virus_guard import (
        guard_enabled,
        preflight_url_download,
        preflight_video_download,
    )

    if not guard_enabled():
        raise HTTPException(status_code=503, detail="Rüzgar virüs koruması kapalı.")
    u = (url or "").strip()
    if not u:
        raise HTTPException(status_code=400, detail="url gerekli.")
    k = (kind or "url").strip().lower()

    mode = (scan_mode or "deep").strip().lower()

    def _run():
        if k == "video":
            return preflight_video_download(u, scan_mode=mode)
        return preflight_url_download(u, filename_hint=filename_hint, scan_mode=mode)

    try:
        data = await run_in_threadpool(_run)
    except Exception as exc:
        raise HTTPException(
            status_code=400,
            detail=str(exc)[:2000] or "Virüs koruması: işlem başarısız",
        ) from exc
    if not data.get("ok"):
        raise HTTPException(
            status_code=400,
            detail=str(data.get("error") or "Virüs koruması: işlem reddedildi"),
        )
    return data


@app.post("/api/ruzgar/virus-guard/commit")
def api_virus_guard_commit(
    pending_id: str = Form(""),
    target_dir_rel: str = Form(""),
    target_abs: str = Form(""),
) -> dict[str, Any]:
    """Kullanıcı onayı sonrası dosyayı hedefe taşı."""
    from ilim_assistant.motorlar.ruzgar_virus_guard import commit_pending

    pid = (pending_id or "").strip()
    if not pid:
        raise HTTPException(status_code=400, detail="pending_id gerekli.")
    res = commit_pending(pid, target_dir_rel=target_dir_rel, target_abs=target_abs)
    if not res.get("ok"):
        raise HTTPException(status_code=400, detail=str(res.get("error") or "Onay tamamlanamadı"))
    return res


@app.post("/api/ruzgar/virus-guard/reject")
def api_virus_guard_reject(pending_id: str = Form("")) -> dict[str, Any]:
    from ilim_assistant.motorlar.ruzgar_virus_guard import reject_pending

    pid = (pending_id or "").strip()
    if not pid:
        raise HTTPException(status_code=400, detail="pending_id gerekli.")
    return reject_pending(pid)


@app.post("/api/ruzgar/virus-guard/scan-file")
async def api_virus_guard_scan_file(
    file: UploadFile = File(...),
    scan_mode: str = Form("deep"),
) -> dict[str, Any]:
    """Yerel dosyayı Rüzgar Virüs Kalkanı ile tara (indirme dışı)."""
    import re as _re
    import uuid as _uuid

    from ilim_assistant.motorlar.ruzgar_antivirus import ruzgar_scan_file

    staging = REPO_ROOT / "ilim-assistant" / "arsiv" / "_virus_guard_staging" / "scan_upload"
    staging.mkdir(parents=True, exist_ok=True)
    name = (file.filename or "upload.bin").replace("\\", "/").split("/")[-1]
    safe = _re.sub(r"[^a-zA-Z0-9._\-]+", "_", name).strip("._") or f"up_{_uuid.uuid4().hex[:8]}"
    target = staging / safe
    data = await file.read()
    if len(data) > 64 * 1024 * 1024:
        raise HTTPException(status_code=400, detail="Dosya çok büyük (64 MB).")
    target.write_bytes(data)
    mode = (scan_mode or "deep").strip().lower()

    def _run():
        return ruzgar_scan_file(target, mode=mode).to_dict()

    return {"ok": True, "scan": await run_in_threadpool(_run)}


@app.get("/api/ruzgar/virus-guard/threat-log")
def api_virus_guard_threat_log(limit: int = Query(40, ge=1, le=200)) -> dict[str, Any]:
    from ilim_assistant.motorlar.ruzgar_antivirus import list_threat_log

    return list_threat_log(limit=limit)


@app.post("/api/tercume/import-url")
def api_tercume_import_url(
    url: str = Form(""),
    filename_hint: str = Form(""),
    target_dir_rel: str = Form(""),
    target_abs: str = Form(""),
    guard_bypass: str = Form("0"),
) -> dict[str, Any]:
    """URL'den akışlı indirme — Rüzgar virüs koruması zorunlu (önce /virus-guard/preflight)."""
    from ilim_assistant.motorlar.arsiv_indirme import download_url_to_folder
    from ilim_assistant.motorlar.ruzgar_virus_guard import guard_enabled

    if guard_enabled() and str(guard_bypass or "").strip() not in ("1", "true", "yes"):
        raise HTTPException(
            status_code=403,
            detail=(
                "Doğrudan indirme kapalı. Rüzgar virüs koruması: önce tarama ve sesli onay "
                "(/api/ruzgar/virus-guard/preflight → commit)."
            ),
        )

    u = (url or "").strip()
    if not u:
        raise HTTPException(status_code=400, detail="url gerekli.")
    out_dir = _tercume_download_target_dir(target_dir_rel, target_abs)
    res = download_url_to_folder(u, out_dir, filename_hint=filename_hint, timeout_sec=7200.0)
    if not res.get("ok"):
        raise HTTPException(status_code=400, detail=str(res.get("error") or "İndirme başarısız"))
    rel = str(res.get("rel") or "")
    if rel and not rel.startswith("ilim-assistant/"):
        try:
            rel = Path(res["abs"]).resolve().relative_to(REPO_ROOT.resolve()).as_posix()
        except Exception:
            rel = str(res.get("abs") or rel)
    return {
        "ok": True,
        "rel": rel,
        "abs": res.get("abs"),
        "bytes": res.get("bytes"),
        "content_type": res.get("content_type"),
        "skipped": res.get("skipped"),
    }


@app.post("/api/tercume/download-url")
def api_tercume_download_url(
    url: str = Form(""),
    filename_hint: str = Form(""),
    target_dir_rel: str = Form(""),
    target_abs: str = Form(""),
) -> dict[str, Any]:
    """Tercüme: seçilen klasöre büyük dosya indir (import-url ile aynı motor)."""
    return api_tercume_import_url(
        url=url,
        filename_hint=filename_hint,
        target_dir_rel=target_dir_rel,
        target_abs=target_abs,
    )


def _tercume_save_rel_allowed(raw: str) -> Path:
    rel = (raw or "").strip().replace("\\", "/").lstrip("/")
    if not rel:
        raise HTTPException(status_code=400, detail="rel gerekli.")
    allowed_prefixes = (
        "ilim-assistant/arsiv/",
        "arsiv/",
    )
    if not any(rel.startswith(p) for p in allowed_prefixes):
        raise HTTPException(
            status_code=400,
            detail="Kayıt yalnızca ilim-assistant/arsiv/ altına (örn. tercume-output/).",
        )
    if Path(rel).suffix.lower() not in {".txt", ".md", ".html", ".htm"}:
        raise HTTPException(status_code=400, detail="Uzantı: txt, md veya html olmalı.")
    root = REPO_ROOT.resolve()
    target = (root / rel.replace("/", os.sep)).resolve()
    try:
        target.relative_to(root)
    except ValueError:
        raise HTTPException(status_code=400, detail="Geçersiz yol (proje dışı).") from None
    return target


@app.get("/api/tercume/save-prefs")
def api_tercume_save_prefs() -> dict[str, Any]:
    """Faz 14D — son kayıt klasörü ve varsayılan yol."""
    from ilim_assistant.motorlar.tercume_save_prefs import get_save_prefs

    return get_save_prefs()


@app.post("/api/tercume/save-target")
def api_tercume_save_target(
    rel: str = Form(""),
    text: str = Form(""),
    copy_rel: str = Form(""),
    avoid_collision: str = Form("1"),
    export_format: str = Form(""),
    source_file: str = Form(""),
    tgt_lang: str = Form("tr"),
) -> dict[str, Any]:
    """Blok J94 — hedef çeviri metnini arşive kaydet (isteğe bağlı ikinci kopya)."""
    from ilim_assistant.motorlar.tercume_save_prefs import (
        collision_versioning_enabled,
        prepare_save_path,
        remember_save_rel,
    )

    body = (text or "")
    if not body.strip():
        raise HTTPException(status_code=400, detail="Kaydedilecek metin boş.")
    if len(body) > 2_500_000:
        raise HTTPException(status_code=400, detail="Metin çok büyük (2.5 MB sınır).")

    fmt = (export_format or "").strip().lower()
    if not fmt and (rel or "").strip():
        ext = Path((rel or "").strip()).suffix.lower()
        if ext == ".md":
            fmt = "md"
        elif ext in (".html", ".htm"):
            fmt = "html"
    if fmt in ("md", "html"):
        from ilim_assistant.motorlar.tercume_export_format import format_export_body

        body = format_export_body(
            body,
            fmt,
            source_rel=(source_file or "").strip(),
            tgt_lang=(tgt_lang or "tr").strip(),
        )

    docx_bytes: bytes | None = None
    if fmt == "docx":
        from ilim_assistant.motorlar.tercume_export_docx import build_docx_bytes, docx_available

        if not docx_available():
            raise HTTPException(
                status_code=503,
                detail="DOCX için python-docx gerekli: pip install python-docx",
            )
        try:
            docx_bytes = build_docx_bytes(
                body,
                source_rel=(source_file or "").strip(),
                tgt_lang=(tgt_lang or "tr").strip(),
            )
        except Exception as exc:
            raise HTTPException(status_code=500, detail=str(exc)[:200]) from exc

    raw_rel = (rel or "").strip()
    _tercume_save_rel_allowed(raw_rel)
    use_version = str(avoid_collision or "1").strip().lower() not in ("0", "false", "no")
    root = REPO_ROOT.resolve()
    try:
        if use_version and collision_versioning_enabled():
            prepared = prepare_save_path(raw_rel, root=root)
            target = prepared["path"]
            final_rel = prepared["rel"]
            versioned = bool(prepared.get("versioned"))
            requested_rel = prepared.get("requested_rel") or raw_rel
        else:
            target = _tercume_save_rel_allowed(raw_rel)
            final_rel = target.relative_to(root).as_posix()
            versioned = False
            requested_rel = raw_rel
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    except OSError as exc:
        raise HTTPException(status_code=500, detail=str(exc)[:200]) from exc

    target.parent.mkdir(parents=True, exist_ok=True)
    try:
        if docx_bytes is not None:
            target.write_bytes(docx_bytes)
            nbytes = len(docx_bytes)
        else:
            target.write_text(body, encoding="utf-8")
            nbytes = len(body.encode("utf-8"))
    except OSError as exc:
        raise HTTPException(status_code=500, detail=f"Yazılamadı: {str(exc)[:200]}") from exc

    remember_save_rel(final_rel)
    out: dict[str, Any] = {
        "ok": True,
        "rel": final_rel,
        "bytes": nbytes,
        "export_format": fmt or "txt",
        "versioned": versioned,
    }
    if versioned:
        out["requested_rel"] = requested_rel
        out["hint_tr"] = f"Dosya zaten vardı — {final_rel} olarak kaydedildi."

    copy_raw = (copy_rel or "").strip()
    if copy_raw:
        try:
            _tercume_save_rel_allowed(copy_raw)
            if use_version and collision_versioning_enabled():
                cp = prepare_save_path(copy_raw, root=root)
                copy_path = cp["path"]
                copy_final = cp["rel"]
                copy_versioned = bool(cp.get("versioned"))
            else:
                copy_path = _tercume_save_rel_allowed(copy_raw)
                copy_final = copy_path.relative_to(root).as_posix()
                copy_versioned = False
            copy_path.parent.mkdir(parents=True, exist_ok=True)
            if docx_bytes is not None:
                copy_path.write_bytes(docx_bytes)
            else:
                copy_path.write_text(body, encoding="utf-8")
            out["copy_rel"] = copy_final
            if copy_versioned:
                out["copy_versioned"] = True
        except HTTPException as exc:
            out["copy_error"] = str(exc.detail)
        except OSError as exc:
            out["copy_error"] = f"İkinci kopya yazılamadı: {str(exc)[:200]}"
    return out


@app.get("/api/arsiv/download-catalog")
def api_arsiv_download_catalog() -> dict[str, Any]:
    """Manifest + indirilmiş dosya durumu (tercüme atölyesi arşiv paneli)."""
    from ilim_assistant.motorlar.arsiv_indirme import ARSIV_INDIRME_VERSION, catalog_with_status

    return {**catalog_with_status(), "module": ARSIV_INDIRME_VERSION}


@app.post("/api/arsiv/download-item")
def api_arsiv_download_item(item_id: str = Form("")) -> dict[str, Any]:
    from ilim_assistant.motorlar.arsiv_indirme import import_manifest_item

    iid = (item_id or "").strip()
    if not iid:
        raise HTTPException(status_code=400, detail="item_id gerekli.")
    return import_manifest_item(iid)


@app.post("/api/arsiv/download-next")
def api_arsiv_download_next(limit: int = Form(1)) -> dict[str, Any]:
    from ilim_assistant.motorlar.arsiv_indirme import import_next_pending

    return import_next_pending(limit=max(1, min(int(limit or 1), 5)))


@app.get("/api/tercume/local-find")
def api_tercume_local_find(
    q: str = Query(""),
    limit: int = Query(8),
) -> dict[str, Any]:
    """Yerel arşivde fuzzy dosya eşleşmesi (sohbet «kitap aç» için)."""
    raw = (q or "").strip()
    if not raw:
        raise HTTPException(status_code=400, detail="q gerekli.")
    if len(raw) > 200:
        raise HTTPException(status_code=400, detail="Sorgu çok uzun (200 karakter).")
    from ilim_assistant.motorlar.tercume_analyst import find_local_archive_matches

    items = find_local_archive_matches(raw, limit=max(1, min(int(limit or 8), 20)))
    best = items[0] if items else None
    return {"ok": True, "query": raw, "items": items, "best": best}


@app.get("/api/tercume/eser-search")
def api_tercume_eser_search(
    q: str = Query(""),
    scored: int = Query(1),
    web: int = Query(1),
) -> dict[str, Any]:
    """B planı arama; scored=1 (varsayılan) analist skorları + önerilen indirme."""
    raw = (q or "").strip()
    if not raw:
        raise HTTPException(status_code=400, detail="q gerekli.")
    if len(raw) > 300:
        raise HTTPException(status_code=400, detail="Arama metni çok uzun (300 karakter).")
    if int(scored or 0) == 0:
        from ilim_assistant.motorlar.tercume_eser_arama import search_eser_merged

        return search_eser_merged(raw)
    from ilim_assistant.motorlar.tercume_analyst import analyze_tercume_query

    include_web = str(web or "1").strip().lower() not in ("0", "false", "no")
    return analyze_tercume_query(raw, include_web=include_web)


@app.get("/api/tercume/analyze")
def api_tercume_analyze_get(q: str = Query("")) -> dict[str, Any]:
    """Faz 1 — tercüme analist raporu (skorlu kaynaklar, Scholar, önerilen URL)."""
    raw = (q or "").strip()
    if not raw:
        raise HTTPException(status_code=400, detail="q gerekli.")
    if len(raw) > 300:
        raise HTTPException(status_code=400, detail="Sorgu çok uzun (300 karakter).")
    from ilim_assistant.motorlar.tercume_analyst import analyze_tercume_query

    return analyze_tercume_query(raw)


@app.post("/api/tercume/analyze")
def api_tercume_analyze_post(q: str = Form("")) -> dict[str, Any]:
    return api_tercume_analyze_get(q=q)


@app.post("/api/tercume/pipeline")
def api_tercume_pipeline(
    q: str = Form(""),
    download: str = Form("0"),
    download_url: str = Form(""),
    target_dir_rel: str = Form("ilim-assistant/arsiv/tercume-imports"),
    read_preview_pages: str = Form("0"),
    translate: str = Form("0"),
    src_lang: str = Form("auto"),
    tgt_lang: str = Form("tr"),
    workspace_root: str = Form(""),
) -> dict[str, Any]:
    """Faz 1 — analyze → isteğe indir → isteğe oku → isteğe çevir (UI değişmeden API)."""
    raw = (q or "").strip()
    if not raw:
        raise HTTPException(status_code=400, detail="q gerekli.")
    from ilim_assistant.motorlar.tercume_analyst import run_tercume_pipeline

    def _flag(v: str) -> bool:
        return str(v or "").strip().lower() in ("1", "true", "yes", "on")

    try:
        rpp = max(0, min(25, int((read_preview_pages or "0").strip() or 0)))
    except ValueError:
        rpp = 0

    return run_tercume_pipeline(
        raw,
        download=_flag(download),
        download_url=(download_url or "").strip(),
        target_dir_rel=(target_dir_rel or "ilim-assistant/arsiv/tercume-imports").strip(),
        read_preview_pages=rpp,
        translate=_flag(translate),
        src_lang=(src_lang or "auto").strip(),
        tgt_lang=(tgt_lang or "tr").strip(),
        workspace_root=(workspace_root or "").strip() or None,
    )


@app.post("/api/tercume/import-from-search")
def api_tercume_import_from_search(
    q: str = Form(""),
    download_url: str = Form(""),
    title: str = Form(""),
    read_preview_pages: str = Form("0"),
    translate: str = Form("0"),
    src_lang: str = Form("auto"),
    tgt_lang: str = Form("tr"),
    workspace_root: str = Form(""),
) -> dict[str, Any]:
    """Faz 5 — eser aramasından arşive al (yerel veya arka plan indirme)."""
    from ilim_assistant.motorlar.tercume_analyst import prepare_import_from_search
    from ilim_assistant.motorlar.tercume_analyst_jobs import start_analyst_job

    plan = prepare_import_from_search(
        query=(q or "").strip(),
        download_url=(download_url or "").strip(),
        title=(title or "").strip(),
    )
    if not plan.get("ok"):
        raise HTTPException(status_code=400, detail=str(plan.get("error") or "import planı başarısız"))

    if plan.get("mode") == "local":
        return {
            "ok": True,
            "mode": "local",
            "rel": plan.get("rel"),
            "title": plan.get("title"),
            "message": plan.get("message"),
        }

    def _flag(v: str) -> bool:
        return str(v or "").strip().lower() in ("1", "true", "yes", "on")

    try:
        rpp = max(0, min(25, int((read_preview_pages or "0").strip() or 0)))
    except ValueError:
        rpp = 0

    hit = start_analyst_job(
        query=str(plan.get("query") or q or title or "import"),
        download=True,
        download_url=str(plan.get("download_url") or ""),
        read_preview_pages=rpp,
        translate=_flag(translate),
        src_lang=(src_lang or "auto").strip(),
        tgt_lang=(tgt_lang or "tr").strip(),
        workspace_root=(workspace_root or "").strip() or None,
        title=(title or "").strip(),
    )
    if not hit.get("ok"):
        raise HTTPException(status_code=400, detail=str(hit.get("error") or "iş başlatılamadı"))
    return {
        "ok": True,
        "mode": "download",
        "job_id": hit.get("job_id"),
        "job_type": hit.get("job_type"),
        "status": hit.get("status"),
        "download_url": plan.get("download_url"),
    }


@app.post("/api/tercume/pipeline-start")
def api_tercume_pipeline_start(
    q: str = Form(""),
    download: str = Form("1"),
    download_url: str = Form(""),
    target_dir_rel: str = Form("ilim-assistant/arsiv/tercume-imports"),
    read_preview_pages: str = Form("0"),
    translate: str = Form("0"),
    src_lang: str = Form("auto"),
    tgt_lang: str = Form("tr"),
    workspace_root: str = Form(""),
) -> dict[str, Any]:
    """Faz 5 — analyze→indir pipeline arka planda."""
    raw = (q or "").strip()
    if not raw and not (download_url or "").strip():
        raise HTTPException(status_code=400, detail="q veya download_url gerekli.")
    from ilim_assistant.motorlar.tercume_analyst_jobs import start_analyst_job

    def _flag(v: str) -> bool:
        return str(v or "").strip().lower() in ("1", "true", "yes", "on")

    try:
        rpp = max(0, min(25, int((read_preview_pages or "0").strip() or 0)))
    except ValueError:
        rpp = 0

    hit = start_analyst_job(
        query=raw or "pipeline",
        download=_flag(download),
        download_url=(download_url or "").strip(),
        target_dir_rel=(target_dir_rel or "ilim-assistant/arsiv/tercume-imports").strip(),
        read_preview_pages=rpp,
        translate=_flag(translate),
        src_lang=(src_lang or "auto").strip(),
        tgt_lang=(tgt_lang or "tr").strip(),
        workspace_root=(workspace_root or "").strip() or None,
    )
    if not hit.get("ok"):
        raise HTTPException(status_code=400, detail=str(hit.get("error") or "iş başlatılamadı"))
    return hit


@app.get("/api/tercume/jobs/{job_id}")
def api_tercume_job_status(job_id: str) -> dict[str, Any]:
    """Faz 5 — analyst / read / batch iş durumu."""
    from ilim_assistant.motorlar.tercume_analyst_jobs import resolve_tercume_job

    hit = resolve_tercume_job(job_id)
    if not hit.get("ok"):
        raise HTTPException(status_code=404, detail=str(hit.get("error") or "İş yok"))
    return hit


@app.post("/api/tercume/pipeline-cancel")
def api_tercume_pipeline_cancel(job_id: str = Form("")) -> dict[str, Any]:
    from ilim_assistant.motorlar.tercume_analyst_jobs import cancel_analyst_job

    hit = cancel_analyst_job(job_id)
    if not hit.get("ok"):
        raise HTTPException(status_code=404, detail=str(hit.get("error") or "İş yok"))
    return hit


@app.post("/api/tercume/bridge-preview")
def api_tercume_bridge_preview(
    source_text: str = Form(""),
    translated_text: str = Form(""),
    source_file: str = Form(""),
    page_index: str = Form(""),
    tgt_lang: str = Form("tr"),
    src_lang: str = Form("auto"),
) -> dict[str, Any]:
    """Faz 6 — onay öncesi ana hafıza önizlemesi."""
    from ilim_assistant.motorlar.tercume_hafiza_bridge import build_bridge_preview

    pi: int | None = None
    raw_pi = (page_index or "").strip()
    if raw_pi:
        try:
            pi = int(raw_pi)
        except ValueError:
            pi = None

    hit = build_bridge_preview(
        source_text,
        translated_text,
        source_file=(source_file or "").strip(),
        page_index=pi,
        tgt_lang=(tgt_lang or "tr").strip(),
        src_lang=(src_lang or "auto").strip(),
    )
    if not hit.get("ok"):
        raise HTTPException(status_code=400, detail=str(hit.get("error") or "önizleme başarısız"))
    return hit


@app.post("/api/tercume/bridge-save")
def api_tercume_bridge_save(
    source_text: str = Form(""),
    translated_text: str = Form(""),
    source_file: str = Form(""),
    page_index: str = Form(""),
    tgt_lang: str = Form("tr"),
    src_lang: str = Form("auto"),
    approved: str = Form("0"),
    save_knowledge: str = Form("1"),
) -> dict[str, Any]:
    """Faz 6 — onaylı çeviri → genel hafıza + knowledge RAG."""
    from ilim_assistant.motorlar.tercume_hafiza_bridge import save_bridge_entry

    def _flag(v: str) -> bool:
        return str(v or "").strip().lower() in ("1", "true", "yes", "on")

    pi: int | None = None
    raw_pi = (page_index or "").strip()
    if raw_pi:
        try:
            pi = int(raw_pi)
        except ValueError:
            pi = None

    hit = save_bridge_entry(
        source_text,
        translated_text,
        source_file=(source_file or "").strip(),
        page_index=pi,
        tgt_lang=(tgt_lang or "tr").strip(),
        src_lang=(src_lang or "auto").strip(),
        approved=_flag(approved),
        save_knowledge=_flag(save_knowledge),
    )
    if not hit.get("ok"):
        raise HTTPException(status_code=400, detail=str(hit.get("error") or "kayıt başarısız"))
    return hit


@app.get("/api/tercume/bridge-log")
def api_tercume_bridge_log(limit: int = Query(12, ge=1, le=40)) -> dict[str, Any]:
    from ilim_assistant.motorlar.tercume_hafiza_bridge import read_bridge_log

    return {"ok": True, "items": read_bridge_log(limit=limit)}


@app.post("/api/tercume/report")
def api_tercume_report(
    q: str = Form(""),
    rel: str = Form(""),
    read_pages: str = Form("5"),
    save: str = Form("1"),
) -> dict[str, Any]:
    """Faz 7 — analist araştırma raporu (senkron)."""
    from ilim_assistant.motorlar.tercume_analyst_report import generate_analyst_report, save_report_file

    def _flag(v: str) -> bool:
        return str(v or "").strip().lower() in ("1", "true", "yes", "on")

    try:
        rp = max(1, min(25, int((read_pages or "5").strip() or 5)))
    except ValueError:
        rp = 5

    hit = generate_analyst_report((q or "").strip(), rel=(rel or "").strip(), read_pages=rp)
    if not hit.get("ok"):
        raise HTTPException(status_code=400, detail=str(hit.get("error") or "rapor başarısız"))
    if _flag(save):
        saved = save_report_file(hit)
        hit["saved"] = saved
    return hit


@app.post("/api/tercume/report-start")
def api_tercume_report_start(
    q: str = Form(""),
    rel: str = Form(""),
    read_pages: str = Form("5"),
    auto_import: str = Form("0"),
    download_url: str = Form(""),
) -> dict[str, Any]:
    """Faz 7 — analist raporu arka planda (isteğe indirme)."""
    from ilim_assistant.motorlar.tercume_analyst_jobs import start_report_job

    def _flag(v: str) -> bool:
        return str(v or "").strip().lower() in ("1", "true", "yes", "on")

    try:
        rp = max(1, min(25, int((read_pages or "5").strip() or 5)))
    except ValueError:
        rp = 5

    hit = start_report_job(
        query=(q or "").strip(),
        rel=(rel or "").strip(),
        read_pages=rp,
        auto_import=_flag(auto_import),
        download_url=(download_url or "").strip(),
    )
    if not hit.get("ok"):
        raise HTTPException(status_code=400, detail=str(hit.get("error") or "rapor işi başlatılamadı"))
    return hit


@app.post("/api/tercume/super-start")
def api_tercume_super_start(
    q: str = Form(""),
    rel: str = Form(""),
    read_pages: str = Form("5"),
    translate: str = Form("1"),
    tgt_lang: str = Form("tr"),
    src_lang: str = Form("auto"),
    checkpoint: str = Form(""),
) -> dict[str, Any]:
    """Faz 8/12 — tam analist zinciri arka planda."""
    from ilim_assistant.motorlar.tercume_analyst_jobs import start_super_job

    def _flag(v: str) -> bool:
        return str(v or "").strip().lower() in ("1", "true", "yes", "on")

    try:
        rp = max(1, min(25, int((read_pages or "5").strip() or 5)))
    except ValueError:
        rp = 5

    ck: bool | None = None
    if str(checkpoint or "").strip():
        ck = _flag(checkpoint)

    hit = start_super_job(
        query=(q or "").strip(),
        rel=(rel or "").strip(),
        read_pages=rp,
        translate=_flag(translate),
        tgt_lang=(tgt_lang or "tr").strip(),
        src_lang=(src_lang or "auto").strip(),
        checkpoint=ck,
    )
    if not hit.get("ok"):
        raise HTTPException(status_code=400, detail=str(hit.get("error") or "süper analist başlatılamadı"))
    return hit


@app.post("/api/tercume/super-resume")
def api_tercume_super_resume(
    job_id: str = Form(""),
    action: str = Form("continue"),
) -> dict[str, Any]:
    """Faz 12 — checkpoint: devam / atla / iptal."""
    from ilim_assistant.motorlar.tercume_analyst_jobs import resume_super_job

    hit = resume_super_job((job_id or "").strip(), action=(action or "continue").strip())
    if not hit.get("ok"):
        raise HTTPException(status_code=400, detail=str(hit.get("error") or "devam edilemedi"))
    return hit


@app.get("/api/tercume/config")
def api_tercume_config():
    from ilim_assistant.motorlar.tercume_atolye import workbench_config

    return {"ok": True, **workbench_config()}


@app.get("/api/tercume/readiness")
def api_tercume_readiness(
    need_internet: int = Query(0),
) -> dict[str, Any]:
    """Faz 13 — çeviri + OCR + arama hazırlık özeti."""
    from ilim_assistant.motorlar.tercume_readiness import collect_tercume_readiness

    net = str(need_internet or "0").strip().lower() in ("1", "true", "yes")
    return collect_tercume_readiness(need_internet=net)


@app.get("/api/tercume/capabilities")
def api_tercume_capabilities() -> dict[str, Any]:
    """Tercüme yetenek özeti: opsiyonel yerel araçlar + format desteği."""
    from ilim_assistant.motorlar.tercume_ebook_read import calibre_available, djvu_available
    from ilim_assistant.motorlar.tercume_export_docx import docx_available
    from ilim_assistant.motorlar.tercume_pdf_preview import pymupdf_available
    from ilim_assistant.motorlar.tercume_preflight import run_tercume_preflight

    pre = run_tercume_preflight()
    ocr_ok = False
    for c in pre.get("checks") or []:
        if isinstance(c, dict) and str(c.get("id") or "") == "ocr_(tesseract)":
            ocr_ok = bool(c.get("ok"))
            break
    from ilim_assistant.motorlar.ruzgar_virus_guard import capabilities as virus_guard_caps

    vg = virus_guard_caps()
    return {
        "ok": True,
        "capabilities": {
            "calibre": calibre_available(),
            "djvu": djvu_available(),
            "docx_export": docx_available(),
            "ocr": ocr_ok,
            "pdf_preview": pymupdf_available(),
            "virus_guard": vg.get("enabled"),
        },
        "virus_guard": vg,
        "notes": {
            "calibre": "MOBI/AZW okumak için Calibre (ebook-convert)",
            "djvu": "DjVu okumak için DjVuLibre (djvutxt)",
            "docx_export": "DOCX çıktı için python-docx",
            "ocr": "Taranmış görsel/PDF için Tesseract",
            "pdf_preview": "PDF sayfa görüntüsü için PyMuPDF (pip install pymupdf)",
            "virus_guard": "Rüzgar Virüs Kalkanı: çok katmanlı tarama + nötralize + sesli onay",
        },
    }


@app.post("/api/tercume/batch-start")
async def api_tercume_batch_start(request: Request) -> dict[str, Any]:
    from ilim_assistant.motorlar.tercume_batch_jobs import start_batch_job, start_page_range_job

    try:
        body = await request.json()
    except Exception:
        body = {}
    if not isinstance(body, dict):
        body = {}

    rel = str(body.get("rel") or "").strip()
    if rel:
        pf = body.get("page_from")
        pt = body.get("page_to")
        try:
            page_from = int(pf) if pf is not None and str(pf).strip() != "" else None
        except (TypeError, ValueError):
            page_from = None
        try:
            page_to = int(pt) if pt is not None and str(pt).strip() != "" else None
        except (TypeError, ValueError):
            page_to = None
        skip_empty = str(body.get("skip_empty", "1")).strip().lower() not in ("0", "false", "no")
        hit = start_page_range_job(
            rel,
            page_from=page_from,
            page_to=page_to,
            skip_empty=skip_empty,
            tgt_lang=str(body.get("tgt_lang") or "tr"),
            src_lang=str(body.get("src_lang") or "auto"),
            read_level=str(body.get("system_level") or body.get("read_level") or "akademik"),
            output_dir_rel=str(body.get("output_dir_rel") or "ilim-assistant/arsiv/tercume-output/page-range"),
        )
    else:
        hit = start_batch_job(
            str(body.get("folder_rel") or ""),
            tgt_lang=str(body.get("tgt_lang") or "tr"),
            src_lang=str(body.get("src_lang") or "auto"),
            read_level=str(body.get("system_level") or body.get("read_level") or "akademik"),
            output_dir_rel=str(body.get("output_dir_rel") or "ilim-assistant/arsiv/tercume-output/batch"),
            file_filter=str(body.get("file_filter") or ""),
        )
    if not hit.get("ok"):
        raise HTTPException(status_code=400, detail=str(hit.get("error") or "batch başlatılamadı"))
    return hit


@app.get("/api/tercume/batch-status")
def api_tercume_batch_status(job_id: str = Query("")) -> dict[str, Any]:
    from ilim_assistant.motorlar.tercume_batch_jobs import get_batch_job

    hit = get_batch_job(job_id)
    if not hit.get("ok"):
        raise HTTPException(status_code=404, detail=str(hit.get("error") or "İş yok"))
    return hit


@app.get("/api/tercume/batch-jobs")
def api_tercume_batch_jobs(limit: int = Query(15, ge=1, le=50)) -> dict[str, Any]:
    """Faz 14A — son arka plan çeviri işleri (sayfa aralığı / cilt sırası)."""
    from ilim_assistant.motorlar.tercume_batch_jobs import list_batch_jobs

    return list_batch_jobs(limit=limit)


@app.post("/api/tercume/batch-cancel")
def api_tercume_batch_cancel(job_id: str = Form("")) -> dict[str, Any]:
    from ilim_assistant.motorlar.tercume_batch_jobs import cancel_batch_job

    hit = cancel_batch_job(job_id)
    if not hit.get("ok"):
        raise HTTPException(status_code=404, detail=str(hit.get("error") or "İş yok"))
    return hit


@app.get("/api/tercume/apprentice-log")
def api_tercume_apprentice_log(
    workspace_root: str | None = None,
    limit: int = Query(12, ge=1, le=40),
):
    from ilim_assistant.motorlar.tercume_atolye import read_apprentice_log

    root = (workspace_root or "").strip() or None
    return {"ok": True, "items": read_apprentice_log(root, limit=limit)}


@app.get("/api/tercume/preflight")
def api_tercume_preflight(
    rel: str = Query(""),
    need_internet: int = Query(0),
) -> dict[str, Any]:
    """Faz 9 — OCR, beyin, arşiv kapı kontrolü."""
    from ilim_assistant.motorlar.tercume_preflight import run_tercume_preflight

    return run_tercume_preflight(
        rel=(rel or "").strip(),
        need_internet=bool(int(need_internet or 0)),
    )


@app.get("/api/tercume/source-pages")
def api_tercume_source_pages(
    rel: str = Query(""),
    page_from: int | None = Query(None, ge=0),
    page_to: int | None = Query(None, ge=0),
) -> dict[str, Any]:
    """Kaynak metni sayfa/parça listesine böler + Faz 3 kalite skoru."""
    raw = (rel or "").strip().replace("\\", "/").lstrip("/")
    if not raw:
        raise HTTPException(status_code=400, detail="rel gerekli.")
    target = _repo_resolve_under_root(raw, must_be_dir=False)
    ext = target.suffix.lower()

    _pipeline_ext = {
        ".pdf",
        ".txt",
        ".md",
        ".markdown",
        ".html",
        ".htm",
        ".docx",
        ".epub",
        ".fb2",
        ".mobi",
        ".azw",
        ".azw3",
        ".kfx",
        ".djvu",
        ".djv",
        ".png",
        ".jpg",
        ".jpeg",
        ".webp",
        ".bmp",
        ".tif",
        ".tiff",
    }
    if ext in _pipeline_ext:
        from ilim_assistant.motorlar.tercume_read_pipeline import extract_source_pages

        hit = extract_source_pages(raw, page_from=page_from, page_to=page_to)
        if not hit.get("ok"):
            raise HTTPException(status_code=400, detail=str(hit.get("error") or "Okuma hatası"))
        return hit

    from ilim_assistant.motorlar.tercume_atolye import split_text_into_pages
    from ilim_assistant.motorlar.tercume_read_pipeline import (
        enrich_pages,
        summarize_page_quality_meta,
        READ_PIPELINE_VERSION,
    )

    pages: list[dict[str, Any]] = []
    meta: dict[str, Any] = {"ext": ext, "read_pipeline": READ_PIPELINE_VERSION}

    if ext == ".docx":
        if not docx_text_runtime_available():
            raise HTTPException(status_code=503, detail="pip install python-docx")
        full = _docx_file_to_plain(target)
    elif ext in {".epub", ".fb2", ".mobi", ".azw", ".azw3", ".kfx", ".djvu", ".djv", ".rtf"}:
        hit = api_workspace_read_ebook(rel=raw)
        full = str(hit.get("text") or "")
    else:
        hit = api_workspace_read_text(rel=raw)
        full = str(hit.get("text") or "")
    for p in split_text_into_pages(full):
        pages.append(
            {
                "index": p["index"],
                "text": p["text"],
                "label": f"Bölüm {int(p['index']) + 1}",
            }
        )
    meta["pages_total"] = len(pages)
    pages = enrich_pages(pages, source_kind="text")
    meta.update(summarize_page_quality_meta(pages))
    return {"ok": True, "rel": raw, "pages": pages, "meta": meta}


@app.post("/api/tercume/read-start")
async def api_tercume_read_start(request: Request) -> dict[str, Any]:
    from ilim_assistant.motorlar.tercume_read_jobs import start_read_job

    try:
        body = await request.json()
    except Exception:
        body = {}
    if not isinstance(body, dict):
        body = {}
    hit = start_read_job(str(body.get("rel") or ""))
    if not hit.get("ok"):
        raise HTTPException(status_code=400, detail=str(hit.get("error") or "read job başlatılamadı"))
    return hit


@app.get("/api/tercume/read-status")
def api_tercume_read_status(job_id: str = Query("")) -> dict[str, Any]:
    from ilim_assistant.motorlar.tercume_read_jobs import get_read_job

    hit = get_read_job(job_id)
    if not hit.get("ok"):
        raise HTTPException(status_code=404, detail=str(hit.get("error") or "İş yok"))
    return hit


@app.post("/api/tercume/read-cancel")
def api_tercume_read_cancel(job_id: str = Form("")) -> dict[str, Any]:
    from ilim_assistant.motorlar.tercume_read_jobs import cancel_read_job

    hit = cancel_read_job(job_id)
    if not hit.get("ok"):
        raise HTTPException(status_code=404, detail=str(hit.get("error") or "İş yok"))
    return hit


@app.get("/api/tercume/review-queue")
def api_tercume_review_queue(job_id: str = Query("")) -> dict[str, Any]:
    """Faz 16B — iş çıktısından inceleme kuyruğu (düşük skor / hata)."""
    from ilim_assistant.motorlar.tercume_batch_jobs import get_batch_job
    from ilim_assistant.motorlar.tercume_review_queue import build_review_queue

    jid = (job_id or "").strip()
    if not jid:
        raise HTTPException(status_code=400, detail="job_id gerekli")
    job = get_batch_job(jid)
    if not job.get("ok"):
        raise HTTPException(status_code=404, detail=str(job.get("error") or "İş yok"))
    return build_review_queue(job)


@app.get("/api/tercume/user-glossary")
def api_tercume_user_glossary_list(limit: int = Query(60, ge=1, le=200)) -> dict[str, Any]:
    from ilim_assistant.motorlar.tercume_user_glossary import list_entries

    return list_entries(limit=limit)


@app.post("/api/tercume/user-glossary/add")
def api_tercume_user_glossary_add(
    src: str = Form(""),
    tr: str = Form(""),
    en: str = Form(""),
    ar: str = Form(""),
    scope: str = Form(""),
    note: str = Form(""),
) -> dict[str, Any]:
    from ilim_assistant.motorlar.tercume_user_glossary import add_entry

    hit = add_entry(src=src, tr=tr, en=en, ar=ar, scope=scope, note=note)
    if not hit.get("ok"):
        raise HTTPException(status_code=400, detail=str(hit.get("error") or "Eklenemedi"))
    return hit


@app.post("/api/tercume/user-glossary/delete")
def api_tercume_user_glossary_delete(entry_id: str = Form("")) -> dict[str, Any]:
    from ilim_assistant.motorlar.tercume_user_glossary import delete_entry

    hit = delete_entry(entry_id)
    if not hit.get("ok"):
        raise HTTPException(status_code=404, detail=str(hit.get("error") or "Silinemedi"))
    return hit


@app.post("/api/tercume/user-glossary/import")
def api_tercume_user_glossary_import(
    text: str = Form(""),
    fmt: str = Form(""),
    merge: str = Form("1"),
) -> dict[str, Any]:
    """Faz 14F — CSV/JSON terim içe aktarma."""
    from ilim_assistant.motorlar.tercume_user_glossary import import_from_text

    do_merge = str(merge or "1").strip().lower() not in ("0", "false", "no")
    hit = import_from_text(text or "", fmt or "", merge=do_merge)
    if not hit.get("ok"):
        raise HTTPException(status_code=400, detail=str(hit.get("error") or "İçe aktarılamadı"))
    return hit


@app.post("/api/tercume/tmx/export")
def api_tercume_tmx_export(
    rel: str = Form(""),
    src_lang: str = Form("auto"),
    tgt_lang: str = Form("tr"),
    include_glossary: str = Form("1"),
    include_tm: str = Form("1"),
) -> dict[str, Any]:
    from ilim_assistant.motorlar.tercume_tmx import collect_tmx_units, build_tmx

    raw = (rel or "").strip().replace("\\", "/").lstrip("/")
    tl = (tgt_lang or "tr").strip().lower()[:8] or "tr"
    sl = (src_lang or "auto").strip().lower()
    if sl in ("", "auto"):
        sl = "tr"
    ig = str(include_glossary or "1").strip().lower() not in ("0", "false", "no")
    itm = str(include_tm or "1").strip().lower() not in ("0", "false", "no")
    units = collect_tmx_units(
        source_file=raw,
        tgt_lang=tl,
        include_glossary=ig,
        include_tm=itm and bool(raw),
    )
    if not units:
        raise HTTPException(status_code=400, detail="Dışa aktarılacak terim çifti yok.")
    tmx = build_tmx(units, src_lang=sl, tgt_lang=tl)
    return {
        "ok": True,
        "tmx": tmx,
        "units": len(units),
        "rel": raw,
        "src_lang": sl,
        "tgt_lang": tl,
    }


@app.post("/api/tercume/tmx/export-segments")
def api_tercume_tmx_export_segments(
    segments_json: str = Form(""),
    src_lang: str = Form("auto"),
    tgt_lang: str = Form("tr"),
) -> dict[str, Any]:
    """CAT-lite: seçili/onaylı segmentlerden TMX üret."""
    import json

    from ilim_assistant.motorlar.tercume_tmx import build_tmx

    blob = (segments_json or "").strip()
    if not blob:
        raise HTTPException(status_code=400, detail="segments_json gerekli")
    try:
        data = json.loads(blob)
    except json.JSONDecodeError as exc:
        raise HTTPException(status_code=400, detail=f"Geçersiz JSON: {exc}") from exc
    if not isinstance(data, list):
        raise HTTPException(status_code=400, detail="segments_json liste olmalı")
    units: list[tuple[str, str]] = []
    for it in data[:500]:
        if not isinstance(it, dict):
            continue
        src = str(it.get("source") or "").strip()
        tgt = str(it.get("target") or "").strip()
        if src and tgt:
            units.append((src, tgt))
    if not units:
        raise HTTPException(status_code=400, detail="TMX için geçerli source/target segment yok")
    sl = (src_lang or "auto").strip().lower()
    if sl in ("", "auto"):
        sl = "tr"
    tl = (tgt_lang or "tr").strip().lower()[:8] or "tr"
    tmx = build_tmx(units, src_lang=sl, tgt_lang=tl, creation_tool="ruzgar-tercume-cat-lite")
    return {"ok": True, "tmx": tmx, "units": len(units), "src_lang": sl, "tgt_lang": tl}


@app.post("/api/tercume/tmx/import")
def api_tercume_tmx_import(
    text: str = Form(""),
    tgt_lang: str = Form("tr"),
    merge: str = Form("1"),
) -> dict[str, Any]:
    from ilim_assistant.motorlar.tercume_user_glossary import import_from_tmx

    do_merge = str(merge or "1").strip().lower() not in ("0", "false", "no")
    hit = import_from_tmx(text or "", tgt_lang=tgt_lang, merge=do_merge)
    if not hit.get("ok"):
        raise HTTPException(status_code=400, detail=str(hit.get("error") or "TMX içe aktarılamadı"))
    return hit


@app.post("/api/tercume/aligned-diff")
def api_tercume_aligned_diff(
    source_text: str = Form(""),
    target_text: str = Form(""),
    rel: str = Form(""),
) -> dict[str, Any]:
    """Faz 17 — kaynak/hedef hizalı segment tablosu."""
    from ilim_assistant.motorlar.tercume_aligned_diff import (
        build_aligned_diff,
        build_aligned_from_pages,
    )

    tgt = (target_text or "").strip()
    raw = (rel or "").strip().replace("\\", "/").lstrip("/")
    if raw:
        from ilim_assistant.motorlar.tercume_read_pipeline import extract_source_pages

        hit = extract_source_pages(raw)
        if hit.get("ok") and isinstance(hit.get("pages"), list):
            return build_aligned_from_pages(hit["pages"], tgt)
    src = (source_text or "").strip()
    if not src and not tgt:
        raise HTTPException(status_code=400, detail="Kaynak veya hedef metin gerekli.")
    return build_aligned_diff(src, tgt)


@app.get("/api/tercume/aligned-notes")
def api_tercume_aligned_notes(rel: str = Query("")) -> dict[str, Any]:
    from ilim_assistant.motorlar.tercume_aligned_notes import load_notes

    raw = (rel or "").strip().replace("\\", "/").lstrip("/")
    if not raw:
        raise HTTPException(status_code=400, detail="rel gerekli")
    hit = load_notes(raw)
    if not hit.get("ok"):
        raise HTTPException(status_code=400, detail=str(hit.get("error") or "Notlar okunamadı"))
    return hit


@app.post("/api/tercume/aligned-notes/save")
def api_tercume_aligned_notes_save(
    rel: str = Form(""),
    notes_json: str = Form(""),
) -> dict[str, Any]:
    import json

    from ilim_assistant.motorlar.tercume_aligned_notes import save_notes

    raw = (rel or "").strip().replace("\\", "/").lstrip("/")
    if not raw:
        raise HTTPException(status_code=400, detail="rel gerekli")
    try:
        notes = json.loads(notes_json or "{}")
    except json.JSONDecodeError as exc:
        raise HTTPException(status_code=400, detail=f"Geçersiz notes_json: {exc}") from exc
    hit = save_notes(raw, notes if isinstance(notes, dict) else {})
    if not hit.get("ok"):
        raise HTTPException(status_code=400, detail=str(hit.get("error") or "Notlar kaydedilemedi"))
    return hit


@app.get("/api/tercume/memory-status")
def api_tercume_memory_status(
    rel: str = Query(""),
    tgt_lang: str = Query("tr"),
) -> dict[str, Any]:
    """Faz 14C — kalıcı terim belleği (TM) özeti."""
    from ilim_assistant.motorlar.tercume_translate_memory import memory_status

    raw = (rel or "").strip().replace("\\", "/").lstrip("/")
    if not raw:
        raise HTTPException(status_code=400, detail="rel gerekli.")
    return memory_status(raw, tgt_lang=tgt_lang)


@app.post("/api/tercume/memory-clear")
def api_tercume_memory_clear(
    source_file: str = Form(""),
    tgt_lang: str = Form(""),
) -> dict[str, Any]:
    """Faz 4/14C — TM temizle (tek dosya+tüm diller veya tümü)."""
    from ilim_assistant.motorlar.tercume_translate_memory import clear_session

    sf = (source_file or "").strip()
    tl = (tgt_lang or "").strip()
    if sf:
        clear_session(sf, tgt_lang=tl or None)
    else:
        clear_session("")
    return {"ok": True, "cleared": sf or "all", "tgt_lang": tl or None}


@app.post("/api/tercume/translate-chunk")
def api_tercume_translate_chunk(
    text: str = Form(""),
    src_lang: str = Form("auto"),
    tgt_lang: str = Form("en"),
    source_file: str = Form(""),
    page_index: str = Form(""),
    read_level: str = Form("akademik"),
    system_level: str = Form(""),
    workspace_root: str = Form(""),
):
    from ilim_assistant.motorlar.tercume_atolye import append_apprentice_log, translate_chunk

    pidx: int | None = None
    if (page_index or "").strip().isdigit():
        pidx = int(page_index.strip())
    sl = (system_level or read_level or "akademik").strip()
    result = translate_chunk(
        text,
        src_lang=src_lang,
        tgt_lang=tgt_lang,
        source_file=source_file,
        page_index=pidx,
        read_level=read_level,
        system_level=sl,
    )
    root = (workspace_root or "").strip() or None
    if result.get("ok"):
        append_apprentice_log(
            root,
            {
                "lesson": "translate_chunk",
                "source_file": source_file,
                "tgt_lang": tgt_lang,
                "page_index": pidx,
                "chars_in": len(text or ""),
                "chars_out": len(result.get("text") or ""),
                "note": "Tercüme: hedef dil imla kurallarina uygun parca cevirisi.",
            },
        )
    return result


@app.post("/api/tercume/quality-score")
def api_tercume_quality_score(
    source_text: str = Form(""),
    target_text: str = Form(""),
    tgt_lang: str = Form("tr"),
) -> dict[str, Any]:
    """CAT-lite: düzenlenen segment için hızlı kalite skoru."""
    from ilim_assistant.motorlar.tercume_translate_quality import score_translation

    src = (source_text or "").strip()
    tgt = (target_text or "").strip()
    if not src and not tgt:
        raise HTTPException(status_code=400, detail="source_text veya target_text gerekli")
    return {"ok": True, "quality": score_translation(src, tgt, tgt_lang=(tgt_lang or "tr").strip())}


@app.post("/api/tercume/academic-check")
def api_tercume_academic_check(
    source_text: str = Form(""),
    target_text: str = Form(""),
) -> dict[str, Any]:
    """Calsima paneli: tez/akademik kaynak izi hizli denetimi."""
    from ilim_assistant.motorlar.tercume_academic_check import analyze_academic_support

    src = (source_text or "").strip()
    tgt = (target_text or "").strip()
    if not src and not tgt:
        raise HTTPException(status_code=400, detail="source_text veya target_text gerekli")
    return analyze_academic_support(src, tgt)


@app.post("/api/tercume/import-file")
async def api_tercume_import_file(
    file: UploadFile = File(...),
) -> dict[str, Any]:
    """Dosya / ekran görüntüsü → arsiv/tercume-imports."""
    import uuid

    out_dir = REPO_ROOT / "ilim-assistant" / "arsiv" / "tercume-imports"
    out_dir.mkdir(parents=True, exist_ok=True)
    name = (file.filename or "import.bin").replace("\\", "/").split("/")[-1]
    safe = re.sub(r"[^a-zA-Z0-9._\-]+", "_", name).strip("._") or f"import_{uuid.uuid4().hex[:8]}"
    target = out_dir / safe
    data = await file.read()
    if len(data) > 30_000_000:
        raise HTTPException(status_code=400, detail="Dosya çok büyük (30 MB).")
    target.write_bytes(data)
    rel = target.relative_to(REPO_ROOT).as_posix()
    return {"ok": True, "rel": rel, "bytes": len(data)}


# —— Mimar · Fotoğraf (Faz 4F-1) ——
@app.get("/api/mimar/fotograf/capabilities")
def api_mimar_fotograf_capabilities() -> dict[str, Any]:
    from ilim_assistant.motorlar.mimar_fotograf import capabilities

    return capabilities()


@app.get("/api/mimar/fotograf/list")
def api_mimar_fotograf_list() -> dict[str, Any]:
    from ilim_assistant.motorlar.mimar_fotograf import list_photos

    return list_photos(REPO_ROOT)


@app.get("/api/mimar/fotograf/file")
def api_mimar_fotograf_file(rel: str = Query("")):
    from ilim_assistant.motorlar.mimar_fotograf import resolve_foto_rel

    try:
        target = resolve_foto_rel(rel, REPO_ROOT)
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e)) from e
    except FileNotFoundError as e:
        raise HTTPException(status_code=404, detail=str(e)) from e
    media = "image/jpeg"
    if target.suffix.lower() == ".png":
        media = "image/png"
    elif target.suffix.lower() == ".webp":
        media = "image/webp"
    return FileResponse(str(target), media_type=media, filename=target.name)


@app.post("/api/mimar/fotograf/upload")
async def api_mimar_fotograf_upload(file: UploadFile = File(...)) -> dict[str, Any]:
    from ilim_assistant.motorlar.mimar_fotograf import upload_bytes

    data = await file.read()
    try:
        return upload_bytes(data, file.filename or "foto.jpg", REPO_ROOT)
    except RuntimeError as e:
        raise HTTPException(status_code=503, detail=str(e)) from e
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e)) from e


@app.post("/api/mimar/fotograf/moderate")
async def api_mimar_fotograf_moderate(body: MimarFotoModerateBody) -> dict[str, Any]:
    from ilim_assistant.motorlar.mimar_fotograf import apply_moderation

    try:
        return await run_in_threadpool(apply_moderation, body.rel, body.op, REPO_ROOT)
    except RuntimeError as e:
        raise HTTPException(status_code=503, detail=str(e)) from e
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e)) from e
    except FileNotFoundError as e:
        raise HTTPException(status_code=404, detail=str(e)) from e


@app.post("/api/mimar/fotograf/restore")
async def api_mimar_fotograf_restore(body: MimarFotoRelBody) -> dict[str, Any]:
    from ilim_assistant.motorlar.mimar_fotograf import restore_original

    try:
        return await run_in_threadpool(restore_original, body.rel, REPO_ROOT)
    except FileNotFoundError as e:
        raise HTTPException(status_code=404, detail=str(e)) from e
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e)) from e


@app.post("/api/mimar/fotograf/restoration/preview")
async def api_mimar_fotograf_restoration_preview(body: MimarFotoModerateBody) -> dict[str, Any]:
    from ilim_assistant.motorlar.mimar_fotograf import preview_restoration

    try:
        return await run_in_threadpool(preview_restoration, body.rel, body.op, REPO_ROOT)
    except RuntimeError as e:
        raise HTTPException(status_code=503, detail=str(e)) from e
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e)) from e
    except FileNotFoundError as e:
        raise HTTPException(status_code=404, detail=str(e)) from e


@app.post("/api/mimar/fotograf/restoration/apply")
async def api_mimar_fotograf_restoration_apply(body: MimarFotoModerateBody) -> dict[str, Any]:
    from ilim_assistant.motorlar.mimar_fotograf import apply_restoration

    try:
        return await run_in_threadpool(apply_restoration, body.rel, body.op, REPO_ROOT)
    except RuntimeError as e:
        raise HTTPException(status_code=503, detail=str(e)) from e
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e)) from e
    except FileNotFoundError as e:
        raise HTTPException(status_code=404, detail=str(e)) from e


# —— Mimar · Sanat (Faz 4S-1) ——
@app.post("/api/mimar/atolye/parse")
def api_mimar_atolye_parse(body: MimarAtolyeParseBody) -> dict[str, Any]:
    from ilim_assistant.motorlar import mimar_faz5  # noqa: F401 — ROK kaydı

    return mimar_faz5.atolye_parse_payload(body.message)


@app.get("/api/mimar/sanat/capabilities")
def api_mimar_sanat_capabilities() -> dict[str, Any]:
    from ilim_assistant.motorlar.mimar_sanat import capabilities

    return capabilities()


@app.get("/api/mimar/sanat/list")
def api_mimar_sanat_list() -> dict[str, Any]:
    from ilim_assistant.motorlar.mimar_sanat import list_works

    return list_works(REPO_ROOT)


@app.get("/api/mimar/sanat/file")
def api_mimar_sanat_file(rel: str = Query("")):
    from ilim_assistant.motorlar.mimar_sanat import resolve_sanat_rel

    try:
        target = resolve_sanat_rel(rel, REPO_ROOT)
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e)) from e
    except FileNotFoundError as e:
        raise HTTPException(status_code=404, detail=str(e)) from e
    media = "image/jpeg"
    if target.suffix.lower() == ".png":
        media = "image/png"
    elif target.suffix.lower() == ".webp":
        media = "image/webp"
    elif target.suffix.lower() == ".gif":
        media = "image/gif"
    elif target.suffix.lower() == ".svg":
        media = "image/svg+xml"
    return FileResponse(str(target), media_type=media, filename=target.name)


@app.post("/api/mimar/sanat/upload")
async def api_mimar_sanat_upload(file: UploadFile = File(...)) -> dict[str, Any]:
    from ilim_assistant.motorlar.mimar_sanat import upload_bytes

    data = await file.read()
    try:
        return upload_bytes(data, file.filename or "eser.jpg", REPO_ROOT)
    except RuntimeError as e:
        raise HTTPException(status_code=503, detail=str(e)) from e
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e)) from e


@app.post("/api/mimar/sanat/meta")
async def api_mimar_sanat_meta(body: MimarSanatMetaBody) -> dict[str, Any]:
    from ilim_assistant.motorlar.mimar_sanat import update_metadata

    try:
        return await run_in_threadpool(
            update_metadata,
            body.rel,
            title=body.title,
            artist=body.artist,
            period=body.period,
            technique=body.technique,
            notes=body.notes,
            repo_root=REPO_ROOT,
        )
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e)) from e
    except FileNotFoundError as e:
        raise HTTPException(status_code=404, detail=str(e)) from e


@app.post("/api/mimar/sanat/identify")
async def api_mimar_sanat_identify(body: MimarSanatRelBody) -> dict[str, Any]:
    from ilim_assistant.motorlar.mimar_sanat import identify_work

    try:
        return await run_in_threadpool(identify_work, body.rel, REPO_ROOT)
    except RuntimeError as e:
        raise HTTPException(status_code=503, detail=str(e)) from e
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e)) from e
    except FileNotFoundError as e:
        raise HTTPException(status_code=404, detail=str(e)) from e


@app.post("/api/mimar/sanat/analyze")
async def api_mimar_sanat_analyze(body: MimarSanatAnalyzeBody) -> dict[str, Any]:
    from ilim_assistant.motorlar.mimar_sanat import analyze_work

    try:
        return await run_in_threadpool(analyze_work, body.rel, body.depth, REPO_ROOT)
    except RuntimeError as e:
        raise HTTPException(status_code=503, detail=str(e)) from e
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e)) from e
    except FileNotFoundError as e:
        raise HTTPException(status_code=404, detail=str(e)) from e


@app.post("/api/mimar/sanat/sketch")
async def api_mimar_sanat_sketch(body: MimarSanatRelBody) -> dict[str, Any]:
    from ilim_assistant.motorlar.mimar_sanat import sketch_work

    try:
        return await run_in_threadpool(sketch_work, body.rel, REPO_ROOT)
    except RuntimeError as e:
        raise HTTPException(status_code=503, detail=str(e)) from e
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e)) from e
    except FileNotFoundError as e:
        raise HTTPException(status_code=404, detail=str(e)) from e


@app.post("/api/mimar/sanat/copy")
async def api_mimar_sanat_copy(body: MimarSanatCopyBody) -> dict[str, Any]:
    from ilim_assistant.motorlar.mimar_sanat import copy_work

    try:
        return await run_in_threadpool(copy_work, body.rel, body.mode, REPO_ROOT)
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e)) from e
    except FileNotFoundError as e:
        raise HTTPException(status_code=404, detail=str(e)) from e


# —— Mimar · Tasarım (Faz 4T-1) ——
@app.get("/api/mimar/tasarim/capabilities")
def api_mimar_tasarim_capabilities() -> dict[str, Any]:
    from ilim_assistant.motorlar.mimar_tasarim import capabilities

    return capabilities()


@app.get("/api/mimar/tasarim/project/list")
def api_mimar_tasarim_project_list() -> dict[str, Any]:
    from ilim_assistant.motorlar.mimar_tasarim import list_projects

    return list_projects(REPO_ROOT)


@app.post("/api/mimar/tasarim/project/new")
async def api_mimar_tasarim_project_new(name: str = Query("Yeni plan")) -> dict[str, Any]:
    from ilim_assistant.motorlar.mimar_tasarim import new_project

    return await run_in_threadpool(new_project, name, 960, 540, REPO_ROOT)


@app.post("/api/mimar/tasarim/project/save")
async def api_mimar_tasarim_project_save(body: MimarTasarimSaveBody) -> dict[str, Any]:
    from ilim_assistant.motorlar.mimar_tasarim import save_project

    try:
        return await run_in_threadpool(save_project, body.project, REPO_ROOT)
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e)) from e


@app.post("/api/mimar/tasarim/project/load")
async def api_mimar_tasarim_project_load(body: MimarTasarimLoadBody) -> dict[str, Any]:
    from ilim_assistant.motorlar.mimar_tasarim import load_project

    try:
        return await run_in_threadpool(load_project, body.project_id, REPO_ROOT)
    except FileNotFoundError as e:
        raise HTTPException(status_code=404, detail=str(e)) from e


@app.post("/api/mimar/tasarim/reference/upload")
async def api_mimar_tasarim_reference_upload(file: UploadFile = File(...)) -> dict[str, Any]:
    from ilim_assistant.motorlar.mimar_tasarim import upload_reference

    data = await file.read()
    try:
        return upload_reference(data, file.filename or "ref.jpg", REPO_ROOT)
    except RuntimeError as e:
        raise HTTPException(status_code=503, detail=str(e)) from e
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e)) from e


@app.get("/api/mimar/tasarim/reference/file")
def api_mimar_tasarim_reference_file(rel: str = Query("")):
    from ilim_assistant.motorlar.mimar_tasarim import resolve_tasarim_rel

    try:
        target = resolve_tasarim_rel(rel, REPO_ROOT)
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e)) from e
    except FileNotFoundError as e:
        raise HTTPException(status_code=404, detail=str(e)) from e
    media = "image/jpeg"
    if target.suffix.lower() == ".png":
        media = "image/png"
    elif target.suffix.lower() == ".webp":
        media = "image/webp"
    return FileResponse(str(target), media_type=media, filename=target.name)


@app.post("/api/mimar/tasarim/sketch/from-image")
async def api_mimar_tasarim_sketch_from_image(body: MimarTasarimSketchImageBody) -> dict[str, Any]:
    from ilim_assistant.motorlar.mimar_tasarim import sketch_from_image

    try:
        return await run_in_threadpool(
            sketch_from_image, body.rel, body.width, body.height, REPO_ROOT
        )
    except RuntimeError as e:
        raise HTTPException(status_code=503, detail=str(e)) from e
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e)) from e
    except FileNotFoundError as e:
        raise HTTPException(status_code=404, detail=str(e)) from e


@app.post("/api/mimar/tasarim/sketch/from-text")
async def api_mimar_tasarim_sketch_from_text(body: MimarTasarimSketchTextBody) -> dict[str, Any]:
    from ilim_assistant.motorlar.mimar_tasarim import sketch_from_text

    try:
        return await run_in_threadpool(
            sketch_from_text, body.text, body.width, body.height, REPO_ROOT
        )
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e)) from e


@app.post("/api/mimar/tasarim/sketch/from-chat-handoff")
async def api_mimar_tasarim_sketch_from_chat_handoff(
    body: MimarTasarimChatHandoffBody,
) -> dict[str, Any]:
    from ilim_assistant.motorlar.mimar_tasarim import sketch_from_chat_handoff

    try:
        return await run_in_threadpool(
            sketch_from_chat_handoff,
            body.user,
            body.assistant,
            body.notes,
            body.width,
            body.height,
            body.project_id,
            REPO_ROOT,
        )
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e)) from e
    except FileNotFoundError as e:
        raise HTTPException(status_code=404, detail=str(e)) from e


@app.post("/api/mimar/tasarim/project/duplicate")
async def api_mimar_tasarim_project_duplicate(body: MimarTasarimDuplicateBody) -> dict[str, Any]:
    from ilim_assistant.motorlar.mimar_tasarim import duplicate_project

    try:
        return await run_in_threadpool(duplicate_project, body.project_id, body.name, REPO_ROOT)
    except FileNotFoundError as e:
        raise HTTPException(status_code=404, detail=str(e)) from e


@app.post("/api/mimar/tasarim/project/regenerate")
async def api_mimar_tasarim_project_regenerate(body: MimarTasarimLoadBody) -> dict[str, Any]:
    from ilim_assistant.motorlar.mimar_tasarim import regenerate_project

    try:
        return await run_in_threadpool(regenerate_project, body.project_id, REPO_ROOT)
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e)) from e
    except FileNotFoundError as e:
        raise HTTPException(status_code=404, detail=str(e)) from e


@app.get("/api/mimar/tasarim/project/export-png")
def api_mimar_tasarim_project_export_png(project_id: str = Query("")):
    from ilim_assistant.motorlar.mimar_tasarim import export_project_png, resolve_tasarim_rel

    pid = (project_id or "").strip()
    if not pid:
        raise HTTPException(status_code=400, detail="project_id gerekli.")
    try:
        info = export_project_png(pid, REPO_ROOT)
        target = resolve_tasarim_rel(info["rel"], REPO_ROOT)
    except FileNotFoundError as e:
        raise HTTPException(status_code=404, detail=str(e)) from e
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e)) from e
    return FileResponse(str(target), media_type="image/png", filename=target.name)


@app.post("/api/mimar/fotograf/ocr")
async def api_mimar_fotograf_ocr(body: MimarFotoVoiceBody) -> dict[str, Any]:
    from ilim_assistant.motorlar.mimar_fotograf import ocr_photo

    try:
        return await run_in_threadpool(ocr_photo, body.rel, REPO_ROOT, body.lang)
    except RuntimeError as e:
        raise HTTPException(status_code=503, detail=str(e)) from e
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e)) from e
    except FileNotFoundError as e:
        raise HTTPException(status_code=404, detail=str(e)) from e


@app.post("/api/mimar/fotograf/voice")
async def api_mimar_fotograf_voice(body: MimarFotoVoiceBody) -> dict[str, Any]:
    from ilim_assistant.motorlar.mimar_fotograf import prepare_voice_text

    try:
        return await run_in_threadpool(prepare_voice_text, body.rel, body.mode, REPO_ROOT)
    except RuntimeError as e:
        raise HTTPException(status_code=503, detail=str(e)) from e
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e)) from e
    except FileNotFoundError as e:
        raise HTTPException(status_code=404, detail=str(e)) from e


@app.post("/api/code/run")
async def api_code_run(body: CodeRunBody):
    """Programlama Atölyesi: Python/JavaScript kodunu kontrollü çalıştırır.

    Kod geçici dosyaya yazılır; workspace_root + cwd_rel ile çalışma dizini proje içine alınır.
    """
    code = (body.code or "").strip()
    if not code:
        raise HTTPException(status_code=400, detail="Çalıştırılacak kod boş.")
    if len(code) > 80_000:
        raise HTTPException(status_code=400, detail="Kod çok büyük (80 KB sınırı).")

    lang = (body.language or "python").strip().lower()
    timeout = max(1.0, min(float(body.timeout_sec or 8.0), 20.0))
    if lang in ("py", "python", "python3"):
        filename = "main.py"
        runner = sys.executable
    elif lang in ("js", "javascript", "node"):
        filename = "main.js"
        runner = "node"
    else:
        return {
            "ok": False,
            "stdout": "",
            "stderr": f"{lang} dili bu çalıştırma sürümünde desteklenmiyor. İlk sürüm: Python + JavaScript.",
            "exit_code": None,
            "timeout": False,
        }

    wr_norm = _normalize_workspace_root(body.workspace_root)
    run_dir = None
    cwd_used_rel = None
    if wr_norm:
        run_dir = _resolve_workspace_run_dir(body.workspace_root or "", body.cwd_rel)
        if run_dir:
            try:
                cwd_used_rel = os.path.relpath(run_dir, wr_norm).replace("\\", "/")
            except ValueError:
                cwd_used_rel = ""

    def _run() -> dict[str, Any]:
        work_td = tempfile.mkdtemp(prefix="ruzgar_code_")
        src = os.path.join(work_td, filename)
        try:
            with open(src, "w", encoding="utf-8", newline="\n") as f:
                f.write(code)
            abs_src = os.path.abspath(src)
            proc_cwd = run_dir if run_dir else work_td
            cmd = [runner, abs_src]
            try:
                proc = subprocess.run(
                    cmd,
                    cwd=proc_cwd,
                    capture_output=True,
                    text=True,
                    encoding="utf-8",
                    errors="replace",
                    timeout=timeout,
                    shell=False,
                )
                out: dict[str, Any] = {
                    "ok": proc.returncode == 0,
                    "stdout": (proc.stdout or "")[:20_000],
                    "stderr": (proc.stderr or "")[:20_000],
                    "exit_code": proc.returncode,
                    "timeout": False,
                }
                if cwd_used_rel is not None:
                    out["cwd_used_rel"] = cwd_used_rel
                return out
            except subprocess.TimeoutExpired as e:
                out = {
                    "ok": False,
                    "stdout": (e.stdout or "")[:20_000] if isinstance(e.stdout, str) else "",
                    "stderr": f"Zaman aşımı: {timeout:.1f} saniye içinde bitmedi.",
                    "exit_code": None,
                    "timeout": True,
                }
                if cwd_used_rel is not None:
                    out["cwd_used_rel"] = cwd_used_rel
                return out
            except FileNotFoundError as e:
                out = {
                    "ok": False,
                    "stdout": "",
                    "stderr": f"Çalıştırıcı bulunamadı: {runner} ({e})",
                    "exit_code": None,
                    "timeout": False,
                }
                if cwd_used_rel is not None:
                    out["cwd_used_rel"] = cwd_used_rel
                return out
        finally:
            shutil.rmtree(work_td, ignore_errors=True)

    return await run_in_threadpool(_run)


def _normalize_stt_lang(lang: str | None) -> str | None:
    if not lang:
        return None
    s = lang.strip().lower()
    if s in ("auto", "detect", "none"):
        return None
    return lang.strip()[:16]


@app.post("/api/stt")
async def api_stt(
    file: UploadFile = File(...),
    lang: Annotated[str | None, Query()] = "tr",
):
    """Tarayıcıdan gelen ses dosyasını yerel Whisper ile metne çevirir (internet şart değil)."""
    if not stt_runtime_available():
        raise HTTPException(
            status_code=503,
            detail="Yerel STT kullanılamıyor: pip install faster-whisper veya RUZGAR_STT=0 kaldırın.",
        )
    raw = await file.read()
    if not raw or len(raw) < 64:
        raise HTTPException(status_code=400, detail="Ses verisi çok kısa veya boş.")

    sfx = ".webm"
    fn = (file.filename or "").lower()
    if fn.endswith(".wav"):
        sfx = ".wav"
    elif fn.endswith(".ogg"):
        sfx = ".ogg"
    elif fn.endswith(".mp4") or fn.endswith(".m4a"):
        sfx = ".m4a"

    wl = _normalize_stt_lang(lang)

    tmp_path = None
    try:
        fd, tmp_path = tempfile.mkstemp(suffix=sfx)
        os.write(fd, raw)
        os.close(fd)

        def _run():
            return transcribe_file(tmp_path, language=wl)

        text, detected = await run_in_threadpool(_run)
        return {"text": text, "language": detected}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=400, detail=str(e)) from e
    finally:
        if tmp_path and os.path.isfile(tmp_path):
            try:
                os.unlink(tmp_path)
            except OSError:
                pass


@app.get("/api/ses/settings")
def api_ses_settings_get():
    """Aktif ses profili ve hız/huzur çarpanları (.ruzgar_ses_ayarlari.json)."""
    from ilim_assistant.tts_service import read_ses_ayarlari

    return read_ses_ayarlari()


@app.post("/api/ses/settings")
def api_ses_settings_patch(body: SesSettingsPatchBody):
    from ilim_assistant.motorlar.ses_motoru import normalize_ses_karakteri
    from ilim_assistant.tts_service import read_ses_ayarlari, write_ses_ayarlari

    patch: dict[str, Any] = {}
    if body.karakter is not None:
        patch["karakter"] = normalize_ses_karakteri(body.karakter).value
    if body.hiz is not None:
        patch["hiz"] = max(0.45, min(1.0, float(body.hiz)))
    if body.huzur is not None:
        patch["huzur"] = max(0.45, min(1.0, float(body.huzur)))
    if patch:
        write_ses_ayarlari(patch)
    return read_ses_ayarlari()


@app.post("/api/tts")
async def api_tts(body: TtsBody):
    """Edge-TTS ile MP3 üretir (Ses atölyesi + sohbet sesli yanıt)."""
    from ilim_assistant.motorlar.ses_motoru import (
        EDGE_VOICES,
        analiz_icerik_yolu,
        edge_pitch_string,
        edge_rate_yuzdesi,
        normalize_ses_karakteri,
        tts_metadata_kimlik,
    )
    from ilim_assistant.tts_service import (
        edge_available,
        read_ses_ayarlari,
        synthesize_edge_mp3_in_process,
    )

    text = (body.text or "").strip()
    if len(text) < 2:
        raise HTTPException(status_code=400, detail="Metin çok kısa.")
    if len(text) > 12_000:
        text = text[:12_000]

    if (body.backend or "edge").strip().lower() != "edge":
        raise HTTPException(status_code=400, detail="Yalnızca edge backend destekleniyor.")
    if not edge_available():
        raise HTTPException(
            status_code=503,
            detail="Edge-TTS yok: pip install edge-tts",
        )

    ayar = read_ses_ayarlari()
    kar = normalize_ses_karakteri(body.karakter or ayar.get("karakter"))
    icerik = analiz_icerik_yolu(text)
    voice = EDGE_VOICES[kar]
    rate = edge_rate_yuzdesi(
        karakter=kar,
        icerik=icerik,
        hiz_carpani=float(ayar.get("hiz", 0.92)),
        huzur_carpani=float(ayar.get("huzur", 0.88)),
    )
    pitch = edge_pitch_string(kar, icerik)
    meta = tts_metadata_kimlik(
        karakter=kar.value,
        icerik_yolu=icerik.value,
        edge_voice=voice,
        ek={"emotion": body.emotion} if body.emotion else None,
    )
    try:
        out_path = await synthesize_edge_mp3_in_process(
            text,
            voice=voice,
            rate=rate,
            pitch=pitch,
            meta=meta,
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e)) from e

    return FileResponse(
        str(out_path),
        media_type="audio/mpeg",
        filename="ruzgar-tts.mp3",
    )


def _iter_instant_chat_events(
    reply: str,
    user_message: str,
    *,
    session_wake_used: bool,
    msg_for_wake: str,
    orch: dict[str, Any] | None = None,
    instant_gundelik: bool = False,
    instant_clarify: bool = False,
    egitim_instant: bool = False,
    programlama_instant: bool = False,
    programlama_focus_rel: str | None = None,
    programlama_project_rel: str | None = None,
    programlama_expand_tree: bool = False,
) -> Iterator[dict]:
    """Ollama/RAG beklemeden tek tur bitir (SSE/WS)."""
    full_out = finalize_assistant_reply(reply)
    yield {"type": "token", "text": full_out}
    new_wake = session_wake_used or message_calls_wake_name(msg_for_wake)
    done: dict[str, Any] = {
        "type": "done",
        "full_reply": full_out,
        "user_message": user_message,
        "new_wake_used": new_wake,
    }
    if orch:
        try:
            from ilim_assistant.ana_motor_agent import mark_agent_answer_done

            if orch.get("agent_steps"):
                orch["agent_steps"] = mark_agent_answer_done(orch["agent_steps"])
        except Exception:
            pass
        done["orchestra"] = orch
    if instant_gundelik:
        done["instant_gundelik"] = True
    if egitim_instant:
        done["egitim_instant"] = True
    if programlama_instant:
        done["programlama_instant"] = True
    if programlama_focus_rel:
        done["programlama_focus_rel"] = programlama_focus_rel
    if programlama_project_rel:
        done["programlama_project_rel"] = programlama_project_rel
    if programlama_expand_tree:
        done["programlama_expand_tree"] = True
    if instant_clarify:
        done["instant_clarify"] = True
    yield done


def _yield_programlama_instant(
    raw: str | dict[str, Any] | None,
    user_message: str,
    *,
    session_wake_used: bool,
    msg_for_wake: str,
    orch: dict[str, Any] | None = None,
) -> Iterator[dict]:
    from ilim_assistant.motorlar.programlama_motoru import unpack_programlama_instant

    text, meta = unpack_programlama_instant(raw)
    if not text:
        return
    yield from _iter_instant_chat_events(
        text,
        user_message,
        session_wake_used=session_wake_used,
        msg_for_wake=msg_for_wake,
        orch=orch or {},
        instant_gundelik=True,
        programlama_instant=True,
        programlama_focus_rel=str(meta.get("focus_rel") or "") or None,
        programlama_project_rel=str(meta.get("project_rel") or "") or None,
        programlama_expand_tree=bool(meta.get("expand_tree")),
    )


def _iter_chat_turn_events_impl(req: ChatRequest) -> Iterator[dict]:
    """SSE/WS: Ana motor durumları → prepare_turn (İdrak + orkestra) → LLM."""
    orch_early: dict[str, Any] = {}
    msg_early = (req.message or "").strip()
    if msg_early:
        try:
            from ilim_assistant.kullanici_baglami import ingest_message

            ingest_message(msg_early)
        except Exception:
            pass
        _skip_early_hub = False
        try:
            from ilim_assistant.ana_motor_plan import looks_like_casual_social_chat

            _skip_early_hub = looks_like_casual_social_chat(msg_early)
        except Exception:
            pass
        if not _skip_early_hub:
            try:
                from ilim_assistant.chat_core import normalize_mode
                from ilim_assistant.motorlar.ana_motor_hub_faz76 import maybe_hub_instant

                if normalize_mode(_effective_chat_mode_raw(req)) == "genel":
                    from ilim_assistant.idrak_entegrasyon import motor_niyeti_heuristic

                    hub_early = maybe_hub_instant(
                        msg_early,
                        motor_flags=motor_niyeti_heuristic(msg_early),
                        workspace_root=req.workspace_root,
                    )
                    if hub_early:
                        yield from _iter_instant_chat_events(
                            hub_early,
                            msg_early,
                            session_wake_used=req.session_wake_used,
                            msg_for_wake=req.message,
                            orch=orch_early,
                            instant_gundelik=True,
                        )
                        return
            except Exception:
                pass
        try:
            from ilim_assistant.nebula_kitap_hafiza import try_consume_nebula_kitap_command

            kitap_reply = try_consume_nebula_kitap_command(msg_early)
            if kitap_reply:
                yield from _iter_instant_chat_events(
                    kitap_reply,
                    msg_early,
                    session_wake_used=req.session_wake_used,
                    msg_for_wake=req.message,
                    orch=orch_early,
                    instant_gundelik=True,
                )
                return
        except Exception:
            pass
        try:
            from ilim_assistant.kisisel_hafiza import try_consume_memory_command

            mem_reply = try_consume_memory_command(msg_early)
            if mem_reply:
                yield from _iter_instant_chat_events(
                    mem_reply,
                    msg_early,
                    session_wake_used=req.session_wake_used,
                    msg_for_wake=req.message,
                    orch=orch_early,
                    instant_gundelik=True,
                )
                return
        except Exception:
            pass
        try:
            from ilim_assistant.chat_core import normalize_mode
            from ilim_assistant.motorlar.programlama_motoru import (
                is_programlama_reserved_command,
                maybe_programlama_instant_reply,
            )

            _mode_early = normalize_mode(_effective_chat_mode_raw(req))
            _coding_early = bool(req.coding_mode) or _mode_early == "programlama"
            if _coding_early or _mode_early == "programlama":
                if is_programlama_reserved_command(msg_early):
                    _prog_early = maybe_programlama_instant_reply(
                        msg_early,
                        "programlama",
                        workspace_root=req.workspace_root,
                        active_file=req.programlama_active_file,
                        editor_snippet=req.programlama_editor_snippet,
                    )
                    if _prog_early:
                        yield from _yield_programlama_instant(
                            _prog_early,
                            msg_early,
                            session_wake_used=req.session_wake_used,
                            msg_for_wake=req.message,
                            orch=orch_early,
                        )
                        return
        except Exception:
            pass
        try:
            from ilim_assistant.ruzgar_egitim import try_consume_egitim_command

            egitim_reply = try_consume_egitim_command(msg_early, req.history)
            if egitim_reply:
                yield from _iter_instant_chat_events(
                    egitim_reply,
                    msg_early,
                    session_wake_used=req.session_wake_used,
                    msg_for_wake=req.message,
                    orch=orch_early,
                    instant_gundelik=True,
                )
                return
        except Exception:
            pass

    from ilim_assistant.idrak_on_islem import pretreat_user_turn

    pt = pretreat_user_turn(req.message, req.history)
    if pt.text != (req.message or "").strip():
        req = req.model_copy(update={"message": pt.text})
    mode_raw = _effective_chat_mode_raw(req)
    mode_norm = normalize_mode(mode_raw)
    coding = req.coding_mode or mode_norm == "programlama"
    _delegated_from_genel = bool(
        (req.mode or "").strip().lower() in ("genel", "gelisim", "uretim", "")
        and not req.coding_mode
        and mode_norm == "programlama"
    )
    if mode_norm == "programlama":
        try:
            from ilim_assistant.motorlar.programlama_faz5 import sync_client_editor_state

            sync_client_editor_state(
                req.workspace_root,
                active_file=req.programlama_active_file,
                editor_snippet=req.programlama_editor_snippet,
                language=req.programlama_language,
            )
        except Exception:
            pass
    route_meta: dict[str, Any] = {
        "mode_raw": (req.mode or "").strip() or None,
        "mode_effective": mode_norm,
        "coding_mode": coding,
        "idrak_surface": {
            "changed": pt.changed,
            "continuation": pt.continuation,
            "replacements": list(pt.replacements)[:8],
        },
    }
    _p92_block_msg = ""
    if mode_norm == "programlama":
        try:
            from ilim_assistant.motorlar.programlama_faz101_report_read import (
                maybe_instant_report_read,
                wants_report_read,
            )

            if wants_report_read(req.message or ""):
                _rep101 = maybe_instant_report_read(
                    req.message or "",
                    req.workspace_root,
                )
                if _rep101:
                    _p92_block_msg = _rep101
        except Exception:
            pass
    if mode_norm == "programlama" and not _p92_block_msg:
        try:
            from ilim_assistant.motorlar.programlama_faz92 import (
                assess_risk,
                build_task_plan,
                render_plan_directive,
                risk_confirmation_text,
            )
            from ilim_assistant.motorlar.programlama_faz93 import (
                build_decision_context,
                build_refactor_pilot_plan,
                looks_like_refactor_pilot,
                render_refactor_pilot_directive,
            )
            from ilim_assistant.motorlar.programlama_faz94 import build_pre_turn_test_directive
            from ilim_assistant.motorlar.programlama_faz95 import (
                build_prompt_cache_key,
                get_prompt_cache,
                optimize_prompt_blocks,
                prompt_cache_metrics,
                set_prompt_cache,
            )
            from ilim_assistant.kullanici_baglami import build_programming_style_directive

            p92 = build_task_plan(req.message or "")
            r92 = assess_risk(req.message or "")
            route_meta["task_plan_v92"] = p92
            route_meta["task_risk_v92"] = r92
            try:
                _sel_prog = select_brain_chain(
                    message=req.message or "",
                    mode_norm="programlama",
                    coding_mode=bool(req.coding_mode),
                )
                route_meta["p9_chain_v99"] = [
                    e.profile_id for e in list(_sel_prog.chain or [])[:6]
                ]
                route_meta["p9_local_first_v99"] = bool(
                    route_meta["p9_chain_v99"]
                    and route_meta["p9_chain_v99"][0] in ("kod", "denge", "hizli")
                )
            except Exception:
                pass
            if bool(r92.get("requires_confirmation")):
                _p92_block_msg = risk_confirmation_text(r92)
            p92_dir = render_plan_directive(p92).strip()
            p93_ctx = build_decision_context(req.workspace_root).strip()
            p5_style = build_programming_style_directive().strip()
            route_meta["style_pref_v95"] = bool(p5_style)
            p6_dir = ""
            p7_pre = ""
            _p_scope = ""
            try:
                from ilim_assistant.motorlar.programlama_faz10 import resolve_scope_rel

                _p_scope = resolve_scope_rel(
                    req.workspace_root,
                    active_file=req.programlama_active_file,
                )
                if looks_like_refactor_pilot(req.message or ""):
                    _pilot = build_refactor_pilot_plan(
                        req.message or "",
                        req.workspace_root,
                        _p_scope,
                    )
                    route_meta["refactor_pilot_v96"] = _pilot
                    p6_dir = render_refactor_pilot_directive(_pilot).strip()
                p7_pre = build_pre_turn_test_directive(
                    req.message or "",
                    req.workspace_root,
                    _p_scope,
                ).strip()
                if p7_pre:
                    route_meta["regression_shield_pre_v97"] = True
            except Exception:
                pass
            if p92_dir or p93_ctx or p5_style or p6_dir or p7_pre:
                _blocks = [p92_dir, p93_ctx, p5_style, p6_dir, p7_pre]
                _cache_key = build_prompt_cache_key(
                    workspace_root=str(req.workspace_root or ""),
                    scope_rel=str(_p_scope or ""),
                    user_message=req.message or "",
                    blocks=_blocks,
                )
                pre = get_prompt_cache(_cache_key, ttl_sec=75).strip()
                route_meta["prompt_cache_v98"] = "hit" if pre else "miss"
                if not pre:
                    pre = optimize_prompt_blocks(
                        _blocks,
                        per_block_limit=1700,
                        total_limit=5200,
                    ).strip()
                    if pre:
                        set_prompt_cache(_cache_key, pre)
                route_meta["prompt_optimized_v98"] = bool(pre)
                route_meta["prompt_cache_metrics_v98"] = prompt_cache_metrics()
                req = req.model_copy(
                    update={"message": f"{pre}\n\n[Kullanici istegi]\n{req.message or ''}"}
                )
        except Exception:
            pass
    if _delegated_from_genel:
        route_meta["programlama_delegated"] = True
        try:
            from ilim_assistant.ana_motor_faz59 import (
                classify_turn_intent,
                enrich_handoff_with_intent,
            )

            _mf59: dict[str, bool] = {}
            try:
                from ilim_assistant.idrak_entegrasyon import motor_niyeti_heuristic

                _mf59 = motor_niyeti_heuristic(req.message or "")
            except Exception:
                pass
            _intent59 = classify_turn_intent(
                req.message or "",
                mode_norm=(req.mode or "").strip().lower() or "genel",
                coding_mode=bool(req.coding_mode),
                motor_flags=_mf59,
            )
            route_meta["intent_v59"] = _intent59
            if _intent59.get("programming_budget_transferred"):
                route_meta["programming_budget_transferred"] = True
        except Exception:
            _intent59 = {}
        try:
            from ilim_assistant.motorlar.programlama_faz79 import build_handoff_packet_v3

            _hub_meta = route_meta.get("hub_delegate") or {}
            _handoff = build_handoff_packet_v3(
                req.message or "",
                req.workspace_root,
                active_file=getattr(req, "programlama_active_file", None),
                hub_meta=_hub_meta if isinstance(_hub_meta, dict) else None,
            )
            if _handoff.get("ok"):
                route_meta["handoff"] = {
                    "scope_rel": _handoff.get("scope_rel"),
                    "template": _handoff.get("parsed_template"),
                }
                try:
                    from ilim_assistant.ana_motor_faz59 import enrich_handoff_with_intent

                    pkt = enrich_handoff_with_intent(
                        str(_handoff.get("packet_text") or ""),
                        req.message or "",
                        mode_norm=(req.mode or "").strip().lower() or "genel",
                    )
                    _handoff = {**_handoff, "packet_text": pkt}
                except Exception:
                    pass
        except Exception:
            _handoff = {}
    else:
        _handoff = {}
        _intent59 = {}
    yield {"type": "meta", "chat_route": route_meta}
    yield {"type": "status", "text": "Rüzgar hazırlanıyor…"}
    if _p92_block_msg:
        yield from _iter_instant_chat_events(
            _p92_block_msg,
            (req.message or "").strip(),
            session_wake_used=req.session_wake_used,
            msg_for_wake=req.message,
            orch={},
            instant_gundelik=True,
        )
        return

    # Programlama Faz 96 (P11) — otonom analiz + onarım (Faz 2 onayından önce)
    _p11_repair_approved = False
    _p11_analysis_before: dict[str, Any] | None = None
    if mode_norm == "programlama":
        try:
            from ilim_assistant.motorlar.programlama_faz96 import (
                prepare_autonomous_repair_turn,
            )

            _p11_prep = prepare_autonomous_repair_turn(
                req.message or "", req.workspace_root
            )
            if _p11_prep:
                if _p11_prep.get("instant"):
                    yield from _iter_instant_chat_events(
                        str(_p11_prep["instant"]),
                        (req.message or "").strip(),
                        session_wake_used=req.session_wake_used,
                        msg_for_wake=req.message,
                        orch={},
                        instant_gundelik=True,
                    )
                    return
                if _p11_prep.get("augmented_message"):
                    req = req.model_copy(
                        update={"message": str(_p11_prep["augmented_message"])}
                    )
                    _p11_repair_approved = True
                    _p11_analysis_before = (
                        _p11_prep.get("analysis")
                        if isinstance(_p11_prep.get("analysis"), dict)
                        else None
                    )
        except Exception:
            pass

    # Programlama Faz 2 — onaylı düzeltme (LLM+pytest; anında yalnızca «önce tara» uyarısı)
    _scan_fix_approved = False
    if mode_norm == "programlama" and not _p11_repair_approved:
        try:
            from ilim_assistant.motorlar.programlama_faz92 import (
                needs_refactor_scope_clarification,
                refactor_scope_question,
            )

            if needs_refactor_scope_clarification(req.message or ""):
                yield from _iter_instant_chat_events(
                    refactor_scope_question(),
                    (req.message or "").strip(),
                    session_wake_used=req.session_wake_used,
                    msg_for_wake=req.message,
                    orch={},
                    instant_gundelik=True,
                )
                return
        except Exception:
            pass
        try:
            from ilim_assistant.motorlar.programlama_faz2 import prepare_scan_fix_turn

            _fix_prep = prepare_scan_fix_turn(
                req.message or "", req.workspace_root
            )
            if _fix_prep:
                if _fix_prep.get("instant"):
                    yield from _iter_instant_chat_events(
                        str(_fix_prep["instant"]),
                        (req.message or "").strip(),
                        session_wake_used=req.session_wake_used,
                        msg_for_wake=req.message,
                        orch={},
                        instant_gundelik=True,
                    )
                    return
                if _fix_prep.get("augmented_message"):
                    req = req.model_copy(
                        update={"message": str(_fix_prep["augmented_message"])}
                    )
                    _scan_fix_approved = True
        except Exception:
            pass

    # Sahip kilidi + Programlama anında yanıt — RAG/LLM/süre bütçesinden ÖNCE
    try:
        if mode_norm == "programlama" and not (_p11_repair_approved or _scan_fix_approved):
            from ilim_assistant.motorlar.programlama_motoru import (
                maybe_programlama_instant_reply,
            )

            _prog_hi = maybe_programlama_instant_reply(
                req.message or "",
                mode_norm,
                workspace_root=req.workspace_root,
                active_file=req.programlama_active_file,
                editor_snippet=req.programlama_editor_snippet,
            )
            if _prog_hi:
                yield from _yield_programlama_instant(
                    _prog_hi,
                    (req.message or "").strip(),
                    session_wake_used=req.session_wake_used,
                    msg_for_wake=req.message,
                    orch={},
                )
                return
        elif mode_norm == "programlama":
            pass
        else:
            from ilim_assistant.ruzgar_owner_lock import maybe_owner_instant_reply

            _owner_hi = maybe_owner_instant_reply(req.message or "", mode_norm)
            if _owner_hi:
                yield from _iter_instant_chat_events(
                    _owner_hi,
                    (req.message or "").strip(),
                    session_wake_used=req.session_wake_used,
                    msg_for_wake=req.message,
                    orch={},
                    instant_gundelik=True,
                )
                return
    except Exception:
        pass

    orch: dict[str, Any] = {}
    if isinstance(route_meta.get("task_plan_v92"), dict):
        orch["task_plan_v92"] = dict(route_meta["task_plan_v92"])
    try:
        from ilim_assistant.ruzgar_umed_cevap_emri import public_meta, umed_emri_applies

        if umed_emri_applies(mode_norm=mode_norm, coding_mode=coding):
            umed_meta = public_meta(req.message or "")
            orch["umed_cevap_emri"] = umed_meta
            yield {"type": "meta", "umed_cevap_emri": umed_meta}
    except Exception:
        pass

    # Bilişsel / empati — hatalı hafızadan ÖNCE («beni anlıyor musun» ≈ «sen beni anlıyor musun»)
    if mode_norm in ("genel", "uretim", "gelisim") and not coding:
        try:
            from ilim_assistant.ruzgar_bilissel_analiz import maybe_bilissel_instant_reply

            bilissel_hi = maybe_bilissel_instant_reply(
                req.message,
                history=req.history,
            )
            if bilissel_hi:
                yield from _iter_instant_chat_events(
                    bilissel_hi,
                    (req.message or "").strip(),
                    session_wake_used=req.session_wake_used,
                    msg_for_wake=req.message,
                    orch=orch,
                    instant_gundelik=True,
                )
                return
        except Exception:
            pass

    # Selam / kısa sohbet — eğitim hafızasından ÖNCE (nasılsın → öğretim şablonu yok)
    if mode_norm in ("genel", "uretim", "gelisim") and not coding:
        try:
            from ilim_assistant.ana_motor_plan import maybe_gundelik_instant_reply

            instant_hi = maybe_gundelik_instant_reply(req.message, mode_norm, {})
            if instant_hi:
                yield from _iter_instant_chat_events(
                    instant_hi,
                    (req.message or "").strip(),
                    session_wake_used=req.session_wake_used,
                    msg_for_wake=req.message,
                    orch=orch,
                    instant_gundelik=True,
                )
                return
        except Exception:
            pass

    # Ümit abi eğitimi — selam rüzgar vb. (empati + gündelik yukarıda ayrıldı)
    if mode_norm in ("genel", "uretim", "gelisim") and not coding:
        try:
            from ilim_assistant.ruzgar_egitim import maybe_egitim_learned_reply

            egitim_hi = maybe_egitim_learned_reply(
                req.message,
                history=req.history,
            )
            if egitim_hi:
                yield from _iter_instant_chat_events(
                    egitim_hi,
                    (req.message or "").strip(),
                    session_wake_used=req.session_wake_used,
                    msg_for_wake=req.message,
                    orch=orch,
                    egitim_instant=True,
                )
                return
        except Exception:
            pass

    # Tarih hızlı yol — plan / prefetch / ağır RAG öncesi (120 sn UI zaman aşımı)
    if mode_norm in ("genel", "uretim", "gelisim") and not coding:
        try:
            from ilim_assistant.tarih_fast import iter_tarih_hafiza_reply

            tarih_early = iter_tarih_hafiza_reply(
                req.message,
                req.history,
                mode_norm=mode_norm,
            )
            if tarih_early is not None:
                yield {
                    "type": "status",
                    "text": "Tarih hafızası — yerel pasaj + Ollama…",
                }
                reply_body = ""
                for piece in tarih_early:
                    reply_body += piece
                    yield {"type": "token", "text": piece}
                if reply_body.strip():
                    full_out = finalize_assistant_reply(reply_body)
                    yield {
                        "type": "done",
                        "full_reply": full_out,
                        "user_message": (req.message or "").strip(),
                        "new_wake_used": req.session_wake_used
                        or message_calls_wake_name(req.message),
                        "orchestra": orch,
                        "tarih_fast": True,
                        "tarih_fast_early": True,
                    }
                    return
                try:
                    from ilim_assistant.chat_core import _tarih_intent

                    if _tarih_intent(req.message):
                        yield from _iter_instant_chat_events(
                            (
                                "Tarih hızlı yol yanıt üretemedi (Ollama zaman aşımı). "
                                "`ollama serve` kontrol edin veya soruyu kısaltıp tekrar deneyin."
                            ),
                            (req.message or "").strip(),
                            session_wake_used=req.session_wake_used,
                            msg_for_wake=req.message,
                            orch=orch,
                        )
                        return
                except Exception:
                    pass
        except Exception:
            pass
    try:
        if mode_norm in ("genel", "uretim", "gelisim") and not coding:
            from ilim_assistant.motorlar.hafiza_faz75 import maybe_instant_faz75_hub

            _hf75 = maybe_instant_faz75_hub(req.message)
        else:
            from ilim_assistant.motorlar.hafiza_faz75 import maybe_instant_faz75

            _hf75 = maybe_instant_faz75(
                req.message, mode_norm=mode_norm, allow_lookup=False
            )
        if _hf75:
            yield from _iter_instant_chat_events(
                _hf75,
                (req.message or "").strip(),
                session_wake_used=req.session_wake_used,
                msg_for_wake=req.message,
                orch=orch,
                instant_gundelik=True,
            )
            return
    except Exception:
        pass
    for _consumer in (
        "ilim_assistant.dinamit_hatirlatici:try_consume_hatirlatici_intent",
        "ilim_assistant.nebula_kitap_hafiza:try_consume_nebula_kitap_command",
        "ilim_assistant.kisisel_hafiza:try_consume_memory_command",
        "ilim_assistant.gorev_yoneticisi:try_consume_task_command",
    ):
        try:
            mod_name, fn_name = _consumer.split(":", 1)
            mod = __import__(mod_name, fromlist=[fn_name])
            reply = getattr(mod, fn_name)(req.message)
            if reply:
                yield from _iter_instant_chat_events(
                    reply,
                    (req.message or "").strip(),
                    session_wake_used=req.session_wake_used,
                    msg_for_wake=req.message,
                    orch=orch,
                    instant_gundelik=True,
                )
                return
        except Exception:
            pass

    turn_plan = None
    motor_flags: dict[str, bool] = {}
    if os.environ.get("RUZGAR_ANA_MOTOR_PLAN", "1").strip().lower() not in (
        "0",
        "false",
        "no",
    ):
        if mode_norm in ("genel", "uretim", "gelisim"):
            try:
                from ilim_assistant.ana_motor_plan import plan_question
                from ilim_assistant.idrak_entegrasyon import motor_niyeti_heuristic

                motor_flags = motor_niyeti_heuristic(req.message)
                turn_plan = plan_question(req.message, mode_norm, motor_flags)
                yield {"type": "status", "text": turn_plan.status_text}
                plan_dict = turn_plan.to_dict()
                orch["plan"] = plan_dict
                yield {"type": "meta", "plan": plan_dict}
                try:
                    from ilim_assistant.ana_motor_plan import maybe_gundelik_instant_reply

                    instant = maybe_gundelik_instant_reply(
                        req.message,
                        mode_norm,
                        motor_flags,
                        question_plan=turn_plan,
                    )
                    if instant:
                        yield from _iter_instant_chat_events(
                            instant,
                            (req.message or "").strip(),
                            session_wake_used=req.session_wake_used,
                            msg_for_wake=req.message,
                            orch=orch,
                            instant_gundelik=True,
                        )
                        return
                except Exception:
                    pass
                try:
                    from ilim_assistant.hafiza_dogal_sentez import (
                        dogal_konus_enabled,
                        iter_hafiza_dogal_reply,
                        lookup_genel_hafiza_hint,
                        should_skip_hafiza_dogal,
                    )

                    if dogal_konus_enabled() and not should_skip_hafiza_dogal(
                        req.message
                    ):
                        hint = lookup_genel_hafiza_hint(req.message)
                        if hint:
                            hstream = iter_hafiza_dogal_reply(
                                req.message,
                                req.history,
                                mode_norm=mode_norm,
                                hint=hint,
                            )
                            if hstream is not None:
                                yield {
                                    "type": "status",
                                    "text": "Yanıtlıyorum…",
                                }
                                reply_body = ""
                                for piece in hstream:
                                    reply_body += piece
                                    yield {"type": "token", "text": piece}
                                if reply_body.strip():
                                    full_out = finalize_assistant_reply(reply_body)
                                    yield {
                                        "type": "done",
                                        "full_reply": full_out,
                                        "user_message": (req.message or "").strip(),
                                        "new_wake_used": req.session_wake_used
                                        or message_calls_wake_name(req.message),
                                        "orchestra": orch,
                                        "hafiza_dogal": True,
                                    }
                                    return
                except Exception:
                    pass
                try:
                    from ilim_assistant.ana_motor_fast import (
                        iter_fast_direct_llm_reply,
                        should_fast_direct_llm,
                    )

                    if should_fast_direct_llm(
                        req.message, mode_norm, turn_plan
                    ):
                        try:
                            from ilim_assistant.gemini_quota_guard import (
                                gemini_cooldown_active,
                            )

                            st = (
                                "Hızlı yanıt — Gemini kota soğumada; Ollama/Groq…"
                                if gemini_cooldown_active()
                                else "Hızlı yanıt — indeks atlandı (Ollama → Groq → Gemini)…"
                            )
                        except Exception:
                            st = "Hızlı yanıt — indeks atlandı (Ollama/Groq/Gemini)…"
                        yield {"type": "status", "text": st}
                        reply_body = ""
                        for piece in iter_fast_direct_llm_reply(
                            req.message,
                            req.history,
                            mode_norm=mode_norm,
                            question_plan=turn_plan,
                        ):
                            reply_body += piece
                            yield {"type": "token", "text": piece}
                        full_out = finalize_assistant_reply(reply_body)
                        yield {
                            "type": "done",
                            "full_reply": full_out,
                            "user_message": (req.message or "").strip(),
                            "new_wake_used": req.session_wake_used
                            or message_calls_wake_name(req.message),
                            "orchestra": orch,
                            "fast_direct_llm": True,
                        }
                        return
                except Exception:
                    pass
            except Exception:
                turn_plan = None

    agent_context = ""
    workspace_step = None
    retrieval_notes: list[str] = []
    if mode_norm in ("genel", "uretim", "gelisim"):
        _skip_agent = False
        try:
            from ilim_assistant.ana_motor_skip import should_skip_agent_workspace

            _skip_agent = should_skip_agent_workspace(
                req.message,
                mode_norm,
                coding=coding,
                question_plan=turn_plan,
            )
        except Exception:
            _skip_agent = False
        _run_agent = not _skip_agent and (
            turn_plan is None
            or bool(getattr(turn_plan, "use_ilim_rag", True))
            or getattr(turn_plan, "primary", "")
            in (
                "bilgi",
                "bilim",
                "dosya",
                "islem",
                "dilbilgisi",
            )
        )
        if _run_agent:
            try:
                from ilim_assistant.ana_motor_agent import run_agent_workspace_phase

                agent_context, workspace_step, ws_events = run_agent_workspace_phase(
                    req.message,
                    mode_norm,
                    turn_plan,
                    workspace_root=req.workspace_root,
                )
                for ev in ws_events:
                    yield ev
            except Exception:
                agent_context = ""
                workspace_step = None

    reuse_b = None
    _skip_prefetch = mode_norm in ("programlama", "hafiza") or coding or (
        turn_plan is not None
        and hasattr(turn_plan, "use_ilim_rag")
        and not bool(turn_plan.use_ilim_rag)
    )
    try:
        from ilim_assistant.ana_motor_skip import should_skip_stream_prefetch

        if should_skip_stream_prefetch(
            req.message,
            req.history,
            mode_norm,
            coding=coding,
            question_plan=turn_plan,
        ):
            _skip_prefetch = True
    except Exception:
        pass
    try:
        from ilim_assistant.ana_motor_plan import is_casual_conversation_turn

        if is_casual_conversation_turn(req.message, mode_norm, turn_plan):
            _skip_prefetch = True
    except Exception:
        pass
    try:
        from ilim_assistant.ana_motor_fast import (
            fast_paths_enabled,
            should_fast_direct_llm,
        )

        if fast_paths_enabled() and should_fast_direct_llm(
            req.message, mode_norm, turn_plan
        ):
            _skip_prefetch = True
    except Exception:
        pass
    try:
        from ilim_assistant.config import light_chat_mode

        if light_chat_mode():
            _skip_prefetch = True
    except Exception:
        pass
    try:
        from ilim_assistant.ruzgar_umed_cevap_emri import should_skip_fast_bypass_paths

        if (
            should_skip_fast_bypass_paths()
            and mode_norm in ("genel", "uretim", "gelisim")
            and not coding
        ):
            _skip_prefetch = False
    except Exception:
        pass
    if mode_norm == "programlama":
        yield {
            "type": "status",
            "text": "Programlama atölyesi — yerel ilim indeksi atlandı…",
        }
    if (
        mode_norm != "hafiza"
        and os.environ.get("RUZGAR_STREAM_PREFETCH_BUNDLE", "1").strip().lower()
        not in (
            "0",
            "false",
            "no",
        )
        and not _skip_prefetch
    ):
        bundle, evs = prefetch_main_engine_bundle_for_stream(
            req.message,
            req.history,
            mode_raw,
            question_plan=turn_plan,
        )
        for ev in evs:
            txt = str(ev.get("text") or "").strip()
            if txt:
                retrieval_notes.append(txt[:72])
            yield ev
        if evs or bundle.hits or (bundle.ilim_citation_tail or "").strip():
            reuse_b = bundle

    if mode_norm == "programlama":
        try:
            from ilim_assistant.motorlar.programlama_faz11 import merge_orchestra_programlama

            orch = merge_orchestra_programlama(
                orch,
                req.message or "",
                req.workspace_root,
                active_file=req.programlama_active_file,
                phase="llm",
            )
        except Exception:
            pass
    elif mode_norm in ("genel", "uretim", "gelisim"):
        try:
            from ilim_assistant.ana_motor_agent import build_agent_steps

            orch["agent_steps"] = [
                s.to_dict()
                for s in build_agent_steps(
                    turn_plan, workspace_step, retrieval_notes
                )
            ]
        except Exception:
            pass

    _casual_fast = False
    try:
        from ilim_assistant.ana_motor_casual import casual_fast_enabled
        from ilim_assistant.ana_motor_plan import is_casual_conversation_turn

        _casual_fast = (
            casual_fast_enabled()
            and not coding
            and is_casual_conversation_turn(req.message, mode_norm, turn_plan)
        )
    except Exception:
        _casual_fast = False
    try:
        from ilim_assistant.ruzgar_umed_cevap_emri import should_disable_casual_fast_path

        if should_disable_casual_fast_path(req.message or ""):
            _casual_fast = False
    except Exception:
        pass

    if _casual_fast:
        try:
            from ilim_assistant.ana_motor_casual import iter_casual_fast_reply
            from ilim_assistant.llm_brain import free_brain_enabled

            status_txt = (
                "Kısa sohbet — Ollama/Groq/Gemini (ücretsiz zincir)…"
                if free_brain_enabled()
                else "Kısa sohbet — Gemini (hızlı yol)…"
            )
            yield {"type": "status", "text": status_txt}
            reply_body = ""
            for piece in iter_casual_fast_reply(
                req.message,
                req.history,
                mode_norm=mode_norm,
            ):
                reply_body += piece
                yield {"type": "token", "text": piece}
            full_out = finalize_assistant_reply(reply_body)
            yield {
                "type": "done",
                "full_reply": full_out,
                "user_message": (req.message or "").strip(),
                "new_wake_used": req.session_wake_used or message_calls_wake_name(req.message),
                "orchestra": orch,
                "instant_gundelik": True,
                "casual_fast": True,
            }
            return
        except Exception as exc:
            yield {
                "type": "status",
                "text": f"Hızlı sohbet yolu atlandı: {str(exc)[:120]}",
            }

    try:
        from ilim_assistant.ruzgar_umed_cevap_emri import (
            begin_turn_budget,
            deadline_exceeded,
            umed_emri_applies,
            umed_miss_reply,
        )

        if umed_emri_applies(mode_norm=mode_norm, coding_mode=coding):
            try:
                from ilim_assistant.ana_motor_faz59 import begin_umed_turn_budget

                _mf_budget: dict[str, bool] = {}
                try:
                    from ilim_assistant.idrak_entegrasyon import motor_niyeti_heuristic

                    _mf_budget = motor_niyeti_heuristic(req.message or "")
                except Exception:
                    pass
                begin_umed_turn_budget(
                    req.message or "",
                    mode_norm=mode_norm,
                    coding_mode=coding,
                    motor_flags=_mf_budget,
                )
            except Exception:
                begin_turn_budget(req.message or "")
        if umed_emri_applies(mode_norm=mode_norm, coding_mode=coding) and deadline_exceeded():
            from ilim_assistant.ruzgar_egitim import (
                maybe_egitim_learned_reply,
                taught_reply_for_message,
            )

            _mdead = (req.message or "").strip()
            _lated = taught_reply_for_message(_mdead) or maybe_egitim_learned_reply(
                _mdead, history=req.history
            )
            yield from _iter_instant_chat_events(
                _lated or umed_miss_reply(),
                _mdead,
                session_wake_used=req.session_wake_used,
                msg_for_wake=req.message,
                orch=orch,
                egitim_instant=bool(_lated),
            )
            return
    except Exception:
        pass

    prep = prepare_turn(
        req.message,
        req.history,
        req.use_web,
        req.fetch_pages,
        coding,
        req.session_wake_used,
        mode=mode_norm,
        workspace_root=req.workspace_root,
        read_message_links=req.read_message_links,
        reuse_main_engine_bundle=reuse_b,
        orchestration_out=orch,
        question_plan=turn_plan,
        agent_context=agent_context or None,
        pazar_kanallari=req.hizir_channels if mode_norm == "hizir" else None,
    )
    if prep is None:
        yield {"type": "error", "text": "Boş mesaj"}
        return

    msg, hits, user_payload, system, model, og_direct = prep
    if _delegated_from_genel and mode_norm == "programlama":
        try:
            from ilim_assistant.motorlar.programlama_faz79 import build_handoff_packet_v3

            ho = build_handoff_packet_v3(
                msg,
                req.workspace_root,
                active_file=getattr(req, "programlama_active_file", None),
            )
            if ho.get("ok") and ho.get("packet_text"):
                msg = ho["packet_text"] + "\n\n[Kullanıcı isteği]\n" + msg
                yield {
                    "type": "status",
                    "text": "Ana Motor → Programlama: handoff v3 paketi (Faz 79).",
                }
        except Exception:
            pass
    new_wake = req.session_wake_used or message_calls_wake_name(msg)

    if og_direct is not None:
        _instant_g = False
        _instant_c = False
        try:
            from ilim_assistant.ana_motor_plan import maybe_gundelik_instant_reply

            _instant_g = bool(
                maybe_gundelik_instant_reply(
                    msg, mode_norm, motor_flags, question_plan=turn_plan
                )
            )
        except Exception:
            pass
        if turn_plan is not None and getattr(turn_plan, "clarification", None):
            clar = str(turn_plan.clarification or "").strip()
            if clar and str(og_direct or "").strip() == clar:
                _instant_c = True
        yield from _iter_instant_chat_events(
            og_direct,
            msg,
            session_wake_used=req.session_wake_used,
            msg_for_wake=msg,
            orch=orch,
            instant_gundelik=_instant_g,
            instant_clarify=_instant_c,
        )
        return

    prior = prior_messages_for_turn(req.history, mode_norm)
    try:
        brain_sel = select_brain_chain(
            message=msg,
            mode_norm=mode_norm,
            coding_mode=coding,
            question_plan=turn_plan,
            legacy_model=model,
        )
        yield {
            "type": "meta",
            "brain": brain_sel.to_public_dict(),
        }
    except Exception:
        pass
    if mode_norm == "programlama":
        try:
            from ilim_assistant.motorlar.programlama_faz47 import (
                iter_proje_uret_events,
                should_run_proje_uret_pipeline,
            )

            if should_run_proje_uret_pipeline(
                msg,
                mode_norm,
                workspace_root=req.workspace_root,
                active_file=getattr(req, "programlama_active_file", None),
            ):
                yield from iter_proje_uret_events(
                    message=msg,
                    req=req,
                    system=system,
                    user_payload=user_payload,
                    model=model,
                    prior=prior,
                    mode_norm=mode_norm,
                    coding=coding,
                    turn_plan=turn_plan,
                    hits=hits,
                    new_wake=new_wake,
                    orch=orch,
                    delegated_from_genel=_delegated_from_genel,
                )
                return
        except Exception as p47_exc:
            import traceback

            tb = traceback.format_exc(limit=6)
            yield {
                "type": "error",
                "text": (
                    "Bağımsız proje üretimi (Faz 47) hata verdi.\n"
                    f"{str(p47_exc)[:300]}\n{tb[-500:]}"
                ),
            }
            return
        try:
            from ilim_assistant.motorlar.programlama_faz20 import (
                iter_unified_programming_agent_events,
                should_run_unified_programming_agent,
            )

            if should_run_unified_programming_agent(
                msg,
                mode_norm,
                workspace_root=req.workspace_root,
                active_file=getattr(req, "programlama_active_file", None),
            ):
                yield from iter_unified_programming_agent_events(
                    message=msg,
                    req=req,
                    system=system,
                    user_payload=user_payload,
                    model=model,
                    prior=prior,
                    mode_norm=mode_norm,
                    coding=coding,
                    turn_plan=turn_plan,
                    hits=hits,
                    new_wake=new_wake,
                    orch=orch,
                    delegated_from_genel=_delegated_from_genel,
                )
                return
        except Exception as agent20_exc:
            import traceback

            tb = traceback.format_exc(limit=6)
            yield {
                "type": "error",
                "text": (
                    "Birleşik kod ajanı (Faz 20) hata verdi.\n"
                    f"{str(agent20_exc)[:300]}\n{tb[-500:]}"
                ),
            }
            return
        try:
            from ilim_assistant.motorlar.programlama_faz14 import (
                iter_code_agent_turn_events,
                should_run_code_agent_loop,
            )

            if should_run_code_agent_loop(
                msg,
                mode_norm,
                workspace_root=req.workspace_root,
                active_file=getattr(req, "programlama_active_file", None),
            ):
                yield from iter_code_agent_turn_events(
                    message=msg,
                    req=req,
                    system=system,
                    user_payload=user_payload,
                    model=model,
                    prior=prior,
                    mode_norm=mode_norm,
                    coding=coding,
                    turn_plan=turn_plan,
                    hits=hits,
                    new_wake=new_wake,
                    orch=orch,
                    delegated_from_genel=_delegated_from_genel,
                )
                return
        except Exception as agent_exc:
            import traceback

            tb = traceback.format_exc(limit=6)
            yield {
                "type": "error",
                "text": (
                    "Otonom görev (Faz 14) hata verdi — sunucu güncel mi? "
                    "`Ruzgar.ps1 -ForceRestart` dene.\n"
                    f"{str(agent_exc)[:300]}\n{tb[-500:]}"
                ),
            }
            return
    reply_body = ""
    try:
        wants_dbg = mode_norm == "programlama" and (
            _scan_fix_approved or _p11_repair_approved or wants_autonomous_code_debug(msg)
        )
        if mode_norm == "programlama" and not wants_dbg:
            try:
                from ilim_assistant.motorlar.programlama_faz2 import (
                    should_force_autonomous_debug,
                )
                from ilim_assistant.motorlar.programlama_faz96 import (
                    should_force_p11_repair_turn,
                )

                wants_dbg = should_force_autonomous_debug(msg) or should_force_p11_repair_turn(msg)
            except Exception:
                pass
        retry_cap = code_debug_max_retries() if wants_dbg else 0
        extras_remaining = retry_cap
        p3_self_heal_remaining = 0
        if mode_norm == "programlama" and not wants_dbg:
            try:
                p3_self_heal_remaining = max(
                    0, min(int(os.environ.get("RUZGAR_P3_SELF_HEAL_RETRIES", "1")), 2)
                )
            except Exception:
                p3_self_heal_remaining = 1
        active_prior: list = list(prior) if prior else []

        if wants_dbg:
            _dbg_st = (
                "P11 otonom onarım — analiz + patch + pytest (Faz 96)…"
                if _p11_repair_approved
                else (
                    "Onaylı düzeltme — patch + pytest (Faz 2)…"
                    if _scan_fix_approved
                    else "Otomatik kod hata ayıklama (Faz 10.4)…"
                )
            )
            yield {"type": "status", "text": _dbg_st}

        while True:
            round_body = ""
            for piece in stream_chat_with_brain(
                system,
                user_payload,
                model=model,
                prior_messages=active_prior,
                mode_norm=mode_norm,
                coding_mode=coding,
                message=msg,
                question_plan=turn_plan,
            ):
                round_body += piece
                reply_body += piece
                yield {"type": "token", "text": piece}

            if mode_norm == "programlama" and not wants_dbg and p3_self_heal_remaining > 0:
                try:
                    from ilim_assistant.motorlar.programlama_faz92 import (
                        self_heal_retry_prompt,
                        should_self_heal_retry,
                    )

                    if should_self_heal_retry(round_body, mode_norm=mode_norm):
                        p3_self_heal_remaining -= 1
                        yield {
                            "type": "status",
                            "text": "P3 self-heal: ilk deneme toparlanıyor, yeniden üretiyorum…",
                        }
                        retry_user = self_heal_retry_prompt(msg, round_body)
                        active_prior = list(active_prior) + [
                            {"role": "assistant", "content": round_body},
                            {"role": "user", "content": retry_user},
                        ]
                        continue
                except Exception:
                    pass

            if not wants_dbg:
                break

            summ, pytest_rep = apply_assistant_reply_tools(
                round_body,
                req.workspace_root,
                run_pytest=True,
            )
            if pytest_rep is None:
                yield {
                    "type": "status",
                    "text": "Otomatik doğrulama atlandı — workspace kökü yok (LOCAL_TOOLS_ROOT).",
                }
                break
            if pytest_rep.ok:
                yield {
                    "type": "meta",
                    "code_debug": {"phase": "pytest_ok", "exit": pytest_rep.exit_code},
                }
                break
            try:
                from ilim_assistant.ruzgar_debug_report import (
                    build_debug_report,
                    format_debug_report,
                    publish_debug_report,
                )

                dbg_report = build_debug_report(
                    message=msg,
                    workspace_root=req.workspace_root,
                    pytest_report=pytest_rep,
                    writes_count=len(summ.writes),
                    retries_left=extras_remaining,
                )
                publish_debug_report(dbg_report)
                yield {
                    "type": "meta",
                    "code_debug": {"phase": "pytest_failed", **dbg_report},
                }
                yield {
                    "type": "status",
                    "text": "Otonom debug v2 raporu terminale yazıldı.",
                }
                yield {
                    "type": "status",
                    "text": format_debug_report(dbg_report)[:1800],
                }
            except Exception:
                pass
            if not summ.writes:
                yield {
                    "type": "status",
                    "text": "Otomatik debug durdu — model @@write ile dosya yazmadı.",
                }
                break
            if extras_remaining <= 0:
                lim_txt = (
                    "Otomatik tekrar kapalı (RUZGAR_CODE_DEBUG_LOOPS=0) — pytest hâlâ kırmızı."
                    if retry_cap == 0
                    else "Otomatik debug: deneme limiti doldu (RUZGAR_CODE_DEBUG_LOOPS)."
                )
                yield {"type": "status", "text": lim_txt}
                break

            extras_remaining -= 1
            snippet = (pytest_rep.output or "").strip()[:14000]
            fail_msg = (
                "[OTOMATIK DEBUG — Ümit & Gökçenur — Faz 10.6]\n"
                "pytest başarısız. Çıktıyı satır satır oku; **yalnızca gerekli** `@@write yol` + "
                "kod bloğu ile düzelt. Gereksiz refaktör yok; kabuk komutu yok.\n\n"
                f"```text\n{snippet}\n```\n"
            )
            active_prior = list(active_prior) + [
                {"role": "assistant", "content": round_body},
                {"role": "user", "content": fail_msg},
            ]
        if mode_norm == "programlama":
            try:
                from ilim_assistant.motorlar.programlama_faz2 import clear_force_debug_turn
                from ilim_assistant.motorlar.programlama_faz96 import (
                    clear_force_p11_repair_turn,
                )

                clear_force_debug_turn()
                clear_force_p11_repair_turn()
            except Exception:
                pass
        footer = rag_footer(hits)
        body_fixed = finalize_assistant_reply(reply_body)
        code_patch_meta: dict[str, Any] = {}
        if mode_norm == "programlama":
            try:
                from ilim_assistant.motorlar.programlama_faz10 import (
                    process_assistant_reply_patches,
                    resolve_scope_rel,
                )

                scope = resolve_scope_rel(
                    req.workspace_root,
                    active_file=req.programlama_active_file,
                )
                code_patch_meta = process_assistant_reply_patches(
                    reply_body,
                    req.workspace_root,
                    scope_rel=scope,
                    skip_if_debug_loop=bool(wants_dbg),
                )
                patch_footer = str(code_patch_meta.get("footer") or "")
                if patch_footer:
                    body_fixed = body_fixed.rstrip() + patch_footer
            except Exception:
                pass
        try:
            from ilim_assistant.ana_motor_reflection import apply_answer_quality_pass

            web_used = bool(
                turn_plan is not None
                and getattr(turn_plan, "prefer_web", False)
                and os.environ.get("ENABLE_WEB_SEARCH", "1") == "1"
            )
            body_fixed = apply_answer_quality_pass(
                body_fixed,
                msg,
                hits=hits,
                question_plan=turn_plan,
                web_was_used=web_used,
            )
        except Exception:
            pass
        full_out = body_fixed + footer
        done_llm: dict[str, Any] = {
            "type": "done",
            "full_reply": full_out,
            "user_message": msg,
            "new_wake_used": new_wake,
        }
        if orch:
            try:
                from ilim_assistant.ana_motor_agent import mark_agent_answer_done

                if orch.get("agent_steps"):
                    orch["agent_steps"] = mark_agent_answer_done(orch["agent_steps"])
            except Exception:
                pass
            done_llm["orchestra"] = orch
        if code_patch_meta and code_patch_meta.get("action") not in ("skip", "none"):
            try:
                from ilim_assistant.motorlar.programlama_faz16 import enrich_code_patch_meta

                code_patch_meta = enrich_code_patch_meta(
                    code_patch_meta, req.workspace_root
                )
            except Exception:
                pass
            done_llm["code_patch"] = {
                "action": code_patch_meta.get("action"),
                "applied": list(code_patch_meta.get("applied") or []),
                "errors": list(code_patch_meta.get("errors") or []),
                "items": list(code_patch_meta.get("items") or []),
                "counts": code_patch_meta.get("counts") or {},
                "count": code_patch_meta.get("count"),
            }
        if mode_norm == "programlama":
            try:
                from ilim_assistant.motorlar.programlama_faz11 import merge_orchestra_programlama

                orch = merge_orchestra_programlama(
                    orch,
                    msg,
                    req.workspace_root,
                    active_file=req.programlama_active_file,
                    phase="done",
                    patch_meta=code_patch_meta,
                )
                done_llm["orchestra"] = orch
            except Exception:
                pass
        if _delegated_from_genel:
            done_llm["programlama_delegated"] = True
            try:
                from ilim_assistant.ana_motor_faz59 import (
                    append_delegation_footer_to_reply,
                    format_delegation_summary_text,
                    summary_from_agent_state,
                )

                _elapsed = float(done_llm.get("elapsed_sec") or 0)
                _summ59 = summary_from_agent_state(
                    req.workspace_root,
                    success=bool(code_patch_meta.get("applied"))
                    or "tamamlandı" in full_out.lower(),
                    turns_used=0,
                    elapsed_sec=_elapsed,
                )
                done_llm["delegate_summary"] = _summ59
                done_llm["delegate_summary_text"] = format_delegation_summary_text(
                    _summ59
                )
                full_out = append_delegation_footer_to_reply(
                    full_out,
                    req.workspace_root,
                    success=bool(_summ59.get("success")),
                    turns_used=int(_summ59.get("turns_used") or 0),
                    elapsed_sec=_elapsed,
                    scope_rel=str(_summ59.get("scope_rel") or ""),
                    goal=str(_summ59.get("goal") or ""),
                )
                done_llm["full_reply"] = full_out
            except Exception:
                pass
        if mode_norm == "programlama":
            try:
                from ilim_assistant.motorlar.programlama_faz93 import (
                    large_change_status_text,
                    record_refactor_checkpoint,
                    render_rollback_directive,
                    rollback_status_text,
                    summarize_large_change,
                )
                from ilim_assistant.motorlar.programlama_faz94 import (
                    build_regression_shield,
                    regression_status_text,
                    render_test_directive,
                    run_regression_shield_check,
                )
                from ilim_assistant.motorlar.programlama_faz10 import resolve_scope_rel

                _lc = summarize_large_change(
                    code_patch_meta if isinstance(code_patch_meta, dict) else None
                )
                done_llm["large_change_v96"] = _lc
                if bool(_lc.get("is_large")):
                    yield {"type": "status", "text": large_change_status_text(_lc)}

                _cp = record_refactor_checkpoint(
                    req.workspace_root,
                    user_message=msg,
                    patch_meta=code_patch_meta if isinstance(code_patch_meta, dict) else None,
                )
                if _cp.get("recorded"):
                    done_llm["refactor_checkpoint_v96"] = _cp
                    _rb_txt = rollback_status_text(_cp.get("rollback") or {})
                    if _rb_txt:
                        yield {"type": "status", "text": _rb_txt}
                    _rb_dir = render_rollback_directive(_cp.get("rollback") or {}).strip()
                    if _rb_dir:
                        done_llm["rollback_directive_v96"] = _rb_dir

                _applied = list(
                    (code_patch_meta or {}).get("applied") or []
                    if isinstance(code_patch_meta, dict)
                    else []
                )
                if _applied:
                    _p_scope_done = resolve_scope_rel(
                        req.workspace_root,
                        active_file=req.programlama_active_file,
                    )
                    _shield = run_regression_shield_check(
                        req.workspace_root,
                        _p_scope_done,
                        _applied,
                    )
                    done_llm["regression_shield_v97"] = _shield
                    _p7_dir = render_test_directive(_shield).strip()
                    if _p7_dir and not bool(_shield.get("verify_ok")):
                        done_llm["regression_test_directive_v97"] = _p7_dir
                    yield {"type": "status", "text": regression_status_text(_shield)}
                elif isinstance(code_patch_meta, dict) and code_patch_meta.get("applied") is not None:
                    _p_scope_done = resolve_scope_rel(
                        req.workspace_root,
                        active_file=req.programlama_active_file,
                    )
                    _shield = build_regression_shield([], req.workspace_root, _p_scope_done)
                    done_llm["regression_shield_v97"] = _shield
            except Exception:
                pass
            if _p11_repair_approved:
                try:
                    from ilim_assistant.motorlar.programlama_faz96 import (
                        format_verification_report,
                        run_post_repair_verification,
                    )

                    _p11_verify = run_post_repair_verification(
                        req.workspace_root,
                        before=_p11_analysis_before,
                    )
                    done_llm["p11_verify_v106"] = _p11_verify
                    _vtxt = format_verification_report(_p11_verify)
                    if _vtxt:
                        yield {"type": "status", "text": _vtxt}
                except Exception:
                    pass
        if mode_norm == "programlama":
            try:
                from ilim_assistant.motorlar.programlama_faz93 import record_decision

                record_decision(
                    req.workspace_root,
                    user_message=msg,
                    assistant_reply=full_out,
                    patch_meta=code_patch_meta if isinstance(code_patch_meta, dict) else None,
                )
            except Exception:
                pass
        yield done_llm
    except Exception as e:
        yield {"type": "error", "text": format_llm_user_error(e)}


def iter_chat_turn_events(req: ChatRequest) -> Iterator[dict]:
    """Tur süresi + eğitim sonrası işlemler (oturum özeti, bulamadım)."""
    import time

    t0 = time.perf_counter()
    for obj in _iter_chat_turn_events_impl(req):
        if obj.get("type") == "done":
            obj = dict(obj)
            obj.setdefault("elapsed_sec", time.perf_counter() - t0)
            try:
                from ilim_assistant.ruzgar_egitim import on_chat_turn_done

                obj = on_chat_turn_done(req, obj)
            except Exception:
                pass
            try:
                _mode_done = normalize_mode(_effective_chat_mode_raw(req))
                if _mode_done == "programlama":
                    from ilim_assistant.motorlar.programlama_faz5 import record_chat_turn

                    cd = obj.get("code_debug")
                    pytest_ok = None
                    pytest_exit = None
                    if isinstance(cd, dict):
                        if cd.get("phase") == "pytest_ok":
                            pytest_ok = True
                            pytest_exit = cd.get("exit")
                        elif cd.get("phase") == "pytest_failed":
                            pytest_ok = False
                            pytest_exit = cd.get("exit")
                    record_chat_turn(
                        req.workspace_root,
                        user_message=req.message or "",
                        assistant_reply=str(obj.get("full_reply") or ""),
                        pytest_ok=pytest_ok,
                        pytest_exit=pytest_exit,
                        active_file=req.programlama_active_file,
                    )
                    if _mode_done == "programlama":
                        from ilim_assistant.motorlar.programlama_faz5 import (
                            sync_client_editor_state,
                        )

                        sync_client_editor_state(
                            req.workspace_root,
                            active_file=req.programlama_active_file,
                            editor_snippet=req.programlama_editor_snippet,
                            language=req.programlama_language,
                        )
            except Exception:
                pass
        yield obj


@app.post("/api/chat/stream")
def chat_stream(req: ChatRequest):
    def generate():
        yield ":ruzgar-events\n\n"
        for obj in iter_chat_turn_events(req):
            yield _sse(obj)

    return StreamingResponse(
        generate(),
        media_type="text/event-stream; charset=utf-8",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
    )


@app.post("/api/chat/full")
def chat_full(req: ChatRequest):
    """Streaming kapalı UI yolu: tüm cevabı üret, tek JSON yanıt olarak döndür."""
    events: list[dict[str, Any]] = []
    full_reply = ""
    done_event: dict[str, Any] | None = None
    for obj in iter_chat_turn_events(req):
        typ = obj.get("type")
        if typ in ("meta", "status"):
            events.append(obj)
        elif typ == "token":
            full_reply += str(obj.get("text") or "")
        elif typ == "done":
            done_event = obj
            full_reply = str(obj.get("full_reply") or full_reply)
            break
        elif typ == "error":
            return {
                "ok": False,
                "error": str(obj.get("text") or "Bilinmeyen hata"),
                "events": events,
            }

    if done_event is None:
        full_reply = finalize_assistant_reply(full_reply)
        new_wake = req.session_wake_used or message_calls_wake_name(req.message)
        done_event = {
            "type": "done",
            "full_reply": full_reply,
            "user_message": (req.message or "").strip(),
            "new_wake_used": new_wake,
        }

    return {
        "ok": True,
        "events": events[-24:],
        "full_reply": full_reply,
        "user_message": done_event.get("user_message", (req.message or "").strip()),
        "new_wake_used": bool(done_event.get("new_wake_used")),
        "orchestra": done_event.get("orchestra"),
        "instant_gundelik": bool(done_event.get("instant_gundelik")),
        "instant_clarify": bool(done_event.get("instant_clarify")),
        "instant_memory": bool(done_event.get("instant_memory")),
    }


@app.websocket("/ws/chat")
async def websocket_chat(ws: WebSocket) -> None:
    """Electron Işık Hızı: SSE ile aynı `iter_chat_turn_events` (önce öğrenme merkezi)."""
    await ws.accept()
    try:
        payload = await ws.receive_json()
    except Exception:
        try:
            await ws.send_json({"type": "error", "text": "Geçersiz veya boş istek"})
        except Exception:
            pass
        await ws.close()
        return
    try:
        req = ChatRequest(**payload)
    except Exception as e:
        try:
            await ws.send_json({"type": "error", "text": f"İstek şeması: {e}"})
        except Exception:
            pass
        await ws.close()
        return
    try:
        async for obj in iterate_in_threadpool(iter_chat_turn_events(req)):
            await ws.send_json(obj)
    except Exception:
        pass
    try:
        await ws.close()
    except Exception:
        pass


if __name__ == "__main__":
    from ilim_assistant.ruzgar_api_port import resolve_api_port

    port = resolve_api_port()
    _default_host = "0.0.0.0" if os.environ.get("PORT") else "127.0.0.1"
    host = os.environ.get("RUZGAR_API_HOST", _default_host)
    # Varsayılan: daha az log I/O (yüksek trafikli token akışı için)
    log_level = os.environ.get("RUZGAR_UVICORN_LOG", "warning")
    access_log = os.environ.get("RUZGAR_ACCESS_LOG", "0").strip() in ("1", "true", "yes")
    uvicorn.run(
        app,
        host=host,
        port=port,
        log_level=log_level,
        access_log=access_log,
    )
