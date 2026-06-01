# Created by Ümit & Gökçenur
"""Tercüme Faz 14G — DOCX çıktı (python-docx)."""

from __future__ import annotations

import re
from typing import Any

DOCX_EXPORT_VERSION = "tercume-export-docx-v17g-2026-05-29"


def docx_available() -> bool:
    try:
        import docx  # type: ignore  # noqa: F401

        return True
    except ImportError:
        return False


def _title_from_rel(source_rel: str) -> str:
    raw = (source_rel or "").strip().replace("\\", "/")
    leaf = raw.split("/")[-1] if raw else ""
    stem = re.sub(r"\.[^.]+$", "", leaf) or "Çeviri"
    return stem.replace("_", " ").strip() or "Çeviri"


def build_docx_bytes(
    text: str,
    *,
    title: str = "",
    source_rel: str = "",
    tgt_lang: str = "tr",
) -> bytes:
    if not docx_available():
        raise RuntimeError("python-docx yok. pip install python-docx")
    from docx import Document  # type: ignore
    from docx.shared import Pt  # type: ignore

    doc = Document()
    doc_title = (title or "").strip() or _title_from_rel(source_rel)
    doc.add_heading(doc_title, level=0)
    if source_rel:
        meta = doc.add_paragraph()
        meta.add_run(f"Kaynak: {source_rel} · Hedef dil: {tgt_lang}").italic = True
    body = (text or "").strip()
    paragraphs = [p.strip() for p in re.split(r"\n\s*\n", body) if p.strip()]
    if not paragraphs and body:
        paragraphs = [body]
    for para in paragraphs:
        if len(para) < 96 and not para.endswith((".", "?", "!", ":", "؛")):
            doc.add_heading(para, level=2)
        else:
            p = doc.add_paragraph(para)
            for run in p.runs:
                run.font.size = Pt(11)
    import io

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def export_docx_meta() -> dict[str, Any]:
    return {
        "version": DOCX_EXPORT_VERSION,
        "available": docx_available(),
        "format": "docx",
    }
